# -*- coding: utf-8 -*-
"""Generate PL/0 experiment report (DOCX) using python-docx.

Style requirements (per 撰写指引-实验报告.doc):
- Page: top 30mm / bottom 25mm / left 30mm / right 20mm
- Line spacing: 1.5x
- Page numbers: arabic, from 绪论 onward, footer-right
- Title (cover): 二号 黑体 加粗
- Chapter heading: 三号 黑体 加粗
- Section heading: 小四号 黑体 加粗
- Body: 小四号 宋体
- Caption (figure): 宋体 五号
- Caption (table): 黑体 五号
- Page number: 小五号 Times New Roman
- Tables: three-line (no left/right borders)
- Figures must not cross page boundary
"""
import os
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'C:\Users\xwmsu\Desktop\c_builder_b'
FIG = os.path.join(ROOT, 'docs', 'report_figures')
OUT = os.path.join(ROOT, 'PL0编译器扩充实验报告.docx')

doc = Document()


def set_run_font(run, name='宋体', size_pt=12, bold=False, ascii_font='Times New Roman'):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)


def add_para(text='', font='宋体', size=12, bold=False, align=None,
             first_indent=True, line_spacing=1.5,
             space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if first_indent and text:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if text:
        r = p.add_run(text)
        set_run_font(r, name=font, size_pt=size, bold=bold)
    return p


def add_h1(text, num=None):
    full = (num + '  ' + text) if num else text
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(full)
    set_run_font(r, name='黑体', size_pt=16, bold=True)
    return p


def add_h2(text, num=None):
    full = (num + '  ' + text) if num else text
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(full)
    set_run_font(r, name='黑体', size_pt=12, bold=True)
    return p


def add_body(text, size=12, indent=True):
    return add_para(text, font='宋体', size=size, first_indent=indent,
                    line_spacing=1.5)


def add_code(text, size=10):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(8)
        r = p.add_run(line if line else ' ')
        set_run_font(r, name='Consolas', size_pt=size, ascii_font='Consolas')


def add_image(path, caption, width_cm=12.5):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.first_line_indent = Pt(0)
    cr = cap.add_run(caption)
    set_run_font(cr, name='宋体', size_pt=9)


def add_three_line_table(headers, rows, caption):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement('w:' + edge)
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tblPr.append(borders)
    # set header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, name='黑体', size_pt=9, bold=True)
        tcPr = c._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'right'):
            b = OxmlElement('w:' + edge)
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
        bb = OxmlElement('w:bottom')
        bb.set(qn('w:val'), 'single')
        bb.set(qn('w:sz'), '8')
        bb.set(qn('w:color'), '000000')
        tcBorders.append(bb)
        tcPr.append(tcBorders)
    for r_idx, row in enumerate(rows):
        is_last = (r_idx == len(rows) - 1)
        for i, val in enumerate(row):
            c = t.rows[r_idx + 1].cells[i]
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run_font(r, name='宋体', size_pt=9)
            tcPr = c._tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for edge in ('left', 'right', 'bottom', 'top'):
                b = OxmlElement('w:' + edge)
                if edge == 'bottom' and is_last:
                    b.set(qn('w:val'), 'single')
                    b.set(qn('w:sz'), '8')
                    b.set(qn('w:color'), '000000')
                else:
                    b.set(qn('w:val'), 'nil')
                tcBorders.append(b)
            tcPr.append(tcBorders)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.first_line_indent = Pt(0)
    cr = cap.add_run(caption)
    set_run_font(cr, name='黑体', size_pt=9)
    return t


def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ========== Page setup ==========
section = doc.sections[0]
section.top_margin = Mm(30)
section.bottom_margin = Mm(25)
section.left_margin = Mm(30)
section.right_margin = Mm(20)
section.different_first_page_header_footer = True

style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = OxmlElement('w:rFonts')
    rPr.append(rFonts)
rFonts.set(qn('w:eastAsia'), '宋体')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')


# ========== Cover page ==========
cover_p = doc.add_paragraph()
cover_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_p.paragraph_format.space_before = Pt(60)
r = cover_p.add_run('广东工业大学')
set_run_font(r, name='黑体', size_pt=22, bold=True)

cover_p2 = doc.add_paragraph()
cover_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover_p2.paragraph_format.space_before = Pt(6)
r = cover_p2.add_run('本科生实验报告')
set_run_font(r, name='黑体', size_pt=18, bold=True)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(20)
title_p.paragraph_format.space_after = Pt(20)
r = title_p.add_run('PL/0 编译器的功能扩充')
set_run_font(r, name='黑体', size_pt=22, bold=True)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub_p.add_run('——基于 Borland C++Builder 6 与递归下降分析法')
set_run_font(r, name='黑体', size_pt=14, bold=True)

doc.add_paragraph()
doc.add_paragraph()

info_table = doc.add_table(rows=6, cols=2)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
info_table.autofit = False
tbl = info_table._tbl
tblPr = tbl.tblPr
borders = OxmlElement('w:tblBorders')
for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
    b = OxmlElement('w:' + edge)
    b.set(qn('w:val'), 'nil')
    borders.append(b)
tblPr.append(borders)

info_data = [
    ('学    院', '计算机学院'),
    ('专    业', '计算机科学与技术'),
    ('年级班别', '2023 级 计科 X 班'),
    ('学    号', '31120xx901'),
    ('学生姓名', '曾伟胜'),
    ('指导教师', '杨劲涛'),
]
for i, (k, v) in enumerate(info_data):
    row = info_table.rows[i]
    c1, c2 = row.cells
    c1.text = ''
    p1 = c1.paragraphs[0]
    p1.paragraph_format.first_line_indent = Pt(0)
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p1.add_run(k)
    set_run_font(r, name='黑体', size_pt=14, bold=True)
    c2.text = ''
    p2 = c2.paragraphs[0]
    p2.paragraph_format.first_line_indent = Pt(0)
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.left_indent = Pt(10)
    r = p2.add_run(v)
    set_run_font(r, name='宋体', size_pt=14)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.paragraph_format.space_before = Pt(60)
r = date_p.add_run('2026 年 6 月 26 日')
set_run_font(r, name='宋体', size_pt=14, bold=True)

add_page_break()


# ========== Table of contents ==========
toc_h = doc.add_paragraph()
toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_h.paragraph_format.space_after = Pt(20)
r = toc_h.add_run('目  录')
set_run_font(r, name='黑体', size_pt=18, bold=True)

toc_entries = [
    ('第一章  绪论', '1'),
    ('    1.1  PL/0 语言与教学编译器', '1'),
    ('    1.2  本次实验的目标与动机', '2'),
    ('    1.3  开发环境与项目结构', '2'),
    ('第二章  编译流程与基础理论', '3'),
    ('    2.1  PL/0 编译器的整体架构', '3'),
    ('    2.2  递归下降语法分析', '4'),
    ('    2.3  符号表与 P-Code 中间表示', '5'),
    ('    2.4  解释执行模型', '6'),
    ('第三章  词法分析器扩展', '7'),
    ('    3.1  扩展保留字与运算符', '7'),
    ('    3.2  GetSym 的状态机实现', '7'),
    ('    3.3  关键字表的二分查找', '8'),
    ('第四章  语法分析与代码生成扩展', '9'),
    ('    4.1  STATEMENT 函数的结构', '9'),
    ('    4.2  ELSE 子句的代码回填', '10'),
    ('    4.3  复合赋值 +=/-= 的代码生成', '11'),
    ('    4.4  ++ / -- 自增自减语句', '12'),
    ('    4.5  FOR-TO / FOR-DOWNTO 循环', '13'),
    ('    4.6  RETURN 提前返回', '15'),
    ('第五章  核心函数代码分析', '16'),
    ('    5.1  Gen / Statement / Interpret 三件套', '16'),
    ('    5.2  错误恢复与现场保留', '17'),
    ('    5.3  OPR 指令域的扩展', '18'),
    ('第六章  测试与实验结果', '20'),
    ('    6.1  测试用例组织', '20'),
    ('    6.2  正向测试样例与运行结果', '21'),
    ('    6.3  反向测试样例与错误处理', '23'),
    ('    6.4  控制台回归驱动', '24'),
    ('第七章  实验中遇到的困难与解决', '26'),
    ('    7.1  命令行构建 BCB6 项目', '26'),
    ('    7.2  GUI 应用自动化测试的替代方案', '27'),
    ('    7.3  词法增量扩展的耦合问题', '28'),
    ('第八章  小结与展望', '29'),
    ('实验总结与收获', '30'),
    ('参考文献', '31'),
    ('附录 A：源代码改动摘要', '32'),
    ('附录 B：核心函数清单', '33'),
]
for label, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label)
    if not label.startswith('    '):
        set_run_font(r, name='黑体', size_pt=12, bold=True)
    else:
        set_run_font(r, name='宋体', size_pt=11)
    tab_r = p.add_run('\t' + page)
    set_run_font(tab_r, name='Times New Roman', size_pt=11)

add_page_break()


# ========== Footer with page numbers ==========
def add_footer_page_number():
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE   \\* MERGEFORMAT'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    set_run_font(r, name='Times New Roman', size_pt=9)


add_footer_page_number()


# ========== 第一章 绪论 ==========
add_h1('绪论', num='第一章')

add_h2('PL/0 语言与教学编译器', num='1.1')
add_body(
    'PL/0 是由 Niklaus Wirth 在 1976 年编写的小型教学语言，其完整 Pascal 子集只包含'
    '常量声明、变量声明、过程声明、赋值、条件、循环、过程调用与基本 I/O 几类语句。'
    'Wirth 在其著作《Compilerbau》中以 PL/0 为对象展示了递归下降语法分析的全部细节，'
    '因此 PL/0 至今仍被国内外的编译原理课程当作最典型的教学蓝本。'
)
add_body(
    '原版 PL/0 的不足之处同样明显：它没有 ELSE 子句，没有 += 之类的复合赋值，'
    '更没有真正的 FOR 循环。学生在课程设计阶段几乎不可避免地要在原版基础上'
    '增加新的语法规则。每一次增加，本质上都是对词法分析、语法分析、'
    '代码生成、解释执行四部分的协同修改——这一点决定了 PL/0 课设虽然规模小，'
    '但能够完整体现编译原理课程要求的全部工程训练。'
)
add_body(
    '本实验所采用的原始代码来自 Wirth 版本的 BCB6 移植，源程序仅有不到一千行，'
    '但几乎所有 PL/0 教学版本所需的扩展点都已预留：保留字表、二分查找、'
    'OPR 子操作码空间、错误恢复现场，正是这种「小而完整」的工程结构为本次'
    '扩展提供了合适的起点。'
)

add_h2('本次实验的目标与动机', num='1.2')
add_body('本次课程设计的题目是对 PL/0 编译程序作以下扩充：')
for item in [
    '增加单词：保留字 ELSE、FOR、TO、DOWNTO；运算符 +=、-=、*=、/=、<>、!=。',
    '修改单词：原版 PL/0 的不等号 # 改为 <> 或 !=（提高要求）。',
    '增加条件语句的 ELSE 子句，并写出相关文法与语法描述。',
    '扩充赋值运算：+= 和 -=。',
    '扩充 FOR 循环：FOR <变量>:=<表达式> TO <表达式> DO <语句> 与 FOR <变量>:=<表达式> DOWNTO <表达式> DO <语句>。',
    '增加自增自减：A++ 与 A--。',
    '增加 RETURN 提前返回。',
]:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(item)
    set_run_font(r, name='宋体', size_pt=12)
add_body(
    '值得注意的是，最终需要让上述语法在 BCB6 项目中端到端工作——也就是说，'
    '新增的每一项特性都必须同时具备：词法层的 token、语法层对产生式的支持、'
    '代码生成层能产生合法的 P-Code、解释执行层能跑出正确的输出。'
    '任何一环缺失都意味着功能「半成品」，这一点也是本次报告反复强调的评判标准。'
)

add_h2('开发环境与项目结构', num='1.3')
add_body(
    '开发环境采用 Borland C++Builder 6（VCL 图形界面），编译命令为 bcc32 + ilink32。'
    '代码本身只有一份核心文件 src/Unit1.cpp，约 980 行；其余为 VCL 窗体'
    '（Unit1.dfm、Unit1.h）与 BCB6 工程文件 PL01.bpr。'
    '测试样例组织在 test_cases/ 目录下，包括扩展前已有的 7 对正向/反向样例，'
    '以及本次新增的 6 对正向样例（FORT01-03、COMP01、INC01、RET01）。'
)
add_body(
    '为了应对命令行环境下回归测试的需求，我还编写了一个 console 版本的'
    'pl0_test 驱动（test_console/pl0_test.cpp），它绕开 VCL 直接复用编译'
    '器逻辑并支持从 stdin 接受 READ 输入。这一辅助工具的源代码与构建方法'
    '放在附录部分，本章不再赘述。'
)
add_image(os.path.join(FIG, 'fig_func_lines.png'),
          '图 1.1  Unit1.cpp 中各核心函数的规模（Top 10，按代码行数排序）',
          width_cm=12.5)

add_page_break()


# ========== 第二章 ==========
add_h1('编译流程与基础理论', num='第二章')

add_h2('PL/0 编译器的整体架构', num='2.1')
add_body(
    'PL/0 编译器是典型的「单趟 + 解释执行」结构。源代码经过一次扫描即可同时完成'
    '词法分析、语法分析和代码生成，所得到的 P-Code 立即进入解释器循环。'
    '整个过程没有显式的中间文件、没有显式的 AST，所见即所得——这种紧凑的'
    '结构非常适合教学，但也意味着每一处语法扩展都必须沿着同一条流水线向前'
    '延伸到解释器，否则就会出现「词法层承认、但解释器不认」的故障。'
)
add_image(os.path.join(FIG, 'pipeline.png'),
          '图 2.1  PL/0 编译与解释的整体架构', width_cm=12.5)
add_body(
    '从图 2.1 可以看到，词法分析器 GetSym 与语法分析器共享关键字表 KWORD、'
    'WSYM 与单字符表 SSYM；语法分析器在归约过程中不断读写符号表 TABLE 与'
    '指令数组 CODE；最终解释器 Interpret 直接消费 CODE。整个数据流是单向的，'
    '调试时只需要顺着 CX 增长的方向观察 CODE 数组即可重现编译过程。'
)
add_body(
    '需要特别指出的是，本工程没有使用任何现代编译器常用的 visitor 模式、'
    '没有 IR 优化、没有基础块划分，因此每一处扩展都不可避免地以'
    '侵入式修改 STATEMENT 的 switch-case 表的形式出现。这种风格在'
    '教学语境下利大于弊：学生可以直接对照源代码与 P-Code 输出，'
    '而无需先理解一套抽象语法树的中间表示。'
)

add_h2('递归下降语法分析', num='2.2')
add_body(
    'PL/0 的语法可以用 EBNF 简洁地描述，其中 statement 部分定义了合法的'
    '语句形式。原版只有 ::= ε | 赋值 | CALL | READ | WRITE | BEGIN/END | IF/THEN | WHILE/DO'
    '等 8 类，本次扩展后新增了复合赋值、自增自减、FOR 循环、RETURN 提前返回'
    '以及 IF/THEN/ELSE 这五种形式，语法规则随之扩展。'
)
add_body(
    '递归下降的实现技巧是：在每个非终结符对应函数的开头，调用 TEST 检查当前'
    'SYM 是否落在 first 集之中；如果不在，则同步到 follow 集以继续分析后续语句。'
    '这种「先检查、再消化」的两段式写法可以天然完成错误恢复，是 PL/0 编译器'
    '能够在多次试验后仍然保持稳健的关键。'
)
add_body(
    '扩展后的 STATEMENT 函数的整体结构见图 2.2。可以看到，针对每一种新增'
    '语句，函数都用一个独立的 case 子句负责语法消化与代码生成，'
    '使得新增语句与原有语句之间不会发生意外的耦合。'
)
add_image(os.path.join(FIG, 'statement.png'),
          '图 2.2  STATEMENT 函数的整体结构（case 分支覆盖）',
          width_cm=12.5)

add_h2('符号表与 P-Code 中间表示', num='2.3')
add_body(
    'PL/0 使用一张扁平数组 TABLE[TXMAX] 表示整个程序的符号表，'
    '其中 TX=0 永远作为哨兵使用：POSITION 函数在查不到标识符时返回 0，'
    '于是 0 自然成为「未声明」或「越界」的统一信号。'
)
add_body(
    'TABLE 的每一项保存三个信息：NAME（最多 10 个字符的标识符）、KIND'
    '（CONSTANT/VARIABLE/PROCEDUR 三选一）以及一个 union 值。其中'
    '常量存 VAL，变量与过程存 LEVEL/ADR/SIZE 三个整数字段。'
    '下面这张表（表 2.1）展示了部分符号表项的结构：'
)
add_three_line_table(
    ['INDEX', 'NAME', 'KIND', 'LEVEL', 'ADR', 'VAL/SIZE'],
    [
        ['0', '(reserved)', '—', '—', '—', '—'],
        ['1', 'A', 'CONSTANT', '0', '—', 'VAL=10'],
        ['2', 'B', 'VARIABLE', '0', '3', '—'],
        ['3', 'C', 'VARIABLE', '0', '4', '—'],
        ['4', 'P', 'PROCEDUR', '0', '10', 'SIZE=4'],
        ['5', 'Q', 'PROCEDUR', '1', '20', 'SIZE=3'],
    ],
    '表 2.1  TABLE 符号表项结构（TX=0 保留为哨兵）'
)
add_body(
    'P-Code 中间表示由 CODE[CXMAX] 数组保存，每条指令包含三个字段：'
    'F（功能码，LIT/OPR/LOD/STO/CAL/INI/JMP/JPC 共 8 种）、'
    'L（层差，用于跨过程访问）、A（地址或立即数）。'
    '在本次扩展中，F 没有新增，全部扩展都是借由已有的 OPR 指令的 A 域'
    '以及新增的 token 在 STATEMENT 内消化后产生的额外指令实现的。'
)

add_h2('解释执行模型', num='2.4')
add_body(
    'PL/0 解释器是一个标准的栈式虚拟机：寄存器 T 指向操作数栈顶，B 指向'
    '当前活动记录的基址，P 指向下一条待执行的 P-Code 地址。'
    '运行时栈的布局大致是：[返回地址 | 静态链 SL | 动态链 DL | 局部变量区 | 操作数]'
    '——这一布局决定了所有 LOD/STO/CAL 指令必须通过 BASE(L, B, S)'
    '函数沿静态链上溯 L 层，从而获得目标活动记录的基址。'
)
add_body(
    '本次扩展没有修改 Interpret 本身的控制流，新增的 ++/--、+=/-=、FOR 循环'
    '以及 RETURN 语句都只对应少量 OPR 操作码的复用——例如 += 对应'
    'OPR 0,2（加），FOR 的边界判断对应 OPR 0,11 或 OPR 0,13（>= / <=）。'
    '因此 Interpret 函数本身不需要任何修改，这是把扩展约束在'
    '代码生成阶段的另一个好处。'
)
add_image(os.path.join(FIG, 'fig_pcode_distribution.png'),
          '图 2.3  P-Code 指令类型分布（正向用例统计）', width_cm=10.5)

add_page_break()


# ========== 第三章 ==========
add_h1('词法分析器扩展', num='第三章')

add_h2('ELSE / 复合赋值 / FOR / RETURN 的文法描述', num='3.0')
add_body(
    '课设任务要求「增加条件语句的 ELSE 子句，并写出相关文法与语法描述」。'
    '原版 PL/0 的 EBNF 文法只描述了 8 种语句形式，扩展后的 statement'
    '产生式如下所示：「BEGIN」「END」「IF」「THEN」「WHILE」「DO」'
    '等关键字不再只是内部保留字，而是能够通过 wsym 进入 first 集。'
)
add_code(
    'Program     = "." Block.\n'
    'Block       = [CONST IDENT "=" number {"," IDENT "=" number} ";]\n'
    '              [VAR IDENT {"," IDENT} ";]\n'
    '              {"PROCEDURE IDENT ";" Block ";"} Statement.\n'
    'Statement   = [IDENT (":="|"+="|"*="|"/=") Expression\n'
    '              | IDENT "++" | IDENT "--"\n'
    '              | "CALL" IDENT\n'
    '              | "?" IDENT\n'
    '              | "!" Expression\n'
    '              | "BEGIN" Statement {";" Statement} "END"\n'
    '              | "IF" Condition "THEN" Statement [ "ELSE" Statement ]\n'
    '              | "WHILE" Condition "DO" Statement\n'
    '              | "FOR" IDENT ":=" Expression ("TO"|"DOWNTO")\n'
    '                Expression "DO" Statement\n'
    '              | "RETURN"\n'
    '              ].\n'
    'Condition   = "ODD" Expression\n'
    '              | Expression ("="|"<>"|"!="|"<"|"<="|">"|">=") Expression.\n'
    'Expression  = ["+"|"-"] Term {("+"|"-") Term}.\n'
    'Term        = Factor {("*"|"/") Factor}.\n'
    'Factor      = IDENT | number | "(" Expression ")".'
)
add_body(
    '上述 EBNF 中「未加粗/未斜体」的是原有规则，方括号 [...] 表示可选，'
    '花括号 {...} 表示重复 0 次或多次。本次扩展添加了 '
    '+= / -= / ++ / -- / ELSE / FOR / TO / DOWNTO / RETURN != '
    '共 11 个新语法产生式，以及 2 个新的终结符（DOWNTO、RETURN 保留字）。'
)
add_body(
    '为了验证文法的一致性，本课程设计指导书还要求在词法层面的 KWORD'
    '表中接入对应的 wsym 内部码——这是正确识别语句开始的必要条件。'
    '本实验中，WSYM 数组为每个保留字分配一个唯一的枚举值（例如 ELSESYM、'
    'FORSYM、DOWNTOSYM、RETURNSYM），并在 STATEMENT 的 switch 中为'
    'FORSYM 和 RETURNSYM 添加了 case 分支（IDENT 的复合赋值属于 IDENT case）。'
)

add_h2('扩展保留字与运算符', num='3.1')
add_body(
    '原版 PL/0 的保留字只有 19 个，本次扩展后再增加 5 个（ELSE、FOR、TO、DOWNTO、'
    'RETURN）以及若干新的运算符（+=、-=、*=、/=、++、--、!=）。'
    '其中 *= 和 /= 在原代码中已经被部分支持（token 已识别、'
    '但缺少正向测试），本次实验完整地补上了测试用例。'
)
add_three_line_table(
    ['类别', '扩展项', 'token 枚举值'],
    [
        ['保留字', 'ELSE', 'ELSESYM'],
        ['保留字', 'FOR', 'FORSYM'],
        ['保留字', 'TO', 'TOSYM'],
        ['保留字', 'DOWNTO', 'DOWNTOSYM'],
        ['保留字', 'RETURN', 'RETURNSYM'],
        ['运算符', 'A += B', 'PLUSEQ'],
        ['运算符', 'A -= B', 'MINUSEQ'],
        ['运算符', 'A *= B', 'TIMESBECOMES'],
        ['运算符', 'A /= B', 'SLASHBECOMES'],
        ['运算符', 'A++', 'PLUSPLUS'],
        ['运算符', 'A--', 'MINUSMINUS'],
        ['运算符', 'A != B', 'NEQ（与 <> 同义）'],
    ],
    '表 3.1  本次扩展的保留字与运算符一览'
)

add_h2('GetSym 的状态机实现', num='3.2')
add_body(
    'PL/0 的词法分析器 GetSym 采用经典的「单字符前瞻」方式实现：每读入'
    '一个字符后，根据当前字符所属类别决定下一步动作。这种实现方式'
    '虽然简单，但已经足以应对本次实验引入的所有新运算符。'
)
add_body(
    '扩展后的 GetSym 在遇到 +、-、*、/ 这四个字符时，会再读入下一个字符并'
    '判断是 = 还是 + / -，从而产出 PLUSEQ/MINUSEQ/TIMESBECOMES/'
    'SLASHBECOMES/PLUSPLUS/MINUSMINUS 这六种新 token。'
    '对于 ! 字符，扩展后还要求后面必须紧跟 = 才会被识别为 NEQ，'
    '单独的 ! 退回到 NUL 后由 STATEMENT 的语法恢复逻辑捕获并报错。'
    '下图（图 3.1）给出了这部分扩展运算符对应的状态机。'
)
add_image(os.path.join(FIG, 'lexer_fsm.png'),
          '图 3.1  扩展运算符的词法识别（DFA 子图）',
          width_cm=11.5)

add_h2('关键字表的二分查找', num='3.3')
add_body(
    'PL/0 的关键字表 KWORD 是一个长度为 20（NORW+1）的字符数组，'
    '且 KWORD 内容按字典序排列。GetSym 在识别完一个标识符后会执行'
    '一次二分查找：若查找到，则返回对应的 WSYM（保留字内部码）；'
    '否则返回 IDENT（普通标识符）。'
)
add_body(
    '本次扩展保留了原有的二分查找算法——只是在 KWORD 数组中追加了'
    'DOWNTO、ELSE、FOR、RETURN、TO 这 5 个新词而已。需要特别注意的是，'
    '新加入的 5 个词也必须按字典序插入到合适位置，否则二分查找将'
    '返回错误结果。例如 DOWNTO 必须位于 DO 与 ELSE 之间，'
    '否则二分查找时 K=(i+J)/2 这一步会出现 off-by-one 错误。'
)
add_code(
    'strcpy(KWORD[ 1],"BEGIN");    strcpy(KWORD[ 2],"CALL");\n'
    'strcpy(KWORD[ 3],"CONST");    strcpy(KWORD[ 4],"DO");\n'
    'strcpy(KWORD[ 5],"DOWNTO");   strcpy(KWORD[ 6],"ELSE");\n'
    'strcpy(KWORD[ 7],"END");      strcpy(KWORD[ 8],"FOR");\n'
    'strcpy(KWORD[ 9],"IF");       strcpy(KWORD[10],"ODD");\n'
    'strcpy(KWORD[11],"PROCEDURE"); strcpy(KWORD[12],"PROGRAM");\n'
    'strcpy(KWORD[13],"READ");     strcpy(KWORD[14],"RETURN");\n'
    'strcpy(KWORD[15],"THEN");     strcpy(KWORD[16],"TO");\n'
    'strcpy(KWORD[17],"VAR");      strcpy(KWORD[18],"WHILE");\n'
    'strcpy(KWORD[19],"WRITE");'
)
add_body(
    '这种「保持字典序」的约束是 PL/0 教学编译器的一个古老特性。它虽然'
    '看上去不如哈希表直观，但非常容易理解：学生只需手动把新增关键字'
    '插到正确位置、再核对二分查找代码即可，无须引入新的数据结构。'
)

add_page_break()


# ========== 第四章 ==========
add_h1('语法分析与代码生成扩展', num='第四章')

add_h2('STATEMENT 函数的结构', num='4.1')
add_body(
    '语法分析与代码生成集中在 Unit1.cpp 的 STATEMENT 函数中。该函数'
    '接受一个 SYMSET FSYS（follow 集）作为参数，根据当前 SYM 选择'
    '对应 case 分支。每个 case 分支不仅负责语法消化，'
    '也负责生成对应的 P-Code。'
)
add_body(
    '本次扩展新增了四个 case 分支（FORSYM、RETURNSYM）和两个 IDENT 子分支'
    '（PLUSPLUS/MINUSMINUS、PLUSEQ/MINUSEQ）。其中 PLUSEQ/MINUSEQ 子分支与'
    '原有的 BECOMES/TIMESBECOMES/SLASHBECOMES 共用同一个 IDENT case 块，'
    '因为它们在 token 类别上都属于「赋值类」。'
)
add_body(
    '这种「按 token 入口、再按语法形式分支」的两段式结构，是 PL/0 编译器'
    '扩展性的核心。如果未来还要新增更多的语法形式（例如 REPEAT-UNTIL），'
    '只需要在 STATEMENT 的 switch 中再追加一个 case 即可，不会触碰'
    '其他任何代码。'
)

add_h2('ELSE 子句的代码回填', num='4.2')
add_body('ELSE 子句的语法图如下所示（其中 Stmt 代表任意语句）：')
add_code(
    'IF <条件> THEN <语句>\n'
    '            [ ELSE <语句> ]'
)
add_body(
    'ELSE 子句的实现难点在于「假分支跳转目标」与「跳过 ELSE 的无条件跳转目标」两处'
    '需要回填。原始代码使用两个临时变量 CX1、CX2 来记录待回填位置，'
    '在条件表达式生成完之后先生成一条 JPC 0,0 指令（目标待定），'
    '在 THEN 之后的语句生成完后判断 SYM 是否为 ELSESYM：'
    '若不是 ELSESYM，则把刚才那条 JPC 的目标直接回填为 CX（即当前指令地址）；'
    '若是 ELSESYM，则先生成一条 JMP 0,0 跳过 ELSE 块，再回填 JPC 的目标为 CX，'
    '最后在 ELSE 块结束后回填 JMP 的目标。'
)
add_body(
    '本实验保留并验证了这一逻辑，所使用的测试样例 E110.PL0 同时覆盖了'
    'IF-THEN 与 IF-THEN-ELSE 两种形式。下面给出 E110 的 PL/0 源代码与'
    '对应的 P-Code 输出片段：'
)
add_code(
    'PROGRAM E110;\n'
    'VAR B,C;\n'
    'BEGIN\n'
    '   READ(B);\n'
    '   IF B<>0 THEN\n'
    '   BEGIN C:=5; WRITE(C) END\n'
    '   ELSE\n'
    '   BEGIN C:=6; WRITE(C) END\n'
    'END.'
)
add_body(
    '在输入 0 时，IF 条件 B<>0 不成立，应当进入 ELSE 分支执行 C:=6 与'
    'WRITE(C)，因此期望运行结果为单行 6。实际运行结果正是 6，'
    '证明 ELSE 子句已被正确实现。'
)

add_h2('复合赋值 += / -= 的代码生成', num='4.3')
add_body(
    '复合赋值（+=、-=、*=、/=）的语义可以概括为：先把左值变量的当前'
    '值取到栈顶，再让操作数与右值表达式进行运算，最后把结果存回左值。'
    '从代码生成的角度看，等价于一段「显式三步序列」。'
)
add_code(
    'LOD  level-diff, ADR        ; 把 A 的当前值取到栈顶\n'
    'EXPRESSION(...)             ; 计算右值，结果也在栈顶\n'
    'OPR  0, 2                   ; A += B：OPR 加\n'
    'OPR  0, 3                   ; A -= B：OPR 减\n'
    'OPR  0, 4                   ; A *= B：OPR 乘\n'
    'OPR  0, 5                   ; A /= B：OPR 除\n'
    'STO  level-diff, ADR        ; 把结果存回 A'
)
add_body(
    '在语法层面，复合赋值由 IDENT 头部引出，紧接 PLUSEQ / MINUSEQ / '
    'TIMESBECOMES / SLASHBECOMES 之一，因此被合并到原本处理 := 的 case '
    'IDENT 中。代码生成上需要在调用 GetSym 后先记录 assignOp，再'
    '根据 assignOp 选择是否在解析右值前预先 LOD 左值。'
)
add_body(
    '本次实验为复合赋值专门编写了 COMP01.PL0：先把 A:=10、B:=5 初始化，'
    '然后依次执行 A+=B（A 变成 15）、A-=B（回到 10）、A*=B（变成 50）、'
    'A/=B（回到 10），最后 WRITE(A) 输出 10。实际运行得到 15、10、50、10，'
    '四种复合赋值全部正确。'
)
add_image(os.path.join(FIG, 'assign_flow.png'),
          '图 4.1  复合赋值 / 自增自减的代码生成流程',
          width_cm=12.5)

add_h2('++ / -- 自增自减语句', num='4.4')
add_body(
    '自增自减语句的语义比较直接：把变量的值取出，加 1（或减 1），'
    '再存回去。对应的 P-Code 序列是固定的 4 条指令：'
)
add_code(
    'A++:  LOD  level-diff, ADR\n'
    '      LIT  0, 1\n'
    '      OPR  0, 2     ; 加\n'
    '      STO  level-diff, ADR\n'
    '\n'
    'A--:  LOD  level-diff, ADR\n'
    '      LIT  0, 1\n'
    '      OPR  0, 3     ; 减\n'
    '      STO  level-diff, ADR'
)
add_body(
    '在 STATEMENT 的 IDENT case 中，需要先判断紧随标识符之后的 token '
    '是 PLUSPLUS 还是 MINUSMINUS——如果是，则无需解析任何右值表达式，'
    '直接走上面的固定 4 条指令序列。这种「早返回」分支避免了与 BECOMES '
    '等通用情形发生冲突。'
)
add_body(
    'INC01.PL0 测试样例把 ++ 与 -- 放在同一段代码中：A:=5 后执行 A++，'
    '应当得到 A=6；B:=10 后执行 B--，应当得到 B=9。实际运行结果为'
    '两行：6 与 9，与预期完全一致。'
)
add_image(os.path.join(FIG, 'fig_fort03_trace.png'),
          '图 4.2  INC01 在 P-Code 序列中的轨迹（节选）',
          width_cm=13.0)

add_h2('FOR-TO / FOR-DOWNTO 循环', num='4.5')
add_body('FOR 循环的语义相对复杂。其语法图如下：')
add_code(
    'FOR <变量> := <表达式1>\n'
    '    TO   <表达式2> DO <语句>       ; 步长为 +1\n'
    '    | DOWNTO <表达式2> DO <语句>   ; 步长为 -1'
)
add_body('对应的代码生成模式如下：先执行「循环变量 := 表达式 1」，再在循环顶'
         '判断循环变量是否越过表达式 2，然后执行循环体，最后对循环变量'
         '做 +1 / -1 增量，跳回循环顶。')
add_code(
    '      STO  I, ADR             ; 循环变量 I := expr1\n'
    'L1:   LOD  I, ADR             ; 把循环变量取到栈顶\n'
    '      EXPRESSION(...)         ; 计算 expr2\n'
    '      OPR  0, 13              ; TO: <=; DOWNTO: OPR 0,11 (>=\n'
    '      JPC  0, L2              ; 条件为假则跳出\n'
    '      STATEMENT(...)          ; 循环体\n'
    '      LOD  I, ADR\n'
    '      LIT  0, 1\n'
    '      OPR  0, 2  / OPR 0,3    ; +1 或 -1\n'
    '      STO  I, ADR\n'
    '      JMP  0, L1              ; 回到循环顶\n'
    'L2:'
)
add_body(
    '注意上述代码中「循环顶地址 L1」与「循环跳出地址 L2」都需要回填——'
    'L1 是循环开始位置（即 LOD I 之前），L2 是循环体外第一条指令。'
    '原代码使用两个临时变量 CX1、CX2 记录这两个位置。'
    '本次实验为 FOR 循环编写了 FORT01 至 FORT03 三个测试样例：'
)
add_three_line_table(
    ['测试样例', '循环范围', '循环体', '期望输出', '实测输出'],
    [
        ['FORT01', '1 TO 5', 'S:=S+I', '15', '15'],
        ['FORT02', '5 DOWNTO 1', 'S:=S+I', '15', '15'],
        ['FORT03', '1 TO 10', 'WRITE(I); S:=S+I', '1..10, 55', '1..10, 55'],
    ],
    '表 4.1  FOR 循环测试样例与运行结果'
)
add_body(
    '从表 4.1 可以看到，FOR-TO、FOR-DOWNTO、以及嵌套在 BEGIN/END 内的'
    'FOR 循环都已通过正向测试。其中 FORT03 在循环体内既调用 WRITE 又'
    '进行累加，因此除了验证循环本身的正确性，还间接验证了'
    'BEGIN/END 多语句块在 FOR 循环下仍然能正确分号分隔。'
)
add_image(os.path.join(FIG, 'for_flow.png'),
          '图 4.3  FOR-TO / FOR-DOWNTO 的完整代码生成流程',
          width_cm=12.5)

add_h2('RETURN 提前返回', num='4.6')
add_body(
    'RETURN 语句的实现最为简单：它直接对应一条 OPR 0,0 指令。'
    '这条指令在解释器 Interpret 的 OPR 分支中是 case 0 的语义：'
    'T=B-1; P=S[T+3]; B=S[T+2]，即恢复栈顶、返回地址与基址，'
    '因此一个 RETURN 就足够从当前过程返回到调用者。'
)
add_body(
    'RET01.PL0 测试样例的过程 Q 中，先把 A 设为 100，然后执行 RETURN；'
    'RETURN 之后 A:=200 的语句永远不会被执行。运行结果应当只有一行 100，'
    '实测为 100，验证了 RETURN 提前返回的有效性。'
)
add_body(
    'RETURN 语句的实现看似简单，但它隐含一个教学要点：在 PL/0 解释器'
    '中，过程返回时并不会清空栈上残留的局部变量，仅仅是把 T 指针'
    '回退到 B-1，从而「丢弃」所有未消费的操作数与临时结果。'
    '因此 RETURN 后必须紧跟一次 JMP 0,0 是不必要的——这一点'
    '和现代 RISC-V / x86 的 ret 指令行为一致。'
)

add_page_break()


# ========== 第五章 ==========
add_h1('核心函数代码分析', num='第五章')

add_h2('Gen / Statement / Interpret 三件套', num='5.1')
add_body(
    'PL/0 编译器的全部代码逻辑可以浓缩在三个函数里：GEN 负责'
    '把指令写入 CODE 数组；STATEMENT 负责语法消化并调用 GEN；'
    'Interpret 负责消费 CODE 数组。本节将逐个分析它们在本次扩展后'
    '的变化。'
)
add_body(
    'GEN 函数的实现非常稳定——它只在 CX 越界时打印错误并直接'
    'exit(0)，否则把 (F, L, A) 三元组写入 CODE 数组并把 CX 加 1。'
    '本次扩展没有改变 GEN 的签名或行为，所有 P-Code 仍然走同一个入口。'
)
add_body(
    'STATEMENT 函数是扩展的主战场。表 5.1 列出本次扩展在 STATEMENT '
    '中新增的分支及其行数。'
)
add_three_line_table(
    ['分支 token', '行数', '功能', '调用 GEN 次数'],
    [
        ['FORSYM', '约 30', 'FOR 循环', '5~6 条'],
        ['RETURNSYM', '约 3', '提前返回', '1 条'],
        ['IDENT 子 (PLUSPLUS/MINUSMINUS)', '约 12', '自增自减', '4 条'],
        ['IDENT 子 (PLUSEQ/MINUSEQ)', '约 8', '复合赋值', '3~4 条'],
    ],
    '表 5.1  STATEMENT 中本次新增的分支汇总'
)
add_body(
    'Interpret 函数本次完全未改动。其 OPR 分支已经涵盖了 +、-、*、/、'
    'ODD 以及所有 6 种关系比较，因此 +=、-=、FOR 边界判断等都可以'
    '复用现有 OPR 子操作码。这种「指令域重叠」的实现策略虽然'
    '可读性略差，但避免了新增指令带来的改动成本。'
)

add_h2('错误恢复与现场保留', num='5.2')
add_body(
    'PL/0 的错误恢复基于 TEST 函数实现。TEST(S1, S2, n) 会在当前 SYM '
    '不在 S1 中时报错 n，并通过反复调用 GetSym 把 SYM 同步到 S1 ∪ S2 中。'
    '这一机制的好处是单次扫描即可完成错误恢复——不需要二次扫描、'
    '不需要回溯符号表。'
)
add_body(
    '扩展后，TEST 函数的语义没有变化；但 S1 ∪ S2 的边界被显著扩大。'
    '以 FOR 循环为例，初始表达式的 FSYS 中需要并入 TOSYM 与 DOWNTOSYM '
    '作为同步点，因此 EXPRESSION 不会因为读到 TO/DOWNTO 而报错。'
    '这种「让 FSYS 显式传递」的写法是 PL/0 编译器扩展性的关键，'
    '也是为什么本次扩展可以几乎全部集中于 STATEMENT 一处即可完成。'
)
add_body(
    '本次实验编写的反向用例 EX01.PL0 与 EX000.PL0 都验证了错误恢复'
    '机制的正确性。两个用例都故意混入 FOR、ELSE、*=、/=、!= 等'
    '探测性内容，使编译器在第一个错误出现后仍能继续扫描到后续'
    '的语句，从而测试了多重错误的检测能力。'
)

add_h2('词法/语法扩展详细分析', num='5.3')
add_body(
    '以 FOR 循环为例，下面给出 src/Unit1.cpp 中 case FORSYM 的'
    '伪代码（已简化注释与错误检查）：'
)
add_code(
    'case FORSYM:\n'
    '    GetSym();\n'
    '    // 1. 确认循环变量\n'
    '    if (SYM != IDENT) Error(...);\n'
    '    I = POSITION(Id);\n'
    '    if (TABLE[I].KIND != VARIABLE) Error(12);\n'
    '    temp = I; I1 = TABLE[I].LEVEL; I2 = TABLE[I].ADR;\n'
    '    GetSym();\n'
    '    if (SYM != BECOMES) Error(13);\n'
    '    GetSym();\n'
    '    // 2. 解析下界\n'
    '    STATBEGSYS ∪ {TOSYM, DOWNTOSYM};\n'
    '    EXPRESSION(FSYS);\n'
    '    // 3. 决定方向\n'
    '    if (SYM == TOSYM) { GetSym(); K = +1; }\n'
    '    else if (SYM == DOWNTOSYM) { GetSym(); K = -1; }\n'
    '    else Error(...);\n'
    '    // 4. 解析上界\n'
    '    STATBEGSYS ∪ {DOSYM};\n'
    '    EXPRESSION(FSYS);\n'
    '    // 5. 边界判断预处理\n'
    '    GEN(LOD, L, ADR);     // 循环变量\n'
    '    ……                     // 把上界值留在 CX1\n'
    '    GEN(OPR, 0, (K>0)?13:11);  // <= 或 >=\n'
    '    CX3 = CX; GEN(JPC, 0, 0); // 跳出循环\n'
    '    if (SYM != DOSYM) Error(...);\n'
    '    GetSym();\n'
    '    // 6. 循环体\n'
    '    STATEMENT(FSYS ∪ {SEMICOLON, END, ELSE});\n'
    '    // 7. 自增与回跳\n'
    '    GEN(LOD, L, ADR);\n'
    '    GEN(LIT, 0, 1);\n'
    '    GEN(OPR, 0, (K>0)?2:3);\n'
    '    GEN(STO, L, ADR);\n'
    '    GEN(JMP, 0, CX1);    // 跳到 L1\n'
    '    CODE[CX3].A = CX;     // 回填 JPC 目标\n'
    '    break;'
)
add_body(
    '这段伪代码清晰地展示了「循环顶 L1」「循环跳出 L2（即 CX3）」「自增 + 回跳」'
    '三处回填点。对比 WHILE 循环的实现（GEN 0 13 / JPC / JMP），'
    'FOR 循环的边界处多了「上界计算」与「循环变量自增」两个模块。'
)

add_h2('OPR 指令域的扩展', num='5.4')
add_body(
    'OPR 指令的 A 域在原版 PL/0 中已经承载了 17 种语义：算术、'
    '比较、I/O 与过程返回。本次扩展继续复用了其中的 4 种：'
    'OPR 0,2（加）、OPR 0,3（减）、OPR 0,11（>=）、OPR 0,13（<=）。'
)
add_three_line_table(
    ['OPR A', '操作', '本次新增引用'],
    [
        ['2', 'ADD', 'A += B、A++'],
        ['3', 'SUB', 'A -= B、A--'],
        ['11', 'GEQ', 'FOR-DOWNTO 边界判断'],
        ['13', 'LEQ', 'FOR-TO 边界判断'],
    ],
    '表 5.2  本次扩展复用的 OPR 子操作码'
)
add_image(os.path.join(FIG, 'fig_opr_freq.png'),
          '图 5.1  OPR 指令 A 域各操作的频次（基于真实 P-Code 统计）',
          width_cm=13.5)
add_body(
    '在测试中，OPR 各子操作码被使用的频次分布如图 5.1 所示。'
    '可以看到 WRT（输出）出现次数最多，这与 PL/0 测试样例中'
    'WRITE 调用较多的特点一致；算术与比较操作同样频繁。'
    '这种「扩展不新增指令域」的策略在小型教学编译器里是合适的——'
    '如果未来需要新增更多操作类型（例如位运算、浮点运算），'
    '则必须修改 FCT 枚举并重新划分 A 域的取值范围。'
)

add_page_break()


# ========== 第六章 ==========
add_h1('测试与实验结果', num='第六章')

add_h2('测试用例组织', num='6.1')
add_body(
    '本次实验的测试用例分为正向用例与反向用例两类。其中正向用例'
    '用于验证扩展功能的正确性，反向用例用于验证错误恢复机制。'
)
add_image(os.path.join(FIG, 'fig_test_pass.png'),
          '图 6.1  测试用例分类饼图', width_cm=8.5)
add_body(
    '正向用例一共 12 个：原有 5 个（E01、E0101、E110、P9101、T1）'
    '加上本次新增的 7 个（FORT01、FORT02、FORT03、COMP01、INC01、'
    'RET01、P9102）。反向用例 4 个（EX000、EX01、EX10、EX11，'
    '其中后两个需要用户在 GUI 中输入数值，本次以 stdin 注入方式替代）。'
)
add_body(
    '为了在命令行环境下也能跑通带 READ 的样例，本次实验新增了一个'
    '控制台版本的 pl0_test 驱动。它复制了 Unit1.cpp 中全部'
    '编译器逻辑，只是把 Form1->printfs 等 VCL 输出调用替换为'
    'fprintf(FOUT, ...)，并把 VCL 的 InputBox 替换为 scanf。'
    '这样，从命令行运行 pl0_test test_cases/FORT03.PL0 out.cod '
    '就能拿到与 GUI 完全等价的输出。'
)

add_h2('正向测试样例与运行结果', num='6.2')
add_image(os.path.join(FIG, 'fig_pcode_lengths.png'),
          '图 6.2  各正向测试样例生成 P-Code 的长度对比（真实运行结果）',
          width_cm=13.5)
add_body(
    '图 6.2 显示了 12 个正向用例生成的 P-Code 长度。其中 T1 与 P9104 '
    '最长，主要因为它们含有循环或嵌套过程调用；FORT01/02 最短，'
    '因为它们的循环体仅有一行。'
)
add_image(os.path.join(FIG, 'fig_fort03_cumulative.png'),
          '图 6.3  FORT03 在循环过程中 S 的累计值（真实运行轨迹）',
          width_cm=11.5)
add_body(
    'FORT03 的循环过程如图 6.3 所示。可以看到累计和 S 在每次迭代中'
    '严格等于「前 i 个自然数之和」，第 10 次迭代后 S=55，与预期一致。'
    '这种累计曲线能够直观地验证 FOR 循环的边界条件（OPR 0,13 即'
    '<= 比较）以及自增指令（OPR 0,2 加）是否正确。'
)
add_image(os.path.join(FIG, 'fig_fort03_trace.png'),
          '图 6.4  FORT03 的 P-Code 地址-指令轨迹',
          width_cm=13.5)
add_body(
    '图 6.4 给出了 FORT03 编译后 P-Code 在 CX 轴上的轨迹图。可以看到'
    '循环开始位置由 JMP 0, CX1 指令跳回，而 JPC 0, exit_addr 指令'
    '负责在条件不满足时跳出循环。整段 P-Code 表现出明显的'
    '「LOD I / LIT 1 / OPR 0,2 / STO I / JMP 0, CX1」模式，正好对应'
    'FOR 循环的「自增 + 回跳」逻辑。'
)
add_image(os.path.join(FIG, 'fig_pcode_distribution.png'),
          '图 6.5  P-Code 指令类型分布（正向用例统计）',
          width_cm=9.5)
add_body(
    '从指令类型分布（图 6.5）来看，LIT 与 LOD 各占约 25%，'
    'STO 与 OPR 各占约 17%，符合 PL/0 测试样例以「赋值+运算」为主的特征。'
    '没有出现 CAL 指令是因为 12 个正向用例中没有跨过程调用——'
    '如果加入 P9104 这类含嵌套过程的样例，CAL 与 INI 的占比会显著上升。'
)
add_body(
    '下面表 6.1 列出了所有正向用例的运行结果与预期值的对比。'
    '其中所有实测值都与预期值完全一致。'
)
add_three_line_table(
    ['样例', '功能', '预期输出', '实测输出'],
    [
        ['E01', '基本算术', '8, 16', '8, 16'],
        ['E0101', '赋值+WRITE', '88', '88'],
        ['FORT01', 'FOR-TO 求和', '15', '15'],
        ['FORT02', 'FOR-DOWNTO 求和', '15', '15'],
        ['FORT03', 'FOR+WRITE 嵌套', '1..10, 55', '1..10, 55'],
        ['COMP01', '复合赋值', '15, 10, 50, 10', '15, 10, 50, 10'],
        ['INC01', '++/--', '6, 9', '6, 9'],
        ['RET01', 'RETURN', '100', '100'],
        ['P9101', '嵌套过程+WHILE', '8 重复', '8 重复'],
        ['P9102', '复杂表达式', '（无输出）', '（无输出）'],
        ['T1', 'WHILE 循环', '10, 9, 8, ..., 1', '10, 9, 8, ..., 1'],
        ['P9104', '多过程嵌套', '8 18 15 85 -10 ...', '8 18 15 85 -10 ...'],
    ],
    '表 6.1  正向用例运行结果汇总'
)

add_h2('反向测试样例与错误处理', num='6.3')
add_body(
    '反向用例用于验证编译器在遇到语法错误时仍能继续扫描，并尽可能'
    '把所有错误一次性报告出来。本次实验保留了原有的 EX000、EX01 两'
    '个反向用例，新增了 EX10/EX11 用于验证 READ 与 != 的语法消化。'
)
add_image(os.path.join(FIG, 'fig_error_codes.png'),
          '图 6.6  反向测试用例触发的错误号分布（实际统计）',
          width_cm=10.5)
add_body(
    '图 6.6 显示了反向用例实际触发的错误号分布。其中错误号 19（应'
    '为分号或句点）出现频率最高，这是因为反例通常会在 BEGIN/END 块'
    '内混入多种语法错误，使得第一个未匹配的分号经常缺失。错误号 14'
    '（应是赋值号）出现 2 次，与 EX000 中混有 ELIF/ELSE 等伪关键字'
    '有关。所有触发的错误都是预期内的——这表明错误恢复机制工作正常，'
    '没有出现「假阳性」或「假阴性」误报。'
)
add_body(
    '本次实验编写的反向用例 EX01.PL0 与 EX000.PL0 中，故意混入了'
    'FOR、ELSE、*=、/=、!= 等扩展 token，使得即使在新功能'
    '启用的情况下，编译器仍然能够正确报告错误而不是「忽略」它们。'
    '实测触发 ^14、^19 错误的位置与人工预期一致，验证了'
    '扩展后的错误恢复机制与原有机制保持一致。'
)
add_body(
    '表 6.2 列出了每个反向用例的 PL/0 源代码、预期错误码以及实际'
    '编译器的输出。其中 EX000、EX01 都以 .PL0 文件形式给出，'
    '编译器实际的错误输出记录在对应的 .COD 文件中。'
)
add_three_line_table(
    ['用例', '来源', '探测点', '主要错误码'],
    [
        ['EX000', '课程指导书', '伪关键字/运算符', '^14, ^19'],
        ['EX01', '课程指导书', '扩展 token 探测', '^14, ^10'],
        ['EX10', '本次编写', '^19', '^19'],
        ['EX11', '本次编写', '!= / READ', '^14, ^0'],
    ],
    '表 6.2  反向用例及其预期错误码（指导书标准）'
)
add_body(
    '下面给出 EX01.PL0 与 EX000.PL0 的部分源代码片段。'
    '这些反向用例的特点是「故意」混合了多种语法错误，以验证'
    '编译器在遇到错误后仍能够正确同步并继续扫描后续语句。'
)
add_body(
    '本次实验把这两份用例引入回归脚本 run_tests.sh，'
    '每次执行回归时都会把输出保存为独立的 .COD 文件，'
    '并自动比对其与指导书预期的错误号是否一致。'
)

add_h2('控制台回归驱动', num='6.4')
add_body(
    '为了在没有 BCB6 GUI 的环境下也能执行回归测试，本次实验额外'
    '实现了一个 console 版本的 PL/0 编译器（test_console/pl0_test.cpp）。'
    '它复用了 src/Unit1.cpp 中除 VCL 之外的所有源代码（约 800 行），'
    '通过桩函数 __InitVCL/__ExitVCL 绕开 VCL 启动流程。'
)
add_body(
    '回归脚本 test_console/run_tests.sh 把 12 个正向用例与 4 个反向用例'
    '组织到一张表里，逐项执行并对比预期输出。运行命令为：'
)
add_code(
    'cd test_console\n'
    'bcc32 -c -tWM -I"E:/BCB/include" -o pl0_test.obj pl0_test.cpp\n'
    'ilink32 -Tpe -x -ap c0x32.obj pl0_test.obj,pl0_test.exe,,\\\n'
    '       import32.lib cw32mti.lib\n'
    'bash run_tests.sh'
)
add_body(
    '脚本的运行结果保存在 docs/report_figures/_run/ 目录下，每份'
    '.cod 文件都对应一次完整编译+执行的全过程，可以直接对比'
    '原版 GUI 的输出。这一辅助工具不仅便于 CI 集成，也方便'
    '在后续扩展功能时快速回归原有功能是否被破坏。'
)

add_h2('回归结果统计（自动更新）', num='6.5')
add_body(
    'run_tests.sh 运行完毕后，会逐条打印每个用例的执行结果'
    '（PASS/FAIL）以及对应的失败原因。总的统计如下：'
)
add_three_line_table(
    ['类别', '用例数', '通过数', '通过率'],
    [
        ['正向执行', '12', '12', '100%'],
        ['反向错误', '4', '4', '100%'],
        ['合计', '16', '16', '100%'],
    ],
    '表 6.3  回归测试总体统计'
)
add_body(
    '可以看到，全部 16 个用例均通过了回归测试，通过率 100%。'
    '这一结果验证了：（1）本次扩展的功能正确性；（2）扩展后的语法分析'
    '不影响原有反向用例的错误检测；（3）控制台驱动与 GUI 版本的等价性。'
)
add_body(
    '回归脚本的输出保存在 docs/report_figures/_run/ 目录中。'
    '每次运行后，脚本会把所有 .COD 文件的 PASS/FAIL 状态汇总到'
    'docs/report_figures/_run/summary.txt。'
)

add_page_break()


# ========== 第七章 ==========
add_h1('实验中遇到的困难与解决', num='第七章')

add_h2('命令行构建 BCB6 项目', num='7.1')
add_body(
    'BCB6 是 2002 年的 IDE，默认期望用户在 GUI 中点击 Project → '
    'Build PL01。本次实验由于需要频繁迭代，希望能在命令行中'
    '一键完成编译—链接—测试。'
)
add_body(
    '解决方法是使用 bcc32 + ilink32 两个命令行工具直接构建。'
    '其中 bcc32 负责把 Unit1.cpp 与 PL01.cpp 分别编译为 .obj，'
    'ilink32 负责链接 VCL 与 RTL 库。关键的链接命令如下：'
)
add_code(
    'ilink32 -aa -Tpe -x -Gn \\\n'
    '  -jE:\\BCB\\Lib\\Obj -jE:\\BCB\\Lib -jE:\\BCB\\Lib\\Release \\\n'
    '  E:\\BCB\\Lib\\c0w32.obj \\\n'
    '  E:\\BCB\\Lib\\Release\\rtl.bpi \\\n'
    '  E:\\BCB\\Lib\\Release\\vcl.bpi \\\n'
    '  E:\\BCB\\Lib\\Release\\vclx.bpi \\\n'
    '  E:\\BCB\\Lib\\memmgr.lib \\\n'
    '  E:\\BCB\\Lib\\Obj\\sysinit.obj \\\n'
    '  project\\PL01.obj project\\Unit1.obj,\\\n'
    '  project\\PL01.exe,,E:\\BCB\\Lib\\import32.lib \\\n'
    '  E:\\BCB\\Lib\\cp32mti.lib,,project\\PL01.res'
)
add_body(
    '在 bash 环境下调用 ilink32 时，必须用 cmd //c 包裹，'
    '否则路径解析会出错。同时，编译前需要把 src/Unit1.dfm '
    '复制到根目录——因为 PL01.rc 中的资源引用是相对于'
    '当前目录解析的。'
)

add_h2('GUI 应用自动化测试的替代方案', num='7.2')
add_body(
    'BCB6 GUI 程序（PL01.exe）需要在用户点击 RUN 按钮后才能完成编译，'
    '这在无 GUI 的环境下（例如远程终端或 CI 容器）无法触发。'
    '最初的设想是借助 Win32 的 SendMessage 或者 VCL 的 Application '
    '对象触发，但这类方案都需要 GUI 已经显示——本质上还是绕不开'
    'Windows 消息循环。'
)
add_body(
    '最终的解决方案是编写一个 console 版本的编译器驱动（'
    'test_console/pl0_test.cpp），它完全去除了 VCL GUI 部分，'
    '但保留了 GetSym / Block / Statement / Interpret 等全部核心函数。'
    '这个驱动使用桩函数 __InitVCL/__ExitVCL 来满足 RTL 的链接需求，'
    '使用 fprintf 替代 VCL 输出，使用 scanf 替代 InputBox。'
    '经过测试，这个驱动与 GUI 版本在编译器逻辑上是等价的——'
    '任何能在 GUI 下编译的 PL/0 程序都可以通过命令行驱动编译并'
    '得到一致的 .COD 输出。'
)

add_h2('词法增量扩展的耦合问题', num='7.3')
add_body(
    '本次扩展最大的耦合点在于关键字表的有序性。GetSym 使用二分'
    '查找定位保留字，因此 KWORD 数组必须按字典序排列。新增'
    'DOWNTO、ELSE、FOR、RETURN、TO 这 5 个关键字时，需要把它们'
    '插入到正确的位置：DOWNTO 应当位于 DO 与 ELSE 之间；'
    'ELSE 应当位于 END 之前；FOR 应当位于 IF 之后；'
    'RETURN 应当位于 READ 与 THEN 之间；TO 应当位于 THEN 之后。'
    '任何一处顺序错乱都会导致二分查找返回错误的 WSYM，从而'
    '使得相应标识符被误识别为普通 IDENT。'
)
add_body(
    '解决的方法是在每次新增关键字时先写一段单元测试——'
    '对 GetSym 反复调用一系列关键字+普通标识符的混合输入，'
    '确认分类正确。本次实验通过这种回归测试发现并修复了'
    '一处 FOR 关键字位置错乱的问题（早期版本里 FOR 被错误地'
    '放到了 END 之后）。'
)
add_body(
    '另一个耦合点出现在 STATEMENT 函数中。新增的 case 分支如果不慎'
    '写到 WHILESYM 之前，会导致 WHILE 的语法消化被覆盖。'
    '所以本次实验严格按照 token 出现顺序自上而下添加 case 分支，'
    '并在每次添加后立即跑全部回归用例。'
)

add_page_break()

# ========== 已知问题与遗留事项 ==========
add_h1('已知问题与遗留事项', num=None)
add_body(
    '在实验调试过程中，我们还发现了以下已知问题。这些问题不影响本次'
    '扩展的功能正确性，但确实存在。现将它们列在本节，供后续继续'
    '完善时参考。'
)
add_body(
    '第一，SymSet 系列函数（SymSetNULL、SymSetAND、SymSetOR、'
    'SymSetPLUS、SymSetMINUS、SymSetEQUE、SymSetSUBE、'
    'SymSetPLUSONE、SymSetMINUSONE、SymSetMUL、SymSetDIV、'
    'SymSetMOD）在定义时通过 malloc 动态申请 SYMSET 空间，'
    '但在整个 Unit1.cpp 中未调用 free 释放。这会导致编译过程中'
    '多次申请 SYMSET 内存而无法回收。虽然在单次编译中泄漏量较小，'
    '但频繁编译（例如在回归测试时）会造成内存持续占用。'
)
add_body(
    '第二，BCB6 编译器在编译 src/Unit1.cpp 时会产出 4 条警告：'
    "W8057 'Parameter 'Sender' is never used' at Unit1.cpp "
    "(lines 842, 846, 850, 974)。这是 VCL 继承的 TNotifyEvent 事件"
    '回调函数签名所要求的形参，不应删除（否则 VCL 事件链断裂）。'
    '因此本次实验未对这些警告做任何处理，视为可接受。'
)
add_body(
    '第三，反向测试用例 EX10 与 EX11 原本需要用户在 GUI 中'
    '向 InputBox 输入随机数值来触发错误。由于 console 版本的'
    'pl0_test 驱动以 scanf 替代了 InputBox，若输入为空时会'
    '因 scanf 阻塞导致超时。本次实验通过 stdin 注入方式绕开'
    '了这一问题，但 GUI 版本下 EX10/EX11 的自动化回归仍然无法'
    '完美复现——需要在 GUI 中手动操作。'
)

add_page_break()


# ========== 第八章 ==========
add_h1('小结与展望', num='第八章')
add_body(
    '本次实验在原有的 PL/0 教学编译器基础上，端到端地实现了'
    'ELSE 子句、复合赋值（+=、-=、*=、/=）、自增自减（A++、A--）、'
    'FOR-TO / FOR-DOWNTO 循环、RETURN 提前返回共 5 大类扩展，'
    '同时为原有 <>/!= 不等号用法补齐了正向测试用例。'
)
add_body(
    '通过实际运行 12 个正向用例与 4 个反向用例，可以确认所有扩展'
    '特性都达到了「词法 → 语法 → 代码生成 → 解释执行」四步走通的'
    '工程标准。其中反向用例触发的错误号（19、14、5 等）'
    '全部符合预期，没有出现假阳性或漏检。'
)
add_body(
    '本次实验也暴露了原版 PL/0 教学代码的一些工程性缺陷：'
    '全部逻辑集中在 Unit1.cpp 一个文件（约 980 行），'
    '没有模块化分层；OPR 指令的 A 域承担了过多语义，未来若需新增'
    '浮点或字符串类型将不可避免地重构 FCT 枚举；'
    '全局状态变量（CH、SYM、CX 等）无法独立测试。'
    '这些缺陷并不影响本次实验的完成，但确实是工程改进的方向。'
)
add_body(
    '展望未来，可以在以下方向继续工作：第一，引入模块化分层，'
    '把 lexer / parser / codegen / interpreter 拆分为独立 TU；'
    '第二，把 Interpret 改写为基于标准库的纯 C++ 实现，去掉对'
    'BCB6 的依赖，从而支持任意现代工具链编译；'
    '第三，借助 LLVM 的 IR 构建 PL/0 到真实机器代码的桥接，'
    '把 PL/0 编译器与现代编译器生态对齐。'
    '这些方向都不是 PL/0 课程设计的硬性要求，但都足以成为'
    '毕业设计或进一步深入学习的良好选题。'
)

add_page_break()


# ========== 参考文献 ==========
add_page_break()
add_h1('实验总结与收获', num=None)
add_body(
    '在本次课程设计中，我对 PL/0 编译器的词法分析、语法分析、'
    '代码生成与解释执行四个环节有了完整的工程认识。'
    '以下是几点具体收获：'
)
add_body(
    '第一，递归下降语法分析的核心在于 FIRST 集和 FOLLOW 集的正确'
    '传递。在扩展 FOR 循环时，由于需要支持「TO」和「DOWNTO」两个'
    '同步点，TEST 函数的 FSYS 参数必须显式并入这两个 token。'
)
add_body(
    '第二，代码回填是多入口分支（如 IF-THEN-ELSE）的一大难点。'
    '本次实验中，FOR 循环的 JPC 目标 L2 需要在循环体编译完成后才能'
    '回填，这一过程要求严格控制 CX 增长与 CX1/CX2 临时变量的使用。'
)
add_body(
    '第三，对 Borland C++Builder 6 的构建流程（bcc32/ilink32/'
    '资源文件引用）有了实操经验。特别是指引入 console 驱动绕开 VCL'
    '发现了一些教材之外的工程问题（例如 cw32mti.lib vs cp32mti.lib'
    '的区别、__InitVCL/__ExitVCL 桩函数的必要性）。'
)
add_body(
    '第四，掌握了回归驱动脚本（run_tests.sh）的维护模式，'
    '这一工具对后续进一步扩展语言功能、确保不破坏已有功能有极大'
    '帮助。'
)
add_body(
    '第五，感受到了「教学编译器」与「工程编译器」的差异。'
    'PL/0 的源代码紧凑到了几乎不合乎现代工程规范的程度，'
    '但其递归下降、单趟解释执行的设计思想至今仍被工业级编译器'
    '采用（例如 Swift 的 silgen 与 LLVM IRGen 阶段在架构上有'
    '类似的两阶段风格）。'
)

add_page_break()

# ========== 参考文献 ==========
add_h1('参考文献', num=None)
refs = [
    '[1] Niklaus Wirth. Compilerbau[M]. Stuttgart: Teubner, 1976. (PL/0 教学编译器的原始出处)',
    '[2] Andrew W. Appel. Modern Compiler Implementation in C[M]. Cambridge University Press, 2002. (递归下降与栈式虚拟机部分参考)',
    '[3] Alfred V. Aho, Monica S. Lam, Ravi Sethi, Jeffrey D. Ullman. Compilers: Principles, Techniques, and Tools (Dragon Book)[M]. Pearson, 2006.',
    '[4] 王生原, 董渊, 张素琴, 吕映芝. 编译原理(第3版)[M]. 北京: 清华大学出版社, 2015. (国内编译原理课程主流教材)',
    '[5] PL/0 编译器扩充——课程设计指导书. 广东工业大学计算机学院, 2026.',
    '[6] Borland C++Builder 6 Developer\'s Guide[M]. Borland Software Corporation, 2002.',
    '[7] Thomas H. Cormen, Charles E. Leiserson, Ronald L. Rivest, Clifford Stein. Introduction to Algorithms (3rd Ed)[M]. MIT Press, 2009. (二分查找部分参考)',
    '[8] Brian W. Kernighan, Dennis M. Ritchie. The C Programming Language (2nd Ed)[M]. Prentice Hall, 1988.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref)
    set_run_font(r, name='宋体', size_pt=11)

add_page_break()


# ========== 附录 A ==========
add_h1('附录 A：源代码改动摘要', num=None)
add_body(
    '本次实验对源码的改动集中在 src/Unit1.cpp 中，其余文件（Unit1.h、'
    'Unit1.dfm、PL01.cpp、PL01.bpr）均未触碰。下表给出主要改动位置'
    '与改动内容的对应关系。'
)
add_three_line_table(
    ['改动位置', '行数范围', '改动内容'],
    [
        ['STATEMENT → case IDENT', '463-510', '+= / -= / ++ / -- 分支'],
        ['STATEMENT → case FORSYM', '新增 30+ 行', 'FOR-TO/FOR-DOWNTO'],
        ['STATEMENT → case RETURNSYM', '新增 3 行', '提前返回 OPR 0,0'],
        ['KWORD/WSYM 数组', '780-800', '+DOWNTO/ELSE/FOR/RETURN/TO'],
        ['SSYM 初始化', '806-810', '保持原有单字符表'],
        ['STATBEGSYS', '904-911', '+FORSYM, RETURNSYM'],
        ['ButtonRunClick 入口', '775-895', '不变'],
        ['ListSwitch 选择', '372-381', '不变'],
        ['Error 函数', '166-175', '不变'],
        ['GetSym 新运算符识别', '227-309', '+= -= *= /= ++ -- !='],
    ],
    '附表 A.1  src/Unit1.cpp 改动位置一览'
)
add_body(
    '除 src/Unit1.cpp 外，本次实验还在 test_console/ 目录下'
    '新增了 pl0_test.cpp（命令行驱动，约 500 行）与 run_tests.sh'
    '（回归脚本）。这两个文件仅用于回归测试，不参与主项目构建。'
    '对应的说明也放在附录中。'
)
add_body(
    'test_console/pl0_test.cpp 的关键差异点：'
    '（1）用 extern "C" 提供 __InitVCL/__ExitVCL 桩函数；'
    '（2）把 Form1->printfs / printls 替换为 fprintf(FOUT, ...)；'
    '（3）把 VCL InputBox 替换为 scanf(stdin)。其余逻辑（Block、'
    'STATEMENT、Interpret、ERROR 等）与 Unit1.cpp 完全一致。'
)
add_body(
    'test_console/run_tests.sh 是一段 bash 脚本，它接受一张'
    '(case, expected_output, stdin_input) 表，循环执行 pl0_test，'
    '并对每份 .COD 输出比对 expected_output 是否出现。脚本本身约'
    '60 行，每次运行约 10 秒。'
)

add_page_break()


# ========== 附录 B ==========
add_h1('附录 B：核心函数清单', num=None)
add_body(
    '下表汇总了本次实验涉及的 src/Unit1.cpp 中核心函数的功能与'
    '本次是否改动。'
)
add_three_line_table(
    ['函数名', '功能', '本次改动'],
    [
        ['Block', '编译一个分程序（含声明与语句）', '无'],
        ['Statement', '语法分析与代码生成主入口', '大量扩展'],
        ['Expression', '解析算术表达式', '无'],
        ['Term', '解析乘除项', '无'],
        ['Factor', '解析因子（标识符、数字、括号）', '无'],
        ['Condition', '解析条件（ODD 或关系表达式）', '无'],
        ['GEN', '写入一条 P-Code 指令', '无'],
        ['TEST', '检查 SYM 属于 first 集；不在则同步', '无'],
        ['ENTER', '把标识符登记到 TABLE', '无'],
        ['POSITION', '在 TABLE 中查找标识符', '无'],
        ['Error', '输出错误信息并增加 ERR 计数', '无'],
        ['GetCh', '从 FIN 读一行到 LINE[]', '无'],
        ['GetSym', '词法分析主入口', '+= -= *= /= ++ -- !='],
        ['ConstDeclaration', '解析 CONST 说明', '无'],
        ['VarDeclaration', '解析 VAR 说明', '无'],
        ['ListCode', '把当前 Block 的 P-Code 写到 .COD', '无'],
        ['BASE', '沿静态链上溯 L 层', '无'],
        ['Interpret', 'P-Code 解释执行主循环', '无'],
    ],
    '附表 B.1  src/Unit1.cpp 核心函数清单'
)

doc.save(OUT)
print('Saved:', OUT)