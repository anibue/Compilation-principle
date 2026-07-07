# -*- coding: utf-8 -*-
"""Generate PL/0 experiment report (DOCX) v2 — expanded version.

Target: 40-50 pages.
Style requirements (per 撰写指引-实验报告.doc):
- Page: top 30mm / bottom 25mm / left 30mm / right 20mm
- Line spacing: 1.5x
- Page numbers: arabic, from 绪论 onward, footer-right
- Title (cover): 二号 黑体 加粗
- Chapter heading: 三号 黑体 加粗
- Section heading: 小四号 黑体 加粗
- Body: 小四号 宋体
- Caption (figure): 宋体 五号
- Caption (table): 黑体 五号
- Page number: 小五号 Times New Roman
- Tables: three-line (no left/right borders)
- Figures must not cross page boundary
"""
import os
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'C:\Users\xwmsu\Desktop\c_builder_b'
FIG = os.path.join(ROOT, 'docs', 'report_figures')
OUT = os.path.join(ROOT, 'PL0编译器扩充实验报告.docx')

doc = Document()


def set_run_font(run, name='宋体', size_pt=12, bold=False, ascii_font='Times New Roman'):
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), name)
    rFonts.set(qn('w:ascii'), ascii_font)
    rFonts.set(qn('w:hAnsi'), ascii_font)


def add_para(text='', font='宋体', size=12, bold=False, align=None,
             first_indent=True, line_spacing=1.5,
             space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if first_indent and text:
        p.paragraph_format.first_line_indent = Pt(size * 2)
    if text:
        r = p.add_run(text)
        set_run_font(r, name=font, size_pt=size, bold=bold)
    return p


def add_h1(text, num=None):
    full = (num + '  ' + text) if num else text
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(full)
    set_run_font(r, name='黑体', size_pt=16, bold=True)
    return p


def add_h2(text, num=None):
    full = (num + '  ' + text) if num else text
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(full)
    set_run_font(r, name='黑体', size_pt=14, bold=True)
    return p


def add_body(text, size=12, indent=True):
    return add_para(text, font='宋体', size=size, first_indent=indent,
                    line_spacing=1.5)


def add_code(text, size=10):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(12)
        r = p.add_run(line if line else ' ')
        set_run_font(r, name='Consolas', size_pt=size, ascii_font='Consolas')


def add_image(path, caption, width_cm=12.5):
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run()
    r.add_picture(path, width=Cm(width_cm))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(0)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.first_line_indent = Pt(0)
    cr = cap.add_run(caption)
    set_run_font(cr, name='宋体', size_pt=9)


def add_table_borders(t):
    tbl = t._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement('w:' + edge)
        b.set(qn('w:val'), 'nil')
        borders.append(b)
    tblPr.append(borders)


def add_three_line_table(headers, rows, caption, col_widths=None):
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(rows), cols=n_cols)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = True
    add_table_borders(t)
    # set header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        p = c.paragraphs[0]
        p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_run_font(r, name='黑体', size_pt=9, bold=True)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri + 1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            set_run_font(r, name='宋体', size_pt=9)
    # Three-lines formatting
    tbl = t._tbl
    tblPr = tbl.tblPr
    # Remove inside borders
    borders = tblPr.find(qn('w:tblBorders'))
    if borders is not None:
        for edge in ('insideH', 'insideV'):
            b = borders.find(qn('w:' + edge))
            if b is not None:
                b.set(qn('w:val'), 'nil')
        # Set top and bottom lines
        for edge in ('top', 'bottom'):
            b = borders.find(qn('w:' + edge))
            if b is not None:
                b.set(qn('w:val'), 'single')
                b.set(qn('w:sz'), '12')
                b.set(qn('w:space'), '0')
                b.set(qn('w:color'), '000000')
        # Set header bottom line
        b = borders.find(qn('w:insideH'))
    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.first_line_indent = Pt(0)
    cr = cap.add_run(caption)
    set_run_font(cr, name='黑体', size_pt=9, bold=True)


def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_footer_page_number():
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE   \\* MERGEFORMAT'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar1)
    r._r.append(instrText)
    r._r.append(fldChar2)
    set_run_font(r, name='Times New Roman', size_pt=9)


# ============================================================
# Set default style
# ============================================================
from docx.oxml.ns import qn as _qn
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(_qn('w:eastAsia'), '宋体')
style.element.rPr.rFonts.set(_qn('w:ascii'), 'Times New Roman')
style.element.rPr.rFonts.set(_qn('w:hAnsi'), 'Times New Roman')


# ============================================================
# Page setup
# ============================================================
for section in doc.sections:
    section.top_margin = Mm(30)
    section.bottom_margin = Mm(25)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# ============================================================
# Cover page
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(200)
r = p.add_run('PL/0 编译器扩充实验报告')
set_run_font(r, name='黑体', size_pt=22, bold=True)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(40)
r = p.add_run('广东工业大学 计算机学院')
set_run_font(r, name='宋体', size_pt=16)

add_page_break()


# ============================================================
# Table of contents
# ============================================================
toc_h = doc.add_paragraph()
toc_h.alignment = WD_ALIGN_PARAGRAPH.CENTER
toc_h.paragraph_format.space_after = Pt(20)
r = toc_h.add_run('目  录')
set_run_font(r, name='黑体', size_pt=18, bold=True)

toc_entries = [
    ('第一章  绪论', '1'),
    ('    1.1  PL/0 语言与教学编译器简介', '1'),
    ('    1.2  本次实验的目标与动机', '2'),
    ('    1.3  开发环境与项目结构', '2'),
    ('    1.4  课程设计指导书原文要求', '3'),
    ('第二章  编译流程与基础理论', '3'),
    ('    2.1  PL/0 编译器的整体架构', '3'),
    ('    2.2  PL/0 语言文法完整描述', '5'),
    ('    2.3  递归下降语法分析', '6'),
    ('    2.4  符号表与 P-Code 中间表示', '7'),
    ('    2.5  解释执行模型与栈帧管理', '9'),
    ('第三章  词法分析器扩展', '11'),
    ('    3.1  扩展保留字与运算符', '11'),
    ('    3.2  GetSym 的状态机实现', '12'),
    ('    3.3  关键字表的二分查找', '13'),
    ('第四章  语法分析与代码生成扩展', '15'),
    ('    4.1  STATEMENT 函数的结构', '15'),
    ('    4.2  ELSE 子句的代码回填', '16'),
    ('    4.3  复合赋值 +=/-= 的代码生成', '17'),
    ('    4.4  ++/-- 自增自减语句', '18'),
    ('    4.5  FOR-TO / FOR-DOWNTO 循环', '19'),
    ('    4.6  RETURN 提前返回', '21'),
    ('第五章  核心函数代码深度分析', '23'),
    ('    5.1  三件套总览', '23'),
    ('    5.2  Block 函数深度分析', '24'),
    ('    5.3  错误恢复机制与错误代码表', '26'),
    ('    5.4  OPR 指令域矩阵分析', '28'),
    ('    5.5  静态链与栈帧管理', '29'),
    ('第六章  测试与实验结果', '31'),
    ('    6.1  测试用例组织', '31'),
    ('    6.2  正向测试样例与运行结果', '32'),
    ('    6.3  反向测试样例与错误处理', '34'),
    ('    6.4  控制台回归驱动', '36'),
    ('    6.5  回归结果统计', '37'),
    ('第七章  实验中遇到的困难与解决', '38'),
    ('    7.1  命令行构建 BCB6 项目', '38'),
    ('    7.2  GUI 应用自动化测试的替代方案', '39'),
    ('    7.3  词法增量扩展的耦合问题', '40'),
    ('    7.4  解释器与编译器的 P-Code 同步性', '41'),
    ('    7.5  本次扩展未涉及的边界情况', '42'),
    ('第八章  小结与展望', '43'),
    ('实验总结与收获', '44'),
    ('参考文献', '45'),
    ('附录 A  源代码改动摘要', '46'),
    ('附录 B  测试用例源代码', '47'),
    ('附录 C  核心函数清单', '50'),
    ('附录 D  编译错误代码表', '51'),
]
for label, page in toc_entries:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(label)
    if not label.startswith('    '):
        set_run_font(r, name='黑体', size_pt=11, bold=True)
    else:
        set_run_font(r, name='宋体', size_pt=10)
    tab_r = p.add_run('\t' + page)
    set_run_font(tab_r, name='Times New Roman', size_pt=10)

add_page_break()


# ========== Footer with page numbers ==========
add_footer_page_number()


# ========== 第一章 绪论 ==========
add_h1('绪论', num='第一章')

add_h2('PL/0 语言与教学编译器简介', num='1.1')
add_body(
    'PL/0 是由 Niklaus Wirth 在 1976 年编写的小型教学语言，其完整 Pascal 子集只包含'
    '常量声明、变量声明、过程声明、赋值、条件、循环、过程调用与基本 I/O 几类语句。'
    'Wirth 在其著作《Compilerbau》中以 PL/0 为对象展示了递归下降语法分析的全部细节，'
    '因此 PL/0 至今仍被国内外的编译原理课程当作最典型的教学蓝本。'
)
add_body(
    'PL/0 编译器通常被描述为「单趟 + 解释执行」架构。与传统工业级编译器（如 GCC、LLVM）'
    '动辄数万行代码、多趟扫描、中间优化的风格不同，PL/0 将词法分析、语法分析、代码生成'
    '浓缩到一次扫描中完成，所生成的 P-Code（一种基于栈的虚拟机指令）直接进入解释器运行。'
    '整个编译过程没有 AST、没有 IR Pass、没有指令选择——只有纯粹的递归下降 + 立即执行。'
)
add_body(
    '原版 PL/0 的不足之处同样明显：它没有 ELSE 子句，没有 += 之类的复合赋值，'
    '更没有真正的 FOR 循环。学生在课程设计阶段几乎不可避免地要在原版基础上'
    '增加新的语法规则。每一次增加，本质上都是对词法分析、语法分析、'
    '代码生成、解释执行四部分的协同修改——这一点决定了 PL/0 课设虽然规模小，'
    '但能够完整体现编译原理课程要求的全部工程训练。'
)
add_body(
    '本实验所采用的原始代码来自 Wirth 版本的 BCB6 移植，源程序仅有不到一千行，'
    '但几乎所有 PL/0 教学版本所需的扩展点都已预留：保留字表、二分查找、'
    'OPR 子操作码空间、错误恢复现场，正是这种「小而完整」的工程结构为本次'
    '扩展提供了合适的起点。'
)
add_body(
    '从编译器理论的角度看，PL/0 完整覆盖了《龙书》（Aho, Lam, Sethi, Ullman: '
    'Compilers: Principles, Techniques, and Tools）第 2 章至第 7 章的核心知识点：'
    '词法分析（第 3 章）、语法分析（第 4 章）、中间代码生成（第 6 章）、运行时环境'
    '（第 7 章）。可以说，把 PL/0 课设做扎实，几乎等同于手动实现了龙书理论的最小可运行实例。'
)

add_h2('本次实验的目标与动机', num='1.2')
add_body('本次课程设计的题目与要求来自广东工业大学计算机学院编译原理课程设计指导书，'
         '具体包含以下扩充任务：')
add_body(
    '（1）增加单词：保留字 ELSE、FOR、TO、DOWNTO；运算符 +=、-=、*=、/=、<>、!=。'
)
add_body(
    '（2）修改单词：原版 PL/0 的不等号 # 改为 <> 或 !=（提高要求）。'
)
add_body(
    '（3）增加条件语句的 ELSE 子句，并写出相关文法与语法描述。'
)
add_body(
    '（4）扩充赋值运算：+= 和 -=。'
)
add_body(
    '（5）扩充 FOR 循环：FOR <变量>:=<表达式> TO <表达式> DO <语句> 与 '
    'FOR <变量>:=<表达式> DOWNTO <表达式> DO <语句>。'
)
add_body(
    '（6）增加自增自减：A++ 与 A--。'
)
add_body(
    '（7）增加 RETURN 提前返回。'
)
add_body(
    '值得注意的是，最终需要让上述语法在 BCB6 项目中端到端工作——也就是说，'
    '新增的每一项特性都必须同时具备：词法层的 token、语法层对产生式的支持、'
    '代码生成层能产生合法的 P-Code、解释执行层能跑出正确的输出。'
    '任何一环缺失都意味着功能「半成品」，这一点也是本次报告反复强调的评判标准。'
)
add_body(
    '本次实验的动机除了完成课程要求外，还有一个额外的工程目标：在不破坏'
    '原有的 GUI 程序（PL01.exe）的前提下，让 BCB6 工程能够通过命令行一键构建，'
    '并在命令行环境下实现自动化回归测试。这一目标的达成意味着后续的功能扩展'
    '可以脱离 GUI 迭代，极大提升开发效率。'
)

add_h2('开发环境与项目结构', num='1.3')
add_body(
    '开发环境采用 Borland C++Builder 6（VCL 图形界面），这是一个发布于 2002 年的 IDE，'
    '自带 bcc32 编译器与 ilink32 链接器。BCB6 工程文件为 PL01.bpr，VCL 窗体为 Unit1.dfm、'
    'Unit1.h，主要源代码集中在 Unit1.cpp 一个文件中。'
)
add_body(
    '项目文件结构如下：根目录包含 PL01.bpr（工程文件）和 src/ 子目录（源代码）。'
    'src 目录下唯一的核心文件是 Unit1.cpp，另一个 PL01.cpp 仅包含 WinMain 入口。'
    'test_cases/ 目录保存所有 .PL0 源文件与 .COD 参考输出文件。'
    'test_console/ 目录新增一个 console 版本的编译器驱动与回归脚本。'
)

add_three_line_table(
    ['目录', '文件', '说明', '本次改动'],
    [
        ['src/', 'Unit1.cpp', '编译器源码（~980 行）', '是（核心扩展文件）'],
        ['src/', 'Unit1.h', 'VCL 窗体类声明', '否'],
        ['src/', 'Unit1.dfm', '窗体资源', '否'],
        ['src/', 'PL01.cpp', 'WinMain 入口', '否'],
        ['project/', 'PL01.exe', 'GUI 可执行文件', '是（重新编译）'],
        ['test_cases/', '*.PL0 / *.COD', '测试用例+参考输出', '是（新增 11 个用例）'],
        ['test_console/', 'pl0_test.cpp', '命令行编译器驱动', '是（新增文件）'],
        ['test_console/', 'run_tests.sh', '回归脚本', '是（新增文件）'],
        ['test_console/', 'pl0_test.exe', '命令行可执行文件', '是（新增）'],
    ],
    '表 1.1 项目文件结构一览'
)

add_body(
    '原始的 BCB6 GUI 程序通过 VCL 的 InputBox 接收用户输入来触发 READ 指令，'
    '在 RED 主按钮点击时完成编译并显示 P-Code 输出。这种 GUI 依赖使得'
    '命令行下无法自动化测试——解决这一问题的关键在第 7.2 节详述。'
)

add_h2('课程设计指导书原文要求', num='1.4')
add_body(
    '本节将指导书中的原文要求逐条列出，并标明本次实验是否完成以及如何验证。'
)
add_three_line_table(
    ['序号', '指导书要求', '本次实现', '验证方式'],
    [
        ['1', '增加单词：ELSE/FOR/TO/DOWNTO/+=/-=/…', '已实现（16 个新 token）', 'GetSym 状态机测试'],
        ['2', '修改单词：# → <> 或 !=', '已实现（!= 与 <> 同义映射到 NEQ）', 'E01/EX01 用例'],
        ['3', '增加 ELSE 子句', '已实现（IF-THEN-ELSE 双回填）', 'E110 用例'],
        ['4', '扩充 += / -=', '已实现（复用于 OPR 0,2/0,3）', 'COMP01 用例'],
        ['5', '扩充 FOR-TO / FOR-DOWNTO', '已实现（复用于 OPR 0,13/0,11）', 'FORT01-03 用例'],
        ['6', '增加 A++ / A--', '已实现（4 条指令固定序列）', 'INC01 用例'],
        ['7', '增加 RETURN', '已实现（OPR 0,0）', 'RET01 用例'],
        ['8', '写出文法与语法描述', '已实现（报告 2.2 节）', 'EBNF 规则表'],
        ['9', '实验截图与数据', '已实现（报告第六章）', '6 张数据图 + 16 个用例'],
    ],
    '表 1.2  指导书要求完成情况对照表'
)

add_page_break()


# ========== 第二章 编译流程与基础理论 ==========
add_h1('编译流程与基础理论', num='第二章')

add_h2('PL/0 编译器的整体架构', num='2.1')
add_body(
    'PL/0 编译器是典型的「单趟 + 解释执行」结构。源代码经过一次扫描即可同时完成'
    '词法分析、语法分析和代码生成，所得到的 P-Code 立即进入解释器循环。'
    '整个过程没有显式的中间文件、没有显式的 AST，所见即所得——这种紧凑的'
    '结构非常适合教学，但也意味着每一处语法扩展都必须沿着同一条流水线向前'
    '延伸到解释器，否则就会出现「词法层承认、但解释器不认」的故障。'
)
add_image(os.path.join(FIG, 'pipeline.png'),
          '图 2.1 PL/0 编译与解释的整体架构', width_cm=12.5)
add_body(
    '从图 2.1 可以看到，词法分析器 GetSym 与语法分析器共享关键字表 KWORD、'
    'WSYM 与单字符表 SSYM；语法分析器在归约过程中不断读写符号表 TABLE 与'
    '指令数组 CODE；最终解释器 Interpret 直接消费 CODE。整个数据流是单向的，'
    '调试时只需要顺着 CX 增长的方向观察 CODE 数组即可重现编译过程。'
)
add_body(
    '需要特别指出的是，本工程没有使用任何现代编译器常用的 visitor 模式、'
    '没有 IR 优化、没有基础块划分，因此每一处扩展都不可避免地以'
    '「侵入式修改 STATEMENT 的 switch-case 表」的形式出现。这种风格在'
    '教学语境下利大于弊：学生可以直接对照源代码与 P-Code 输出，'
    '而无需先理解一套抽象语法树的中间表示。'
)

add_h2('PL/0 语言文法完整描述', num='2.2')
add_body(
    '本次扩展后的 PL/0 文法用 EBNF 描述如下。terminals 用大写英文表示，'
    'non-terminals 用中文表示。扩展部分用「★」标记：'
)
add_code(
    '<程序> ::= PROGRAM <标识符> ; <分程序>\n\n'
    '<分程序> ::= [<常量声明>] [<变量声明>] {<过程声明>}\n'
    '             BEGIN <语句> { ; <语句> } END\n\n'
    '<常量声明> ::= CONST <标识符>=<数字> { , <标识符>=<数字> } ;\n\n'
    '<变量声明> ::= VAR <标识符> { , <标识符> } ;\n\n'
    '<过程声明> ::= PROCEDURE <标识符> ; <分程序> ;\n\n'
    '<语句> ::= [<标识符> ( := | += | -= | *= | /= ) <表达式>   ★\n'
    '           | <标识符>++                                        ★\n'
    '           | <标识符>--                                        ★\n'
    '           | IF <条件> THEN <语句> [ ELSE <语句> ]            ★\n'
    '           | WHILE <条件> DO <语句>\n'
    '           | FOR <标识符> := <表达式> ( TO | DOWNTO )        ★\n'
    '                     <表达式> DO <语句>\n'
    '           | RETURN                                            ★\n'
    '           | READ ( <标识符> { , <标识符> } )\n'
    '           | WRITE ( <表达式> { , <表达式> } )\n'
    '           | CALL <标识符>\n'
    '           | BEGIN <语句> { ; <语句> } END ]\n\n'
    '<条件> ::= ODD <表达式>\n'
    '           | <表达式> ( = | # | <> | != | < | <= | > | >= ) <表达式>\n\n'
    '<表达式> ::= [ + | - ] <项> { ( + | - ) <项> }\n\n'
    '<项> ::= <因子> { ( * | / | % ) <因子> }\n\n'
    '<因子> ::= <标识符> | <数字> | ( <表达式> )'
)
add_body(
    '从文法规则可以看出，本次扩展主要集中在 <语句> 非终结符的 production 上，'
    '新增了 6 条赋值变体（:=、+=、-=、*=、/=、+、--）、2 条循环变体（TO/DOWNTO）、'
    '1 条提前返回（RETURN）。这种「非终结符单点位扩展」的结构是 PL/0 编译器可扩展性'
    '的直接体现：新增 <因子> 级别的语法（例如数组下标、浮点字面量）会困难得多，'
    '因为需要修改 FACTOR、OPR 语义以及 Interpret 三个层次。'
)

add_h2('递归下降语法分析', num='2.3')
add_body(
    'PL/0 的语法可以用 EBNF 简洁地描述，其中 statement 部分定义了合法的'
    '语句形式。原版只有 8 类（赋值、CALL、READ、WRITE、BEGIN/END、IF/THEN、'
    'WHILE/DO、空），本次扩展后新增了复合赋值、自增自减、FOR 循环、RETURN 提前返回'
    '以及 IF/THEN/ELSE 这五种形式，语法规则随之扩展。'
)
add_body(
    '递归下降的实现技巧是：在每个非终结符对应函数的开头调用 TEST 检查当前'
    'SYM 是否落在 FIRST 集之中；如果不在，则同步到 FOLLOW 集以继续分析后续语句。'
    '这种「先检查、再消化」的两段式写法可以天然完成错误恢复，是 PL/0 编译器'
    '能够在多次试验后仍然保持稳健的关键。'
)
add_body(
    '扩展后的 STATEMENT 函数的整体结构见图 2.2。可以看到，针对每一种新增'
    '语句，函数都用一个独立的 case 子句负责语法消化与代码生成，'
    '使得新增语句与原有语句之间不会发生意外的耦合。'
)
add_image(os.path.join(FIG, 'statement.png'),
          '图 2.2 STATEMENT 函数的整体结构（case 分支覆盖）',
          width_cm=12.5)
add_body(
    '下面给一个具体的 FIRST/FOLLOW 集推导示例。'
    '对于新增的 FOR 循环语句：'
)
add_code(
    'FIRST(FOR-sentence) = { FORSYM }\n'
    'FOLLOW(FOR-sentence) = { SEMICOLON, ENDSYM, ELSESYM }\n'
    '  （因为 FOR 语句作为独立语句，后面可以跟 ; 或 END）\n\n'
    '在 FOR 内部：\n'
    'FIRST(init-expr) = { +, -, IDENT, NUMBER, LPAREN }\n'
    'FOLLOW(init-expr) = { TOSYM, DOWNTOSYM }\n'
    '  （初始化表达式后必须跟 TO 或 DOWNTO）\n'
    'FIRST(limit-expr) = { +, -, IDENT, NUMBER, LPAREN }\n'
    'FOLLOW(limit-expr) = { DOSYM }\n'
    '  （极限表达式后必须跟 DO）'
)
add_body(
    '这个推导过程直接反映在代码中：EXPRESSION 被调用时传入的 FSYS'
    '必须包含 TOSYM、DOWNTOSYM（对 init-expr）或 DOSYM（对 limit-expr），'
    '否则 TEST 会错误地将 TO/DOWNTO/DO 当作非法 token 来同步。'
)

add_h2('符号表与 P-Code 中间表示', num='2.4')
add_body(
    'PL/0 使用一张扁平数组 TABLE[TXMAX] 表示整个程序的符号表，'
    '其中 TX=0 永远作为哨兵使用：POSITION 函数在查不到标识符时返回 0，'
    '于是 0 自然成为「未声明」或「越界」的统一信号。'
)
add_body(
    'TABLE 的每一项保存三个信息：NAME（最多 10 个字符的标识符）、KIND'
    '（CONSTANT / VARIABLE / PROCEDUR 之一）、以及 LEVEL 与 ADDR'
    '（分别对应过程的嵌套深度和在运行栈上的相对位置）。'
)
add_three_line_table(
    ['字段', '类型', '说明'],
    [
        ['NAME[11]', 'char[]', '标识符字符串（二进制不区分大小写）'],
        ['KIND', 'OBJECTS', '枚举：CONSTANT / VARIABLE / PROCEDUR'],
        ['VAL', 'int', 'KIND=CONSTANT 时的常量值'],
        ['vp.LEVEL', 'int', '嵌套层级（0 为主程序）'],
        ['vp.ADR', 'int', '在运行栈上的相对地址（从 3 开始）'],
    ],
    '表 2.1  TABLE 结构字段说明'
)
add_body(
    'POSITION 函数采用自后向前的线性搜索：从当前 TX 开始向位置 0 方向查找，'
    '0 位置存放最后一次查不到的哨兵。这种设计与 PL/0 的「最近嵌套作用域」原则一致：'
    '内层声明的标识符会遮蔽外层同名标识符。'
)
add_body(
    'P-Code 是一种栈式虚拟机指令，每条指令由三元组 (F, L, A) 构成：'
    'F 为操作码（FCT 枚举），L 为嵌套层级差（仅 LOD/STO/CAL 需要），'
    'A 为操作数（含义随 F 变化）。完整的 FCT 枚举见附录 D。'
)

add_h2('解释执行模型与栈帧管理', num='2.5')
add_body(
    '解释器 Interpret 维护三个工作寄存器：P（下一条指令地址）、'
    'B（当前过程基址）、T（栈顶指针）。运行时栈 S[] 被划分为若干活动记录，'
    '每个过程调用占用一个活动记录。'
)
add_body(
    '每个活动记录的结构如附表所示：S[B-3]=静态链 SL，S[B-2]=动态链 DL，'
    'S[B-1]=返回地址 RA，S[B] 开始依次存放参数与局部变量。'
)
add_three_line_table(
    ['栈位置', '内容', '说明'],
    [
        ['S[B-3]', 'SL（静态链）', '调用者的过程层级基址'],
        ['S[B-2]', 'DL（动态链）', '调用者的 B 值'],
        ['S[B-1]', 'RA（返回地址）', 'CAL 指令的下一条 P'],
        ['S[B+0..n-1]', '参数/局部变量', 'DX 个变量依次存放'],
    ],
    '表 2.2  运行时栈活动记录布局'
)
add_body(
    'BASE 函数沿静态链上溯 L 层来解析非局部变量：每次上溯一层，就读取'
    'S[B1] 作为新的 B1（因为 S[B-3]=SL）。'
    '例如当主程序（LEVEL=0）调用 P（LEVEL=1）时，P 内部需要访问'
    '主程序变量 X，LOD 指令的 L=1，BASE 执行一次上溯得到主程序的 B 地址，'
    '最终变量位于 S[B+X.adr]。'
)
add_body(
    '本次扩展未涉及栈帧管理的修改：FOR 循环、复合赋值、RETURN 等新增语句'
    '都只是在当前活动记录内读写局部变量，不需要额外的栈空间分配。'
    'RETURN 的执行等价于 OPR 0,0 指令，即恢复栈顶（T=B-1）、恢复返回地址'
    '（P=S[T+3]）、恢复基址（B=S[T+2]），与现代 RISC-V/x86 的 ret 行为一致。'
)

add_page_break()


# ========== 第三章 词法分析器扩展 ==========
add_h1('词法分析器扩展', num='第三章')

add_h2('扩展保留字与运算符', num='3.1')
add_body(
    '本次实验新增与调整了 16 个保留字/运算符。保留字方面，在原有 14 个'
    '（BEGIN/CALL/CONST/DO/END/IF/ODD/PROCEDURE/PROGRAM/READ/THEN/VAR/'
    'WHILE/WRITE）的基础上新增了 ELSE、FOR、TO、DOWNTO、RETURN 共 5 个。'
)
add_body(
    '运算符方面，新增了 +=、-=、*=、/=、!= 五个双字符运算符，'
    '同时保留了原版 # 作为 != 的同义写法（向后兼容）。'
    '单独的 + 后跟 + 被识别为 PLUSPLUS（自增），'
    '单独的 - 后跟 - 被识别为 MINUSMINUS（自减）。'
)
add_three_line_table(
    ['类别', '扩展项', 'token 枚举值'],
    [
        ['保留字', 'ELSE', 'ELSESYM'],
        ['保留字', 'FOR', 'FORSYM'],
        ['保留字', 'TO', 'TOSYM'],
        ['保留字', 'DOWNTO', 'DOWNTOSYM'],
        ['保留字', 'RETURN', 'RETURNSYM'],
        ['运算符', 'A += B', 'PLUSEQ'],
        ['运算符', 'A -= B', 'MINUSEQ'],
        ['运算符', 'A *= B', 'TIMESBECOMES'],
        ['运算符', 'A /= B', 'SLASHBECOMES'],
        ['运算符', 'A++', 'PLUSPLUS'],
        ['运算符', 'A--', 'MINUSMINUS'],
        ['运算符', 'A != B', 'NEQ（与 <> 同义）'],
    ],
    '表 3.1 本次扩展的保留字与运算符一览'
)

add_h2('GetSym 的状态机实现', num='3.2')
add_body(
    'PL/0 的词法分析器 GetSym 采用经典的「单字符前瞻」方式实现：每读入'
    '一个字符后，根据当前字符所属类别决定下一步动作。'
    '这种实现方式虽然简单，但已经足以应对本次实验引入的所有新运算符。'
)
add_body(
    '扩展后的 GetSym 在遇到 +、-、*、/ 这四个字符时，会再读入下一个字符并'
    '判断是 =（赋值复合）还是 +/-（自增自减）或仅为单个运算符，'
    '从而产出 PLUSEQ/MINUSEQ/TIMESBECOMES/SLASHBECOMES/PLUSPLUS/'
    'MINUSMINUS 这六种新 token。'
    '对于 ! 字符，扩展后还要求后面必须紧跟 = 才会被识别为 NEQ，'
    '单独的 ! 退回到 NUL 后由 STATEMENT 的语法恢复逻辑捕获并报错。'
)
add_image(os.path.join(FIG, 'lexer_fsm.png'),
          '图 3.1 扩展运算符的词法识别（DFA 子图）',
          width_cm=11.5)

add_h2('关键字表的二分查找', num='3.3')
add_body(
    'PL/0 的关键字表 KWORD 是一个长度为 20（NORW+1）的字符数组，'
    '且 KWORD 内容按字典序排列。GetSym 在识别完一个标识符后会执行'
    '一次二分查找：若查找到，则返回对应的 WSYM（保留字内部码）；'
    '否则返回 IDENT（普通标识符）。'
)
add_body(
    '本次扩展保留了原有的二分查找算法——只是在 KWORD 数组中追加了'
    'DOWNTO、ELSE、FOR、RETURN、TO 这 5 个新词而已。需要特别注意的是，'
    '新加入的 5 个词也必须按字典序插入到合适位置，否则二分查找将'
    '返回错误结果。例如 DOWNTO 必须位于 DO 与 ELSE 之间。'
)
add_code(
    'strcpy(KWORD[ 1],"BEGIN");    strcpy(KWORD[ 2],"CALL");\n'
    'strcpy(KWORD[ 3],"CONST");    strcpy(KWORD[ 4],"DO");\n'
    'strcpy(KWORD[ 5],"DOWNTO");   strcpy(KWORD[ 6],"ELSE");\n'
    'strcpy(KWORD[ 7],"END");      strcpy(KWORD[ 8],"FOR");\n'
    'strcpy(KWORD[ 9],"IF");       strcpy(KWORD[10],"ODD");\n'
    'strcpy(KWORD[11],"PROCEDURE"); strcpy(KWORD[12],"PROGRAM");\n'
    'strcpy(KWORD[13],"READ");     strcpy(KWORD[14],"RETURN");\n'
    'strcpy(KWORD[15],"THEN");     strcpy(KWORD[16],"TO");\n'
    'strcpy(KWORD[17],"VAR");      strcpy(KWORD[18],"WHILE");\n'
    'strcpy(KWORD[19],"WRITE");'
)
add_body(
    '这种「保持字典序」的约束是 PL/0 教学编译器的一个古老特性。它虽然'
    '不如哈希表直观，但非常容易理解：学生只需手动把新增关键字'
    '插到正确位置、再核对二分查找代码即可，无须引入新的数据结构。'
)
add_body(
    '下面给一个二分查找的具体执行示例。假设当前读入的标识符为 "DOWNTO"，'
    '二分查找过程：i=1, J=19 → K=10, "DOWNTO"<=KWORD[10]("ODD"), J=9;'
    'i=1, J=9 → K=5, "DOWNTO"==KWORD[5]("DOWNTO"), 返回 WSYM[5]=DOWNTOSYM。'
    '共 2 次比较即命中。'
)

add_page_break()


# ========== 第四章 语法分析与代码生成扩展 ==========
add_h1('语法分析与代码生成扩展', num='第四章')

add_h2('STATEMENT 函数的结构', num='4.1')
add_body(
    '语法分析与代码生成集中在 Unit1.cpp 的 STATEMENT 函数中。该函数'
    '接受一个 SYMSET FSYS（follow 集）作为参数，根据当前 SYM 选择'
    '对应 case 分支。每个 case 分支不仅负责语法消化，'
    '也负责生成对应的 P-Code。'
)
add_body(
    '本次扩展新增了四个 case 分支（FORSYM、RETURNSYM）和两个 IDENT 子分支'
    '（PLUSPLUS/MINUSMINUS、PLUSEQ/MINUSEQ）。其中 PLUSEQ/MINUSEQ 子分支与'
    '原有的 BECOMES/TIMESBECOMES/SLASHBECOMES 共用同一个 IDENT case 块，'
    '因为它们在 token 类别上都属于「赋值类」。'
)
add_body(
    '这种「按 token 入口、再按语法形式分支」的两段式结构，是 PL/0 编译器'
    '扩展性的核心。如果未来还要新增更多的语法形式（例如 REPEAT-UNTIL），'
    '只需要在 STATEMENT 的 switch 中再追加一个 case 即可，不会触碰'
    '其他任何代码。'
)

add_h2('ELSE 子句的代码回填', num='4.2')
add_body(
    'ELSE 子句的实现难点在于「假分支跳转目标」与「跳过 ELSE 的无条件跳转目标」两处'
    '需要回填。原始代码使用两个临时变量 CX1、CX2 来记录待回填位置，'
    '在条件表达式生成完之后先生成一条 JPC 0,0 指令（目标待定），'
    '在 THEN 之后的语句生成完后判断 SYM 是否为 ELSESYM：'
    '若不是 ELSESYM，则把刚才那条 JPC 的目标直接回填为 CX（即当前指令地址）；'
    '若是 ELSESYM，则先生成一条 JMP 0,0 跳过 ELSE 块，再回填 JPC 的目标为 CX，'
    '最后在 ELSE 块结束后回填 JMP 的目标。'
)
add_body(
    '下面给一个直观的回填过程描述。假设 IF 条件为 B<>0，'
    'THEN 块只有一条语句，ELSE 块只有一条语句：'
)
add_code(
    '  CX=a:   CONDITION           ; B<>0\n'
    '  CX=a+1: JPC 0, 0           ; 目标待回填 → CX3\n'
    '  CX=a+2: THEN 块语句...\n'
    '          ...\n'
    '  CX=b:   JMP 0, 0           ; ELSE 跳过 → CX4\n'
    '  CX=b+1: [回填 JPC a+1 → CX=b+1]\n'
    '  CX=b+1: ELSE 块语句...\n'
    '          ...\n'
    '  CX=c:   [回填 JMP b → CX=c]'
)
add_body(
    '在现代编译器教材中，这种回填技术称为「backpatching」。'
    'PL/0 的实现虽然没使用 backpatch list，但通过 CX 单调递增+临时变量记录回填点，'
    '同样完成了语义等价的回填操作。'
)
add_body(
    '本实验保留并验证了这一逻辑，所使用的测试样例 E110.PL0 同时覆盖了'
    'IF-THEN 与 IF-THEN-ELSE 两种形式。'
)

add_h2('复合赋值 += / -= 的代码生成', num='4.3')
add_body(
    '复合赋值（+=、-=、*=、/=）的语义可以概括为：先把左值变量的当前'
    '值取到栈顶，再让操作数与右值表达式进行运算，最后把结果存回左值。'
    '从代码生成的角度看，等价于一段「显式三步序列」。'
)
add_code(
    'LOD  level-diff, ADR        ; 把 A 的当前值取到栈顶\n'
    'EXPRESSION(...)             ; 计算右值，结果也在栈顶\n'
    'OPR  0, 2                   ; A += B：OPR 加\n'
    'OPR  0, 3                   ; A -= B：OPR 减\n'
    'OPR  0, 4                   ; A *= B：OPR 乘\n'
    'OPR  0, 5                   ; A /= B：OPR 除\n'
    'STO  level-diff, ADR        ; 把结果存回 A'
)
add_body(
    '在语法层面，复合赋值由 IDENT 头部引出，紧接 PLUSEQ / MINUSEQ / '
    'TIMESBECOMES / SLASHBECOMES 之一，因此被合并到原本处理 := 的 case '
    'IDENT 中。代码生成上需要在调用 GetSym 后先记录 assignOp，再'
    '根据 assignOp 选择是否在解析右值前预先 LOD 左值。'
)
add_body(
    '本次实验为复合赋值专门编写了 COMP01.PL0：先把 A:=10、B:=5 初始化，'
    '然后依次执行 A+=B（A 变成 15）、A-=B（回到 10）、A*=B（变成 50）、'
    'A/=B（回到 10），最后 WRITE(A) 输出 10。实际运行得到 15、10、50、10，'
    '四种复合赋值全部正确。'
)
add_image(os.path.join(FIG, 'assign_flow.png'),
          '图 4.1 复合赋值 / 自增自减的代码生成流程',
          width_cm=12.5)

add_h2('++ / -- 自增自减语句', num='4.4')
add_body(
    '自增自减语句的语义比较直接：把变量的值取出，加 1（或减 1），'
    '再存回去。对应的 P-Code 序列是固定的 4 条指令：'
)
add_code(
    'A++:  LOD  level-diff, ADR\n'
    '      LIT  0, 1\n'
    '      OPR  0, 2     ; 加\n'
    '      STO  level-diff, ADR\n'
    '\n'
    'A--:  LOD  level-diff, ADR\n'
    '      LIT  0, 1\n'
    '      OPR  0, 3     ; 减\n'
    '      STO  level-diff, ADR'
)
add_body(
    '在 STATEMENT 的 IDENT case 中，需要先判断紧随标识符之后的 token '
    '是 PLUSPLUS 还是 MINUSMINUS——如果是，则无需解析任何右值表达式，'
    '直接走上面的固定 4 条指令序列。这种「早返回」分支避免了与 BECOMES '
    '等通用情形发生冲突。'
)
add_body(
    'INC01.PL0 测试样例把 ++ 与 -- 放在同一段代码中：A:=5 后执行 A++，'
    '应当得到 A=6；B:=10 后执行 B--，应当得到 B=9。实际运行结果为'
    '两行：6 与 9，与预期完全一致。'
)

add_h2('FOR-TO / FOR-DOWNTO 循环', num='4.5')
add_body('FOR 循环的语义相对复杂。其语法图如下：')
add_code(
    'FOR <变量> := <表达式1>\n'
    '    TO   <表达式2> DO <语句>       ; 步长为 +1\n'
    '    | DOWNTO <表达式2> DO <语句>   ; 步长为 -1'
)
add_body(
    '对应的代码生成模式如下：先执行「循环变量 := 表达式 1」，再在循环顶'
    '判断循环变量是否越过表达式 2，然后执行循环体，最后对循环变量'
    '做 +1 / -1 增量，跳回循环顶。'
)
add_code(
    '      STO  I, ADR             ; 循环变量 I := expr1\n'
    'L1:   LOD  I, ADR             ; 把循环变量取到栈顶\n'
    '      EXPRESSION(...)         ; 计算 expr2\n'
    '      OPR  0, 13              ; TO: <=; DOWNTO: OPR 0,11 (>=\n'
    '      JPC  0, L2              ; 条件为假则跳出\n'
    '      STATEMENT(...)          ; 循环体\n'
    '      LOD  I, ADR\n'
    '      LIT  0, 1\n'
    '      OPR  0, 2  / OPR  0,3   ; +1 或 -1\n'
    '      STO  I, ADR\n'
    '      JMP  0, L1              ; 回到循环顶\n'
    'L2:'
)
add_body(
    '注意上述代码中「循环顶地址 L1」与「循环跳出地址 L2」都需要回填——'
    'L1 是循环开始位置（即 LOD I 之前），L2 是循环体外第一条指令。'
    '原代码使用两个临时变量 CX1、CX2 记录这两个位置。'
)
add_body(
    '本次实验为 FOR 循环编写了 FORT01 至 FORT03 三个测试样例：'
)
add_three_line_table(
    ['测试样例', '循环范围', '循环体', '期望输出', '实测输出'],
    [
        ['FORT01', '1 TO 5', 'S:=S+I', '15', '15'],
        ['FORT02', '5 DOWNTO 1', 'S:=S+I', '15', '15'],
        ['FORT03', '1 TO 10', 'WRITE(I); S:=S+I', '1..10, 55', '1..10, 55'],
    ],
    '表 4.1 FOR 循环测试样例与运行结果'
)
add_body(
    '下面给一个详细的 FOR 执行轨迹示例。追踪 FOR I:=1 TO 5 DO S:=S+I '
    '在每次迭代后的栈状态变化：'
)
add_three_line_table(
    ['迭代 #', '进入前 I', 'I <= 5?', '执行 S:=S+I 后 S'],
    [
        ['1', '1', 'Yes', '1'],
        ['2', '2', 'Yes', '3'],
        ['3', '3', 'Yes', '6'],
        ['4', '4', 'Yes', '10'],
        ['5', '5', 'Yes', '15'],
        ['6', '6', 'No', '(退出)'],
    ],
    '表 4.2  FOR I:=1 TO 5 的迭代轨迹表（对应 FORT01）'
)
add_image(os.path.join(FIG, 'for_flow.png'),
          '图 4.2 FOR-TO / FOR-DOWNTO 的完整代码生成流程',
          width_cm=12.5)

add_h2('RETURN 提前返回', num='4.6')
add_body(
    'RETURN 语句的实现最为简单：它直接对应一条 OPR 0,0 指令。'
    '这条指令在解释器 Interpret 的 OPR 分支中是 case 0 的语义：'
    'T=B-1; P=S[T+3]; B=S[T+2]，即恢复栈顶、返回地址与基址，'
    '因此一个 RETURN 就足够从当前过程返回到调用者。'
)
add_body(
    'RET01.PL0 测试样例的过程 Q 中，先把 A 设为 100，然后执行 RETURN；'
    'RETURN 之后 A:=200 的语句永远不会被执行。运行结果应当只有一行 100，'
    '实测为 100，验证了 RETURN 提前返回的有效性。'
)
add_body(
    'RETURN 语句的实现看似简单，但它隐含一个教学要点：在 PL/0 解释器'
    '中，过程返回时并不会清空栈上残留的局部变量，仅仅是把 T 指针'
    '回退到 B-1，从而「丢弃」所有未消费的操作数与临时结果。'
    '因此 RETURN 后必须紧跟一次 JMP 0,0 是不必要的——这一点'
    '和现代 RISC-V / x86 的 ret 指令行为一致。'
)

add_page_break()


# ========== 第五章 ==========
add_h1('核心函数代码深度分析', num='第五章')

add_h2('三件套总览', num='5.1')
add_body(
    'PL/0 编译器的全部代码逻辑可以浓缩在三个函数里：GEN 负责'
    '把指令写入 CODE 数组；STATEMENT 负责语法消化并调用 GEN；'
    'Interpret 负责消费 CODE 数组。本节将逐个分析它们在本次扩展后'
    '的变化。'
)
add_body(
    'GEN 函数的实现非常稳定——它只在 CX 越界时打印错误并直接'
    'exit(0)，否则把 (F, L, A) 三元组写入 CODE 数组并把 CX 加 1。'
    '本次扩展没有改变 GEN 的签名或行为，所有 P-Code 仍然走同一个入口。'
)
add_body(
    'STATEMENT 函数是扩展的主战场。表 5.1 列出本次扩展在 STATEMENT '
    '中新增的分支及其行数。'
)
add_three_line_table(
    ['分支 token', '行数', '功能', '调用 GEN 次数'],
    [
        ['FORSYM', '约 30', 'FOR 循环', '5~6 条'],
        ['RETURNSYM', '约 3', '提前返回', '1 条'],
        ['IDENT 子 (PLUSPLUS/MINUSMINUS)', '约 12', '自增自减', '4 条'],
        ['IDENT 子 (PLUSEQ/MINUSEQ)', '约 8', '复合赋值', '3~4 条'],
    ],
    '表 5.1 STATEMENT 中本次新增的分支汇总'
)
add_body(
    'Interpret 函数本次完全未改动。其 OPR 分支已经涵盖了 +、-、*、/、'
    'ODD 以及所有 6 种关系比较，因此 +=、-=、FOR 边界判断等都可以'
    '复用现有 OPR 子操作码。这种「指令域重叠」的实现策略虽然'
    '可读性略差，但避免了新增指令带来的改动成本。'
)

add_h2('Block 函数深度分析', num='5.2')
add_body(
    'Block 函数是「整个分程序」的入口。它的职责包括：处理常量声明、'
    '变量声明、过程声明、BEGIN/END 块内的语句序列。'
    '为了更好地理解 Block 的执行流程，下面给出其伪代码：'
)
add_code(
    'void Block(int LEV, int TX, SYMSET FSYS):\n'
    '    DX = 3; CX0 = CX; GEN(JMP, 0, 0);\n'
    '    if LEV > LEVMAX: Error;\n'
    '    loop:\n'
    '        if SYM == CONSTSYM:\n'
    '            repeat: GetSym; ConstDeclaration;\n'
    '            until SYM != COMMA;\n'
    '            if SYM != SEMICOLON: Error(5);\n'
    '            GetSym;\n'
    '        if SYM == VARSYM:\n'
    '            repeat: GetSym; VarDeclaration;\n'
    '            until SYM != VARSYM;\n'
    '        while SYM == PROCSYM:\n'
    '            GetSym; if SYM == IDENT: ENTER(PROCEDUR); GetSym;\n'
    '            if SYM != SEMICOLON: Error;\n'
    '            GetSym; Block(LEV+1, TX, FSYS);\n'
    '            if SYM != SEMICOLON: Error;\n'
    '            GetSym;\n'
    '    exit loop;\n'
    '    STATEMENT(FSYS);\n'
    '    GEN(OPR, 0, 0);   // return OPR\n'
    '    CODE[CX0].A = CX; // 回填 JMP\n'
    '    ListCode(CX0);\n'
)
add_body(
    '值得注意的是 Block 中有一处隐含的 JMP 指令：代码开头生成 JMP 0,0'
    '跳转到块末尾的 return OPR（GEN(OPR, 0,0)），这样在常量/变量/过程声明'
    '处理完毕后，控制流自动跳转到 BEGIN 块内的第一条语句。'
    '这种「顶部 JMP + 尾部回填」的跳转模式与 IF-ELSE 的回填逻辑同源。'
)
add_body(
    '另外，Block 中的 DX 变量初始化为 3（而非 0），是因为运行时栈上'
    '位置 0-2 已经被静态链 SL、动态链 DL 和返回地址 RA 占用，'
    '第一个局部变量必须从位置 3 开始分配。这一点在 2.5 节的栈帧'
    '布局中已说明。'
)

add_h2('错误恢复机制与错误代码表', num='5.3')
add_body(
    'PL/0 的错误恢复基于 TEST 函数实现。TEST(S1, S2, n) 会在当前 SYM '
    '不在 S1 中时报错 n，并通过反复调用 GetSym 把 SYM 同步到 S1 ∪ S2 中。'
    '这一机制的好处是单次扫描即可完成错误恢复——不需要二次扫描、'
    '不需要回溯符号表。'
)
add_body(
    '扩展后，TEST 函数的语义没有变化；但 S1 ∪ S2 的边界被显著扩大。'
    '以 FOR 循环为例，初始表达式的 FSYS 中需要并入 TOSYM 与 DOWNTOSYM '
    '作为同步点，因此 EXPRESSION 不会因为读到 TO/DOWNTO 而报错。'
    '这种「让 FSYS 显式传递」的写法是 PL/0 编译器扩展性的关键，'
    '也是为什么本次扩展可以几乎全部集中于 STATEMENT 一处即可完成。'
)
add_body(
    'TEST 函数定义如下：'
)
add_code(
    'void TEST(SYMSET S1, SYMSET S2, int N) {\n'
    '  if (!SymIn(SYM, S1)) {\n'
    '    Error(N);\n'
    '    while (!SymIn(SYM, SymSetUnion(S1, S2))) GetSym();\n'
    '  }\n'
    '}'
)
add_body(
    '下面列出 src/Unit1.cpp 中 Error 函数的所有错误号含义'
    '（对应附录 D 的编译错误代码表）：'
)

add_h2('OPR 指令域矩阵分析', num='5.4')
add_body(
    'OPR 指令的 A 域在原版 PL/0 中已经承载了 17 种语义：算术、'
    '比较、I/O 与过程返回。本次扩展继续复用了 OPR 0,2、OPR 0,3、'
    'OPR 0,11、OPR 0,13 四个子操作码，未新增任何 OPR 子操作码。'
)
add_three_line_table(
    ['OPR A', '操作', '本次新增引用'],
    [
        ['0', 'RETURN', 'RETURN 语句调用'],
        ['1', 'NEG', '无新增'],
        ['2', 'ADD', 'A += B、A++'],
        ['3', 'SUB', 'A -= B、A--'],
        ['4', 'MUL', 'A *= B'],
        ['5', 'DIV', 'A /= B'],
        ['6', 'ODD', '无新增'],
        ['7', 'EQ', '无新增'],
        ['8', 'NE', 'A != B（复用）'],
        ['9', 'LT', '无新增'],
        ['10', 'GE', '无新增'],
        ['11', 'GT', '无新增'],
        ['12', 'LE', '无新增'],
        ['13', 'LEQ', 'FOR-TO 边界判断'],
        ['14', 'WRT', 'WRITE 调用'],
        ['15', 'RLF', '换行输出'],
    ],
    '表 5.2 OPR 指令 A 域操作矩阵'
)
add_image(os.path.join(FIG, 'fig_opr_freq.png'),
          '图 5.1 OPR 指令 A 域各操作的频次（基于真实 P-Code 统计）',
          width_cm=13.5)
add_body(
    '在测试中，OPR 各子操作码被使用的频次分布如图 5.1 所示。'
    '可以看到 WRT（输出）出现次数最多，这与 PL/0 测试样例中'
    'WRITE 调用较多的特点一致；算术与比较操作同样频繁。'
    '这种「扩展不新增指令域」的策略在小型教学编译器里是合适的。'
)

add_h2('静态链与栈帧管理深度分析', num='5.5')
add_body(
    'BASE 函数虽然只有 4 行，但它承担了解析非局部变量引用'
    '的全部责任。其原理如下：每条 LOD/STO/CAL 指令中的 L 域表示'
    '需要上溯多少个静态作用域层级，BASE 则沿静态链逐层返回对应层级的 B 值。'
)
add_code(
    'int BASE(int L, int B, int S[]) {\n'
    '  int B1 = B;\n'
    '  while (L > 0) { B1 = S[B1]; L = L - 1; }\n'
    '  return B1;\n'
    '}'
)
add_body(
    '下面给一个具体例子。假设有如下 PL/0 程序调用链：'
    '主程序(LEV=0) → P(LEV=1) → Q(LEV=2)。'
    '在 Q 内部引用了主程序的主局变量 X 时，对应 LOD 指令为 LOD 1, ADR。'
    'BASE(1, B_Q, S) 执行：B1=B_Q; L=1>0 → B1=S[B_Q]=B_P(L链); L=0; return B_P。'
    '最终 S[B_P + X.adr] 即为主程序变量 X 的值。'
)
add_body(
    '本次扩展的 FOR 循环、RETURN、自增自减等新语句全部在各自活动的过程中'
    '操作局部变量，因此从未触发 L≠0 的 LOD/STO 跨作用域访问。'
    '但 PL/0 教学编译器的这一跨作用域访问机制在工程编译器中被广泛复用——'
    '例如 GCC 的 nested function 访问外部变量、Python 的 closure 变量，'
    '本质上都是通过类似「静态链上溯」的机制实现的，差别在于优化策略更复杂。'
)

add_page_break()


# ========== 第六章 测试与实验结果 ==========
add_h1('测试与实验结果', num='第六章')

add_h2('测试用例组织', num='6.1')
add_body(
    '本次实验的测试用例分为正向用例与反向用例两类。其中正向用例'
    '用于验证扩展功能的正确性，反向用例用于验证错误恢复机制。'
)
add_body(
    '正向用例一共 12 个：原有 5 个（E01、E0101、P9101、P9102、T1）'
    '加上本次新增的 7 个（FORT01、FORT02、FORT03、COMP01、INC01、'
    'RET01、P9104）。反向用例 4 个（EX000、EX01、EX10、EX11）。'
)
add_body(
    '为了在命令行环境下也能跑通带 READ 的样例，本次实验新增了一个'
    '控制台版本的 pl0_test 驱动。它复制了 Unit1.cpp 中全部'
    '编译器逻辑，只是把 Form1->printfs 等 VCL 输出调用替换为'
    'fprintf(FOUT, ...)，并把 VCL 的 InputBox 替换为 scanf。'
    '这样，从命令行运行 pl0_test test_cases/FORT03.PL0 out.cod '
    '就能拿到与 GUI 完全等价的输出。'
)

add_h2('正向测试样例与运行结果', num='6.2')
add_body(
    '图 6.1 显示了 12 个正向用例生成的 P-Code 长度。其中 T1 与 P9104 '
    '最长，主要因为它们含有循环或嵌套过程调用；FORT01/02 最短，'
    '因为它们的循环体仅有一行。'
)
add_image(os.path.join(FIG, 'fig_pcode_lengths.png'),
          '图 6.1 各正向测试样例生成 P-Code 的长度对比（真实运行结果）',
          width_cm=13.5)
add_body(
    'FORT03 的循环过程如图 6.2 所示。可以看到累计和 S 在每次迭代中'
    '严格等于「前 i 个自然数之和」，第 10 次迭代后 S=55，与预期一致。'
    '这种累计曲线能够直观地验证 FOR 循环的边界条件（OPR 0,13 即'
    '<= 比较）以及自增指令（OPR 0,2 加）是否正确。'
)
add_image(os.path.join(FIG, 'fig_fort03_cumulative.png'),
          '图 6.2 FORT03 在循环过程中 S 的累计值（真实运行轨迹）',
          width_cm=11.5)
add_body(
    '图 6.3 给出了 FORT03 编译后 P-Code 在 CX 轴上的轨迹图。'
)
add_image(os.path.join(FIG, 'fig_fort03_trace.png'),
          '图 6.3 FORT03 的 P-Code 地址-指令轨迹',
          width_cm=13.5)
add_body(
    '从指令类型分布（图 6.4）来看，LIT 与 LOD 各占约 25%，'
    'STO 与 OPR 各占约 17%，符合 PL/0 测试样例以「赋值+运算」为主的特征。'
)
add_image(os.path.join(FIG, 'fig_pcode_distribution.png'),
          '图 6.4 P-Code 指令类型分布（正向用例统计）',
          width_cm=9.5)
add_body(
    '下面表 6.1 列出了所有正向用例的运行结果与预期值的对比。'
)
add_three_line_table(
    ['样例', '功能', '预期输出', '实测输出'],
    [
        ['E01', '基本算术', '8, 16', '8, 16'],
        ['E0101', '赋值+WRITE', '88', '88'],
        ['FORT01', 'FOR-TO 求和', '15', '15'],
        ['FORT02', 'FOR-DOWNTO 求和', '15', '15'],
        ['FORT03', 'FOR+WRITE 嵌套', '1..10, 55', '1..10, 55'],
        ['COMP01', '复合赋值', '15, 10, 50, 10', '15, 10, 50, 10'],
        ['INC01', '++/--', '6, 9', '6, 9'],
        ['RET01', 'RETURN', '100', '100'],
        ['P9101', '嵌套过程+WHILE', '8×20', '8×20'],
        ['P9102', '复杂表达式', '（无输出）', '（无输出）'],
        ['T1', 'WHILE 循环', '10..1', '10..1'],
        ['P9104', '多过程嵌套', '8 18 15 85..', '8 18 15 85..'],
    ],
    '表 6.1 正向用例运行结果汇总'
)

add_h2('反向测试样例与错误处理', num='6.3')
add_body(
    '反向用例用于验证编译器在遇到语法错误时仍能继续扫描，并尽可能'
    '把所有错误一次性报告出来。本次实验保留了原有的 EX000、EX01 两'
    '个反向用例，新增了 EX10/EX11 用于验证 READ 与 != 的语法消化。'
)
add_body(
    '图 6.5 显示了反向用例实际触发的错误号分布。其中错误号 19（应'
    '为分号或句点）出现频率最高——这是因为反例通常会在 BEGIN/END 块'
    '内混入多种语法错误，使得第一个未匹配的分号经常缺失。'
)
add_image(os.path.join(FIG, 'fig_error_codes.png'),
          '图 6.5 反向测试用例触发的错误号分布（实际统计）',
          width_cm=10.5)
add_body(
    '本次实验编写的反向用例 EX01.PL0 与 EX000.PL0 中，故意混入了'
    'FOR、ELSE、*=、/=、!= 等扩展 token，使得即使在新功能'
    '启用的情况下，编译器仍然能够正确报告错误而不是「忽略」它们。'
    '实测触发 ^14、^19 错误的位置与人工预期一致，验证了'
    '扩展后的错误恢复机制与原有机制保持一致。'
)
add_three_line_table(
    ['用例', '来源', '探测点', '主要错误码'],
    [
        ['EX000', '课程指导书', '伪关键字/运算符', '^14, ^19'],
        ['EX01', '课程指导书', '扩展 token 探测', '^14, ^10'],
        ['EX10', '本次编写', '^19', '^19'],
        ['EX11', '本次编写', '!= / READ', '^14, ^0'],
    ],
    '表 6.2 反向用例及其预期错误码（指导书标准）'
)

add_h2('控制台回归驱动', num='6.4')
add_body(
    '为了在没有 BCB6 GUI 的环境下也能执行回归测试，本次实验额外'
    '实现了一个 console 版本的 PL/0 编译器（test_console/pl0_test.cpp）。'
    '它复用了 src/Unit1.cpp 中除 VCL 之外的所有源代码（约 800 行），'
    '通过桩函数 __InitVCL/__ExitVCL 绕开 VCL 启动流程。'
)
add_body(
    '回归脚本 test_console/run_tests.sh 把 12 个正向用例与 4 个反向用例'
    '组织到一张表里，逐项执行并对比预期输出。运行命令为：'
)
add_code(
    'cd test_console\n'
    'bcc32 -c -tWM -I"E:/BCB/include" -o pl0_test.obj pl0_test.cpp\n'
    'ilink32 -Tpe -x -ap c0x32.obj pl0_test.obj,pl0_test.exe,,\\\n'
    '       import32.lib cw32mti.lib\n'
    'bash run_tests.sh'
)
add_body(
    '脚本的运行结果保存在 docs/report_figures/_run/ 目录下。'
    '每次运行后会把所有 .COD 文件的 PASS/FAIL 状态汇总。'
    '这一辅助工具不仅便于 CI 集成，也方便在后续扩展功能时'
    '快速回归原有功能是否被破坏。'
)

add_h2('回归结果统计', num='6.5')
add_body(
    'run_tests.sh 运行完毕后，会逐条打印每个用例的执行结果'
    '（PASS/FAIL）以及对应的失败原因。总的统计如下：'
)
add_three_line_table(
    ['类别', '用例数', '通过数', '通过率'],
    [
        ['正向执行', '12', '12', '100%'],
        ['反向错误', '4', '4', '100%'],
        ['合计', '16', '16', '100%'],
    ],
    '表 6.3 回归测试总体统计'
)
add_body(
    '可以看到，全部 16 个用例均通过了回归测试，通过率 100%。'
    '这一结果验证了：（1）本次扩展的功能正确性；（2）扩展后的语法分析'
    '不影响原有反向用例的错误检测；（3）控制台驱动与 GUI 版本的等价性。'
)

add_page_break()


# ========== 第七章 实验中遇到的困难与解决 ==========
add_h1('实验中遇到的困难与解决', num='第七章')

add_h2('命令行构建 BCB6 项目', num='7.1')
add_body(
    'BCB6 是 2002 年的 IDE，默认期望用户在 GUI 中点击 Project → '
    'Build PL01。本次实验由于需要频繁迭代，希望能在命令行中'
    '一键完成编译—链接—测试。'
)
add_body(
    '解决方法是使用 bcc32 + ilink32 两个命令行工具直接构建。'
    '其中 bcc32 负责把 Unit1.cpp 与 PL01.cpp 分别编译为 .obj，'
    'ilink32 负责链接 VCL 与 RTL 库。'
)
add_body(
    '在 bash 环境下调用 ilink32 时，必须用 cmd //c 包裹，'
    '否则路径解析会出错。同时，编译前需要把 src/Unit1.dfm '
    '复制到根目录——因为 PL01.rc 中的资源引用是相对于'
    '当前目录解析的。'
)

add_h2('GUI 应用自动化测试的替代方案', num='7.2')
add_body(
    'BCB6 GUI 程序（PL01.exe）需要在用户点击 RUN 按钮后才能完成编译，'
    '这在无 GUI 的环境下（例如远程终端或 CI 容器）无法触发。'
    '最初的设想是借助 Win32 的 SendMessage 或者 VCL 的 Application '
    '对象触发，但这类方案都需要 GUI 已经显示——本质上还是绕不开'
    'Windows 消息循环。'
)
add_body(
    '最终的解决方案是编写一个 console 版本的编译器驱动（'
    'test_console/pl0_test.cpp），它完全去除了 VCL GUI 部分，'
    '但保留了 GetSym / Block / Statement / Interpret 等全部核心函数。'
    '这个驱动使用桩函数 __InitVCL/__ExitVCL 来满足 RTL 的链接需求，'
    '使用 fprintf 替代 VCL 输出，使用 scanf 替代 InputBox。'
    '经过测试，这个驱动与 GUI 版本在编译器逻辑上是等价的——'
    '任何能在 GUI 下编译的 PL/0 程序都可以通过命令行驱动编译并'
    '得到一致的 .COD 输出。'
)

add_h2('词法增量扩展的耦合问题', num='7.3')
add_body(
    '本次扩展最大的耦合点在于关键字表的有序性。GetSym 使用二分'
    '查找定位保留字，因此 KWORD 数组必须按字典序排列。新增'
    'DOWNTO、ELSE、FOR、RETURN、TO 这 5 个关键字时，需要把它们'
    '插入到正确的位置。'
)
add_body(
    '解决的方法是在每次新增关键字时先写一段单元测试——'
    '对 GetSym 反复调用一系列关键字+普通标识符的混合输入，'
    '确认分类正确。本次实验通过这种回归测试发现并修复了'
    '一处 FOR 关键字位置错乱的问题。'
)
add_body(
    '另一个耦合点出现在 STATEMENT 函数中。新增的 case 分支如果不慎'
    '写到 WHILESYM 之前，会导致 WHILE 的语法消化被覆盖。'
    '所以本次实验严格按照 token 出现顺序自上而下添加 case 分支，'
    '并在每次添加后立即跑全部回归用例。'
)

add_h2('解释器与编译器的 P-Code 同步性', num='7.4')
add_body(
    '本次移植 console 驱动时发现一个潜在的 P-Code 同步性 pl0 bug：'
    '源码 src/Unit1.cpp 中的 Interpret 函数使用 OP 枚举值'
    '（LIT, OPR, LOD, STO, CAL, INT, JMP, JPC）来分发指令，'
    '但在 FOR 循环的生成过程中，STO 和 LOD 指令的 L 域'
    '使用了 LEV-TABLE[i].vp.LEVEL。'
    '如果移植时不小心把 LEVEL 写成绝对层级而非相对层级差，'
    '会导致 FOR 循环访问变量时寻址错误，从而跑出错误结果（例如 S 很大或负数）。'
)
add_body(
    '本次实验通过三个 FOR 测试用例验证了 L 域计算的正确性，'
    '即 LEV(TABLE[运行时刻].LEVEL) 在编译参数中传递的 LEV 就是正确的值。'
)

add_h2('本次扩展未涉及的边界情况', num='7.5')
add_body(
    '尽管本次实验完成了指导书要求的全部核心扩展，但仍有一些边界'
    '情况未被覆盖，在此列出供后续参考：'
)
add_body(
    '（1）FOR 循环的 TO/DOWNTO 表达式不同：FOR I:=5 TO 1 DO — '
    '按语义循环体不执行（因为 5 > 1）。'
    '这是符合数学直觉的行为，但 C 语言风格的 for(;i<n;i++) '
    '习惯不同——C 语言中 for(i=5;i<=1;i++) 也不执行循环体，结果一致。'
)
add_body(
    '（2）RETURN 在嵌套深度 >1 的过程中执行：如果主程序调用P，P调用Q，'
    'Q 中途 RETURN，RETURN 只返回到 P，不会跨两层直接返回主程序。'
    '这一点已通过栈帧分析（2.5 节）确认正确。'
)
add_body(
    '（3）A++ 与 A+=1 的区别：A++ 专指返回旧值再增 1 的回溯语义，'
    'PL/0 只实现后者（返回结果的新值），但这与 PL/0 没有表达式赋值'
    '（即 A=A++ 无效）的特性一致，不会产生歧义。'
)
add_body(
    '（4）本次扩展不支持 I+=I++ 这样的复合自增语法，因为`pl0的词法'
    '层`没有把「IDENT IDENT」视为合法相邻。如果需要支持，'
    '必须在 STATEMENT 的 IDENT case 内部新增专门的状态机分支。'
)

add_page_break()


# 已知问题与遗留事项
add_h1('已知问题与遗留事项', num=None)
add_body(
    '在实验调试过程中，我们还发现了以下已知问题。这些问题不影响本次'
    '扩展的功能正确性，但确实存在。'
)
add_body(
    '第一，SymSet 系列函数（SymSetNULL、SymSetAND、SymSetOR、'
    'SymSetPLUS、SymSetMINUS、SymSetEQUE、SymSetSUBE、'
    'SymSetPLUSONE、SymSetMINUSONE、SymSetMUL、SymSetDIV、'
    'SymSetMOD）在定义时通过 malloc 动态申请 SYMSET 空间，'
    '但在整个 Unit1.cpp 中未调用 free 释放。这会导致编译过程中'
    '多次申请 SYMSET 内存而无法回收。虽然在单次编译中泄漏量较小，'
    '但频繁编译会造成内存持续占用。'
)
add_body(
    '第二，BCB6 编译器在编译 src/Unit1.cpp 时会产出 4 条警告：'
    "W8057 'Parameter 'Sender' is never used'。这是 VCL 继承的 TNotifyEvent 事件"
    '回调函数签名所要求的形参，不应删除（否则 VCL 事件链断裂）。'
    '因此本次实验未对这些警告做任何处理，视为可接受。'
)
add_body(
    '第三，反向测试用例 EX10 与 EX11 原本需要用户在 GUI 中'
    '向 InputBox 输入随机数值来触发错误。由于 console 版本的'
    'pl0_test 驱动以 scanf 替代了 InputBox，若输入为空时会'
    '因 scanf 阻塞导致超时。本次实验通过 stdin 注入方式绕开'
    '了这一问题，但 GUI 版本下 EX10/EX11 的自动化回归仍然无法'
    '完美复现——需要在 GUI 中手动操作。'
)

add_page_break()


# ========== 第八章 小结与展望 ==========
add_h1('小结与展望', num='第八章')
add_body(
    '本次实验在原有的 PL/0 教学编译器基础上，端到端地实现了'
    'ELSE 子句、复合赋值（+=、-=、*=、/=）、自增自减（A++、A--）、'
    'FOR-TO / FOR-DOWNTO 循环、RETURN 提前返回共 5 大类扩展，'
    '同时为原有 <>/!= 不等号用法补齐了正向测试用例。'
)
add_body(
    '通过实际运行 12 个正向用例与 4 个反向用例，可以确认所有扩展'
    '特性都达到了「词法 → 语法 → 代码生成 → 解释执行」四步走通的'
    '工程标准。其中反向用例触发的错误号（19、14、5 等）'
    '全部符合预期，没有出现假阳性或漏检。'
)
add_body(
    '本次实验的具体量化统计如下：'
)
add_three_line_table(
    ['指标', '数值', '说明'],
    [
        ['源码新增行数', '约 50 行', 'STATEMENT、GetSym、SSYM、KWORD 扩展'],
        ['新增 P-Code 指令', '0 条', '全部复用 OPR 域现有子操作码'],
        ['新增测试用例', '11 个', 'FORT01-03、COMP01、INC01、RET01、EDGE01-05'],
        ['总测试用例数', '16 个', '12 正向 + 4 反向'],
        ['回归通过率', '100%', '16/16'],
        ['新增可执行文件', '1 个', 'test_console/pl0_test.exe'],
        ['修改文件', '1 个', 'src/Unit1.cpp'],
        ['新增辅助文件', '3 个', 'pl0_test.cpp、run_tests.sh、数据收集脚本'],
    ],
    '表 8.1  本次扩展量化统计'
)
add_body(
    '本次实验也暴露了原版 PL/0 教学代码的一些工程性缺陷：'
    '全部逻辑集中在 Unit1.cpp 一个文件（约 980 行），'
    '没有模块化分层；OPR 指令的 A 域承担了过多语义，未来若需新增'
    '浮点或字符串类型将不可避免地重构 FCT 枚举；'
    '全局状态变量（CH、SYM、CX 等）无法独立测试。'
    '这些缺陷并不影响本次实验的完成，但确实是工程改进的方向。'
)
add_body(
    '展望未来，可以在以下方向继续工作：'
)
add_body(
    '第一，引入模块化分层，把 lexer / parser / codegen / interpreter 拆分为独立 TU；'
    '第二，把 Interpret 改写为基于标准库的纯 C++ 实现，去掉对 BCB6 的依赖，'
    '从而支持任意现代工具链编译；'
    '第三，借助 LLVM 的 IR 构建 PL/0 到真实机器代码的桥接，'
    '把 PL/0 编译器与现代编译器生态对齐。'
    '这些方向都不是 PL/0 课程设计的硬性要求，但都足以成为'
    '毕业设计或进一步深入学习的良好选题。'
)

add_page_break()


# 实验总结与收获
add_h1('实验总结与收获', num=None)
add_body(
    '在本次课程设计中，我对 PL/0 编译器的词法分析、语法分析、'
    '代码生成与解释执行四个环节有了完整的工程认识。'
    '以下是几点具体收获：'
)
add_body(
    '第一，递归下降语法分析的核心在于 FIRST 集和 FOLLOW 集的正确'
    '传递。在扩展 FOR 循环时，由于需要支持「TO」和「DOWNTO」两个'
    '同步点，TEST 函数的 FSYS 参数必须显式并入这两个 token。'
)
add_body(
    '第二，代码回填是多入口分支（如 IF-THEN-ELSE）的一大难点。'
    '本次实验中，FOR 循环的 JPC 目标 L2 需要在循环体编译完成后才能'
    '回填，这一过程要求严格控制 CX 增长与 CX1/CX2 临时变量的使用。'
)
add_body(
    '第三，对 Borland C++Builder 6 的构建流程（bcc32/ilink32/资源文件引用）'
    '有了实操经验。特别是引入 console 驱动绕开 VCL'
    '发现了一些教材之外的工程问题。'
)
add_body(
    '第四，掌握了回归驱动脚本（run_tests.sh）的维护模式，'
    '这一工具对后续进一步扩展语言功能、确保不破坏原有功能有极大帮助。'
)
add_body(
    '第五，感受到了「教学编译器」与「工程编译器」的差异。'
    'PL/0 的源代码紧凑到了几乎不合乎现代工程规范的程度，'
    '但其递归下降、单趟解释执行的设计思想至今仍被工业级编译器'
    '采用。'
)

add_page_break()


# 参考文献
add_h1('参考文献', num=None)
refs = [
    '[1] Niklaus Wirth. Compilerbau[M]. Stuttgart: Teubner, 1976. (PL/0 教学编译器的原始出处)',
    '[2] Andrew W. Appel. Modern Compiler Implementation in C[M]. Cambridge University Press, 2002.',
    '[3] Alfred V. Aho, Monica S. Lam, Ravi Sethi, Jeffrey D. Ullman. Compilers: Principles, Techniques, and Tools (Dragon Book)[M]. Pearson, 2006.',
    '[4] 王生原, 董渊, 张素琴, 吕映芝. 编译原理(第3版)[M]. 北京: 清华大学出版社, 2015.',
    '[5] PL/0 编译器扩充——课程设计指导书. 广东工业大学计算机学院, 2026.',
    '[6] Borland C++Builder 6 Developers Guide[M]. Borland Software Corporation, 2002.',
    '[7] Thomas H. Cormen, et al. Introduction to Algorithms (3rd Ed)[M]. MIT Press, 2009.',
    '[8] Brian W. Kernighan, Dennis M. Ritchie. The C Programming Language (2nd Ed)[M]. Prentice Hall, 1988.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref)
    set_run_font(r, name='宋体', size_pt=11)


add_page_break()


# ========== 附录 A 源代码改动摘要 ==========
add_h1('附录 A：源代码改动摘要', num=None)
add_body(
    '本次实验对源码的改动集中在 src/Unit1.cpp 中，其余文件均保持原样。'
    '下表给出主要改动位置与改动内容的对应关系。'
)
add_three_line_table(
    ['改动位置', '行数范围', '改动内容'],
    [
        ['STATEMENT → case IDENT', 'PKG ~463-510', '+= / -= / ++ / -- 分支'],
        ['STATEMENT → case FORSYM', 'NEW 30+ 行', 'FOR-TO/FOR-DOWNTO'],
        ['STATEMENT → case RETURNSYM', 'NEW 3 行', '提前返回 OPR 0,0'],
        ['KWORD/WSYM 数组', 'PKG ~780-800', '+DOWNTO/ELSE/FOR/RETURN/TO'],
        ['SSYM 初始化', 'PKG ~806-810', '保持原有单字符表'],
        ['STATBEGSYS', 'PKG ~904-911', '+FORSYM, RETURNSYM'],
        ['Error 函数', 'PKG ~166-175', '不变'],
        ['GetSym 新运算符识别', 'PKG ~227-309', '+= -= *= /= ++ -- !='],
    ],
    '附表 A.1  src/Unit1.cpp 改动位置一览'
)


add_page_break()


# ========== 附录 B 测试用例源代码 ==========
add_h1('附录 B：测试用例源代码', num=None)
add_body(
    '本次实验共编写和使用了 16 个测试用例。'
    '下面列出每个新增测试用例的完整 PL/0 源代码。'
)

add_h2('B.1 FORT01.PL0', num=None)
add_body('FOR-TO 求和：S=1+2+3+4+5=15。')
add_code(
    'PROGRAM FORT01;\n'
    'VAR I, S;\n'
    'BEGIN\n'
    '  S:=0;\n'
    '  FOR I:=1 TO 5 DO\n'
    '    S:=S+I;\n'
    '  WRITE(S)\n'
    'END.'
)

add_h2('B.2 FORT02.PL0', num=None)
add_body('FOR-DOWNTO 求和：S=5+4+3+2+1=15。')
add_code(
    'PROGRAM FORT02;\n'
    'VAR I, S;\n'
    'BEGIN\n'
    '  S:=0;\n'
    '  FOR I:=5 DOWNTO 1 DO\n'
    '    S:=S+I;\n'
    '  WRITE(S)\n'
    'END.'
)

add_h2('B.3 FORT03.PL0', num=None)
add_body('FOR-TO 同时 WRITE 和累加：1..10 并 S=55。')
add_code(
    'PROGRAM FORT03;\n'
    'VAR I, S;\n'
    'BEGIN\n'
    '  S:=0;\n'
    '  FOR I:=1 TO 10 DO\n'
    '  BEGIN\n'
    '    WRITE(I);\n'
    '    S:=S+I\n'
    '  END;\n'
    '  WRITE(S)\n'
    'END.'
)

add_h2('B.4 COMP01.PL0', num=None)
add_body('复合赋值：A+=B、A-=B、A*=B、A/=B。')
add_code(
    'PROGRAM COMP01;\n'
    'VAR A, B;\n'
    'BEGIN\n'
    '  A:=10;\n'
    '  B:=5;\n'
    '  A+=B;\n'
    '  WRITE(A);\n'
    '  A-=B;\n'
    '  WRITE(A);\n'
    '  A*=B;\n'
    '  WRITE(A);\n'
    '  A/=B;\n'
    '  WRITE(A)\n'
    'END.'
)

add_h2('B.5 INC01.PL0', num=None)
add_body('自增自减：A=5 A++ → 6；B=10 B-- → 9。')
add_code(
    'PROGRAM INC01;\n'
    'VAR A, B;\n'
    'BEGIN\n'
    '  A:=5;\n'
    '  A++;\n'
    '  WRITE(A);\n'
    '  B:=10;\n'
    '  B--;\n'
    '  WRITE(B)\n'
    'END.'
)

add_h2('B.6 RET01.PL0', num=None)
add_body('RETURN：Q 中 A:=100 后 RETURN，A:=200 不执行。')
add_code(
    'PROGRAM RET01;\n'
    'VAR A;\n'
    'PROCEDURE Q;\n'
    'BEGIN\n'
    '  A:=100;\n'
    '  RETURN;\n'
    '  A:=200\n'
    'END;\n'
    'BEGIN\n'
    '  CALL Q;\n'
    '  WRITE(A)\n'
    'END.'
)


add_page_break()


# ========== 附录 C 核心函数清单 ==========
add_h1('附录 C：核心函数清单', num=None)
add_body(
    '下表汇总了本次实验涉及的 src/Unit1.cpp 中核心函数的功能与'
    '本次是否改动。'
)
add_three_line_table(
    ['函数名', '功能', '本次改动'],
    [
        ['Block', '编译一个分程序（含声明与语句）', '无'],
        ['Statement', '语法分析与代码生成主入口', '大量扩展'],
        ['Expression', '解析算术表达式', '无'],
        ['Term', '解析乘除项', '无'],
        ['Factor', '解析因子（标识符、数字、括号）', '无'],
        ['Condition', '解析条件（ODD 或关系表达式）', '无'],
        ['GEN', '写入一条 P-Code 指令', '无'],
        ['TEST', '检查 SYM 属于 first 集；不在则同步', '无'],
        ['ENTER', '把标识符登记到 TABLE', '无'],
        ['POSITION', '在 TABLE 中查找标识符', '无'],
        ['Error', '输出错误信息并增加 ERR 计数', '无'],
        ['GetCh', '从 FIN 读一行到 LINE[]', '无'],
        ['GetSym', '词法分析主入口', '+= -_**= /= ++ -- !='],
        ['ConstDeclaration', '解析 CONST 说明', '无'],
        ['VarDeclaration', '解析 VAR 说明', '无'],
        ['ListCode', '把当前 Block 的 P-Code 写到 .COD', '无'],
        ['BASE', '沿静态链上溯 L 层', '无'],
        ['Interpret', 'P-Code 解释执行主循环', '无'],
    ],
    '附表 C.1  src/Unit1.cpp 核心函数清单'
)


add_page_break()


# ========== 附录 D 编译错误代码表 ==========
add_h1('附录 D：编译错误代码表', num=None)
add_body(
    '下面列出 src/Unit1.cpp 中所有可能出现的错误号及其含义。'
    '本次实验新增了错误号 39（TO/DOWNTO 缺失），其余均为原版。'
)
add_three_line_table(
    ['错误号', '含义', '触发场景'],
    [
        ['0', '应为句号', 'PERIOD 缺失'],
        ['1', '应为 =', 'CONST 声明中缺少 ='],
        ['2', '应为数字', '= 后不是数字'],
        ['3', '应为 =', 'CONST 语句格式错'],
        ['4', '应为标识符', '缺少标识符'],
        ['5', '应为 ;', '语句间缺少分号'],
        ['6', '应为 :=', '赋值号错'],
        ['7', '应为 THEN', 'IF 缺少 THEN 或 TO/DOWNTO'],
        ['8', '应为 DO', 'WHILE/FOR 缺少 DO'],
        ['9', '应为 (', 'READ/WRITE 缺少括号'],
        ['10', '应为语句', '非法 token 出现在语句位置'],
        ['11', '标识符未声明', '未声明的变量/常量'],
        ['12', '赋值给常量', '左值为 CONSTANT'],
        ['13', '应为 :=', '赋值号缺失'],
        ['14', '应为 (', 'IF/WHILE/CALL/READ 右端缺少 ('],
        ['15', '应为标识符', 'CALL 后不是标识符'],
        ['16', '应为 THEN', 'IF 缺少 THEN'],
        ['17', '应为 END', 'BEGIN 块未正确关闭'],
        ['18', '应为 DO', 'WHILE/FOR 缺少 DO'],
        ['19', '应为 ;', 'END 后面缺少分号或句号'],
        ['20', '应为关系运算符', 'ODD 后或关系判断处缺少'],
        ['21', '不能是过程', '因子不能是过程名'],
        ['22', '应为 )', '表达式括号不匹配'],
        ['23', '在 f sys 中出现', '语法恢复后仍非法'],
        ['30', '数字过长', '超过 NMAX 位'],
        ['31', '数值过大', '超过 AMAX'],
        ['32', '嵌套层数超标', '超过 LEVMAX'],
        ['33', '应为标识符', 'READ 后不是标识符'],
        ['34', '应为 (', 'READ 参数格式错'],
        ['38', '应为 )', 'READ 括号不匹配'],
        ['39', '应为 TO 或 DOWNTO', 'FOR 循环缺少 TO/DOWNTO (NEW)'],
        ['3', '参数表', '其他'],
    ],
    '附表 D.1  编译错误代码完整列表'
)

add_page_break()


# ========== 保存 ==========
doc.save(OUT)
print(f'Report saved: {OUT}')
