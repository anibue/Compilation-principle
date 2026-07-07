# -*- coding: utf-8 -*-
"""Generate PL/0 report v6 — 40 pages. Full source listings + trace tables."""
import os
from docx import Document
from docx.shared import Pt, Mm, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r'C:\Users\xwmsu\Desktop\c_builder_b'
FIG = os.path.join(ROOT, 'docs', 'report_figures')
OUT = os.path.join(ROOT, 'PL0编译器扩充实验报告.docx')

doc = Document()


def set_font(run, name='宋体', sz=12, bold=False):
    run.font.name = 'Times New Roman'; run.font.size = Pt(sz); run.font.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFs = rPr.find(qn('w:rFonts'))
    if rFs is None: rFs = OxmlElement('w:rFonts'); rPr.append(rFs)
    rFs.set(qn('w:eastAsia'), name); rFs.set(qn('w:ascii'), 'Times New Roman'); rFs.set(qn('w:hAnsi'), 'Times New Roman')


def P(text='', font='宋体', sz=12, bold=False, align=None, indent=True, sp=1.5):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = sp
    p.paragraph_format.space_before = Pt(0); p.paragraph_format.space_after = Pt(3)
    if align is not None: p.alignment = align
    if indent: p.paragraph_format.first_line_indent = Pt(sz * 2)
    if text:
        r = p.add_run(text); set_font(r, name=font, sz=sz, bold=bold)
    return p


def body(t, sz=12): return P(t, font='宋体', sz=sz, indent=True, sp=1.5)


def h1(t, num=None):
    f = (num + '  ' + t) if num else t
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f); set_font(r, name='黑体', sz=16, bold=True)


def h2(t, num=None):
    f = (num + '  ' + t) if num else t
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f); set_font(r, name='黑体', sz=14, bold=True)


def code(txt, sz=9):
    for ln in txt.split('\n'):
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = Pt(0); p.paragraph_format.left_indent = Pt(12)
        r = p.add_run(ln if ln else ' '); set_font(r, name='Consolas', sz=sz, bold=False)


def IMG(path, cap, wc=12.5):
    if not os.path.exists(path): return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.first_line_indent = Pt(0)
    r = p.add_run(); r.add_picture(path, width=Cm(wc))
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(0); cp.paragraph_format.space_after = Pt(10)
    cp.paragraph_format.first_line_indent = Pt(0)
    cr = cp.add_run(cap); set_font(cr, name='宋体', sz=9)


def BORD(t):
    tbl = t._tbl; pr = tbl.tblPr
    b = OxmlElement('w:tblBorders')
    for e in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement('w:' + e); el.set(qn('w:val'),'nil'); b.append(el)
    pr.append(b)


def TBL(hdrs, rows, cap):
    n = len(hdrs); t = doc.add_table(rows=1+len(rows), cols=n)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.autofit = True; BORD(t)
    for i,h in enumerate(hdrs):
        c = t.rows[0].cells[i]; c.text=''
        p = c.paragraphs[0]; p.paragraph_format.first_line_indent = Pt(0)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); set_font(r, name='黑体', sz=9, bold=True)
    for ri,row in enumerate(rows):
        for ci,val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text=''
            p = c.paragraphs[0]; p.paragraph_format.first_line_indent = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val)); set_font(r, name='宋体', sz=9)
    cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_before = Pt(4); cp.paragraph_format.space_after = Pt(10)
    cp.paragraph_format.first_line_indent = Pt(0)
    cr = cp.add_run(cap); set_font(cr, name='黑体', sz=9, bold=True)


def PB():
    p = doc.add_paragraph(); p.add_run().add_break(WD_BREAK.PAGE)


def FOOTER():
    sec = doc.sections[0]; ft = sec.footer
    p = ft.paragraphs[0] if ft.paragraphs else ft.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run()
    f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'),'begin')
    it = OxmlElement('w:instrText'); it.set(qn('xml:space'),'preserve'); it.text='PAGE   \\* MERGEFORMAT'
    f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'),'end')
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    set_font(r, name='Times New Roman', sz=9)


st = doc.styles['Normal']; st.font.name='宋体'; st.font.size=Pt(12)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'宋体')
st.element.rPr.rFonts.set(qn('w:ascii'),'Times New Roman')
st.element.rPr.rFonts.set(qn('w:hAnsi'),'Times New Roman')

for sec in doc.sections:
    sec.top_margin = Mm(30); sec.bottom_margin = Mm(25)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2)

# === COVER ===
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(200)
r = p.add_run('PL/0 编译器扩充实验报告'); set_font(r, name='黑体', sz=22, bold=True)
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before = Pt(40)
r = p.add_run('广东工业大学 计算机学院'); set_font(r, name='宋体', sz=16)
PB()

# === TOC ===
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(20)
r = p.add_run('目  录'); set_font(r, name='黑体', sz=18, bold=True)
toc = [
    ('第一章  绪论','1'),('    1.1  PL/0 语言与教学编译器简介','1'),('    1.2  本次实验的目标与动机','2'),
    ('    1.3  开发环境与项目结构','2'),('    1.4  课程设计指导书原文要求','3'),('第二章  编译流程与基础理论','3'),
    ('    2.1  PL/0 编译器的整体架构','3'),('    2.2  PL/0 语言文法完整描述','5'),('    2.3  递归下降语法分析','6'),
    ('    2.4  符号表与 P-Code 中间表示','7'),('    2.5  解释执行模型与栈帧管理','9'),('第三章  词法分析器扩展','11'),
    ('    3.1  扩展保留字与运算符','11'),('    3.2  GetSym 的状态机实现','12'),('    3.3  关键字表的二分查找','13'),
    ('第四章  语法分析与代码生成扩展','15'),('    4.1  STATEMENT 函数的结构','15'),('    4.2  ELSE 子句的代码回填','16'),
    ('    4.3  复合赋值 +=/-= 的代码生成','17'),('    4.4  ++/-- 自增自减语句','18'),('    4.5  FOR-TO / FOR-DOWNTO 循环','19'),
    ('    4.6  RETURN 提前返回','21'),('第五章  核心函数代码深度分析','23'),('    5.1  三件套总览','23'),
    ('    5.2  Block 函数深度分析','24'),('    5.3  错误恢复机制','26'),('    5.4  OPR 指令域分析','27'),
    ('    5.5  静态链与栈帧管理','28'),('第六章  测试与实验结果','30'),('    6.1  测试用例组织','30'),
    ('    6.2  正向测试与运行结果','31'),('    6.3  反向测试与错误处理','33'),('    6.4  控制台回归驱动','34'),
    ('    6.5  回归结果统计','35'),('第七章  小结与展望','36'),('第八章  实验总结与收获','38'),
    ('参考文献','40'),('附录 A  新增 token 完整列表','41'),('附录 B  测试用例源代码','42'),
    ('附录 C  核心函数清单','44'),('附录 D  编译错误代码表','45'),
]
for lbl,pg in toc:
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(3); p.paragraph_format.line_spacing = 1.15
    r = p.add_run(lbl)
    set_font(r, name='黑体' if not lbl.startswith('    ') else '宋体', sz=(11 if not lbl.startswith('    ') else 10), bold=not lbl.startswith('    '))
    tr = p.add_run('\t' + pg); set_font(tr, name='Times New Roman', sz=10)
PB()
FOOTER()


# ================================================================
# 第一章 绪论
# ================================================================
h1('绪论', num='第一章')

h2('PL/0 语言与教学编译器简介', num='1.1')
body('PL/0 是由 Niklaus Wirth 在 1976 年编写的小型教学语言。Wirth 是 Pascal 语言的设计者，也是结构化编程和递归下降语法分析的先驱。PL/0 作为其子集，完整覆盖了编译原理课程的全部核心环节。')
body('从语言设计看，PL/0 是一种典型的「Pascal 子集」。它没有类型系统（只有整型）、没有数组、没有记录、没有函数返回值（通过全局变量传递结果）、没有递归调用。这些限制使得编译器可被压缩到千行级。')
body('原版 PL/0 语法精简：只有常量/变量/过程声明、赋值、条件语句（IF-THEN）、循环（WHILE-DO）、过程调用（CALL）、READ、WRITE 共 9 类语句。严谨的框架给课设带来独特工程约束。')
body('从编译器理论看，PL/0 完整对应龙书第 2 至 7 章。学生通过亲手实现扩展，可验证文法分析、中间代码生成、运行时环境等每个概念。这正是 PL/0 长盛不衰的原因。')
body('本实验的原始代码来自 Wirth 教科书 Compilerbau 中 PL/0 编译器向 BCB6 移植的版本。BCB6 的 VCL 图形界面把编译结果与 P-Code 列表直接显示在窗体上，降低了调试门槛。')
body('此外，PL/0 的 P-Code 设计与 JVM 字节码、CPython 字节码高度相似：都是基于栈的虚拟机指令集，都用操作栈完成运算，都用变量表管理局部状态。理解了 P/0 的 P-Code，再 learning 现代虚拟机技术时会有一种「原来如此」的感觉。')

h2('本次实验的目标与动机', num='1.2')
body('本次课程设计题目与要求来自广东工业大学计算机学院编译原理课程设计指导书，核心任务是在已有 PL/0 编译器基础上扩展 7 类新语法。')
body('第 1 项：增加单词（ELSE、FOR、TO、DOWNTO；+=、-=、*=、/=、<>、!=）。难点在单字符前瞻框架内正确识别双字符运算符，新增保留字需二分查找表保持字典序。')
body("第 2 项：# 改为 <> 或 !=。看似简单，但 # 原为单字符 SSYM ='#'，改为 != 双字符后需词法→语法协同修改——原版移植已通过 != 与 <> 双映射实现。")
body('第 3 项：ELSE 子句。语法分析层最大扩展，需处理双回填（JPC 与 JMP）的 backpatching 逻辑。')
body('第 4-7 项：复合赋值 +=/-=、自增自减 A++/A--、FOR-TO/FOR-DOWNTO、RETURN。每项都需词法→语法→代码生成→解释执行四步协同。')
body('指导书隐含要求：新语法必须在解释器中正确执行——不能只「编译通过」却「解释出错」。这在后续测试验证环节会反复强调。')
body('额外动机：让 BCB6 项目能通过命令行自动构建。GUI 环境手动点击无法满足「反复迭代、快速回归」节奏。最终实现的一键构建脚本可作为持续集成基础。')

h2('开发环境与项目结构', num='1.3')
body('开发环境：Borland C++Builder 6（BCB6），2002 年发布的 C++ IDE，自带 bcc32 编译器与 ilink32 链接器，基于 VCL（Visual Component Library）。')
body('bcc32 支持 C++98 绝大多数特性，ilink32 负责链接 RTL 与 VCL 库。命令行构建需显式指定库路径与链接顺序。')
body('项目结构简明扼要：根目录 PL01.bpr（工程文件），src/ 源代码目录，project/ 编译产物。test_cases/ 存放所有 .PL0 源文件与 .COD 参考输出。test_console/ 为本次新增，存放 console 编译器驱动与回归脚本。')

TBL(['目录','文件','说明','本次改动'],[
    ['src/','Unit1.cpp','编译器源码（约 980 行）','是（核心扩展）'],
    ['src/','Unit1.h','VCL 窗体类声明','否'],
    ['src/','Unit1.dfm','窗体资源','否'],
    ['src/','PL01.cpp','WinMain 入口','否'],
    ['project/','PL01.exe','GUI 可执行文件','是（重新编译）'],
    ['test_cases/','*.PL0 / *.COD','测试用例+参考输出','是（新增 11 个用例）'],
    ['test_console/','pl0_test.cpp','命令行编译器驱动','是（新增）'],
    ['test_console/','run_tests.sh','回归测试脚本','是（新增）'],
    ['test_console/','pl0_test.exe','命令行可执行文件','是（新增）'],
],'表 1.1 项目文件结构一览')

body('原始 GUI 程序通过 VCL InputBox 接收用户输入触发 READ，RED 按钮点击时完成编译并显示 P-Code。此 GUI 依赖使得命令行下无法自动化测试——这是本次实验主要工程挑战之一。')
body('代码组织方面，全部核心逻辑汇集在 Unit1.cpp 单文件（约 980 行）。这种「单文件工程」便于学生整体理解编译流程，但对工程实践不利于分工协作。本次实验通过 console 驱动分层设计缓解此问题。')

h2('课程设计指导书原文要求', num='1.4')
body('以下逐条列出指导书要求，附完成情况与验证方式。')
TBL(['序号','指导书要求','本次实现','验证方式'],[
    ['1','增加单词 ELSE/FOR/TO/DOWNTO/+=/-=/…','已实现（16 个新 token）','GetSym 测试'],
    ['2','# → <> 或 !=','已实现（!= 与 <> 同义映射到 NEQ）','E01/EX01'],
    ['3','增加 ELSE 子句','已实现（双回填）','E110'],
    ['4','扩充 += / -=','已实现（OPR 0,2/0,3）','COMP01'],
    ['5','扩充 FOR-TO / FOR-DOWNTO','已实现（OPR 0,13/0,11）','FORT01-03'],
    ['6','增加 A++ / A--','已实现（4 条固定指令序列）','INC01'],
    ['7','增加 RETURN','已实现（OPR 0,0）','RET01'],
    ['8','写出文法与语法描述','已实现（2.2 节）','EBNF 规则表'],
    ['9','实验截图与数据','已实现（第 6 章）','6 张图 + 16 用例'],
],'表 1.2 完成情况对照表')
body('表 1.2 显示本次实验完成全部 9 项核心要求。')

PB()

# ================================================================
# 第二章 编译流程与基础理论
# ================================================================
h1('编译流程与基础理论', num='第二章')

h2('PL/0 编译器的整体架构', num='2.1')
body('PL/0 是经典的「单趟 + 解释执行」结构。源代码一次扫描即可完成词法分析、语法分析、代码生成，所生成 P-Code 不写入中间文件直接进入解释器循环。这种设计既是教学优势（千行级掌握全流程），也是工程局限（无法支持多趟优化）。')
body('编译器四层：词法层（GetSym）、语法层（STATEMENT）、中间表示层（CODE 数组）、解释执行层（Interpret）。每层消费上层输出、向下一层提供输入。单向数据流使调试时只需顺 CX 增长方向观察 CODE 数组。')
body('与工业编译器对比：GCC 包括数千优化 Pass，LLVM 是可扩展 SSA IR，而 PL/0 只扫描一遍、不做优化。教学角度，「极简」是优势——学生 48 小时可走通全链路而非数月。')
body('本次扩展涉及对全部四层的修改：词法层 16 个新 token，语法层 6 个 case 分支，中间表示层通过复用 OPR 子操作码完成，解释执行层完全不动。这种「顶层改动、底层不动」结构是 PL/0 可扩展性的体现。')

IMG(os.path.join(FIG,'pipeline.png'),'图 2.1 PL/0 编译与解释的整体架构')
body('从图 2.1 可见，源程序从左到右依次通过词法→语法→代码生成，生成 P-Code 被 CODE 数组持有。解释器 Interpret 反向消费 CODE 数组（P 从 0 递增到 CX）。数据流严格单向，无反向依赖。')

h2('PL/0 语言文法完整描述', num='2.2')
body('扩展后的 PL/0 文法用 EBNF 描述，其中扩展部分用「★」标记。')
code("""<程序> ::= PROGRAM <标识符> ; <分程序>

<分程序> ::= [<常量声明>] [<变量声明>] {<过程声明>}
             BEGIN <语句> { ; <语句> } END

<常量声明> ::= CONST <标识符>=<数字> { , <标识符>=<数字> } ;

<变量声明> ::= VAR <标识符> { , <标识符> } ;

<过程声明> ::= PROCEDURE <标识符> ; <分程序> ;

<语句> ::= [<标识符> ( := | += | -= | *= | /= ) <表达式>   ★
           | <标识符>++                                        ★
           | <标识符>--                                        ★
           | IF <条件> THEN <语句> [ ELSE <语句> ]            ★
           | WHILE <条件> DO <语句>
           | FOR <标识符> := <表达式> ( TO | DOWNTO )        ★
                     <表达式> DO <语句>
           | RETURN                                            ★
           | READ ( <标识符> { , <标识符> } )
           | WRITE ( <表达式> { , <表达式> } )
           | CALL <标识符>
           | BEGIN <语句> { ; <语句> } END ]

<条件> ::= ODD <表达式>
           | <表达式> ( = | # | <> | != | < | <= | > | >= ) <表达式>

<表达式> ::= [ + | - ] <项> { ( + | - ) <项> }

<项> ::= <因子> { ( * | / | % ) <因子> }

<因子> ::= <标识符> | <数字> | ( <表达式> )""")

body('本次扩展集中在 <语句> production 上，新增 6 条赋值变体、2 条循环变体、1 条提前返回。这种「非终结符单点位扩展」是 PL/0 可扩展性的直接体现——新增 <因子> 级别语法会困难得多。')
body('下面给出各非终结符 FIRST/FOLLOW 集推导结果：')
body('FIRST(<程序>) = { PROGRAM }，FOLLOW(<程序>) = { $ }。这是递归下降分析的起点。')
body('FIRST(<分程序>) = { CONST, VAR, PROCEDURE, BEGIN }。当 CONST/VAR/PROCEDURE 都不出现时直接以 BEGIN 开始。')
body('FIRST(<语句>) = { IDENT, IF, WHILE, FOR, RETURN, READ, WRITE, CALL, BEGIN }，共 9 种语句形式。IDENT 在 STATEMENT 中由 case IDENT 统一处理，内部再区分。')
body('FIRST(<表达式>) = { +, -, IDENT, NUMBER, LPAREN }。作为 TEST 的 S1 参数，确保表达式解析不在非法 token 处卡住。')
body('FIRST(<项>) = FIRST(<因子>) = { IDENT, NUMBER, LPAREN }。')
body('FOLLOW(<语句>) = { ;, END, ELSE }。语句后可跟分号、END 或 ELSE。作为 FSYS 参数使用。')
body('FOLLOW(<表达式>) 需综合所有上下文：赋值 := 或复合赋值后、条件 THEN/DO 后、WRITE 参数中 ) 或 , 后。完整推导见 4.3 节。')

h2('递归下降语法分析', num='2.3')
body('递归下降语法分析是 PL/0 编译器最核心的分析方法。基本思想：为文法中的每个非终结符编写一个函数，函数体根据当前 SYM 选择产生式并递归调用其他函数。文法→代码的直接映射使代码结构清晰、易于理解。')
body('PL/0 中，每个非终结符对应一个函数：<程序> → Block(0,0,FSYS)；<分程序> → Block；<语句> → STATEMENT；<表达式> → EXPRESSION；<项> → TERM；<因子> → FACTOR；<条件> → CONDITION。一一对应，可直接从函数名推断对应的文法非终结符。')
body('关键技术是「先检查、再消化」的两段式写法：函数入口 TEST(STATBEGSYS, FSYS, N) 检查 SYM 是否在 STATBEGSYS 中；不在则报错 N 并同步到 STATBEGSYS ∪ FSYS。这种写法天然完成错误恢复——输入有语法错误，分析器也能继续扫描并尽可能多报告错误。')
body('扩展后的 STATEMENT 新增 6 个 case 分支：FORSYM（FOR 循环）、RETURNSYM（RETURN）、IDENT 内部的 PLUSEQ/MINUSEQ/TIMESBECOMES/SLASHBECOMES（复合赋值）和 PLUSPLUS/MINUSMINUS（自增自减）。新增分支与原有 BECOMES（:=）共享 IDENT case 块。')

IMG(os.path.join(FIG,'statement.png'),'图 2.2 STATEMENT 函数的整体结构')

body('下面给出 STATEMENT 函数整体结构伪代码，新增部分用「★」标记：')
code("""void STATEMENT(FSYS):
  TEST(STATBEGSYS, FSYS, 10)
  if SYM == IDENT:
    i = POSITION(id); GetSym
    if   SYM == BECOMES:     ; :=
    elif SYM == PLUSEQ:      ; +★
    elif SYM == MINUSEQ:     ; -★
    elif SYM == TIMESBECOMES:*★
    elif SYM == SLASHBECOMES:/★
    elif SYM == PLUSPLUS:    ; ++★
    elif SYM == MINUSMINUS:  ; --★
  elif SYM == IF:
    GetSym; CONDITION; JPC CX1
    TEST(then); GetSym; STATEMENT; CX1 = CX
    if SYM == ELSE: GetSym; STATEMENT
  elif SYM == FOR:           ★
    GetSym; id := EXPRESSION
    if SYM == TO: op=13
    elif SYM == DOWNTO: op=11
    GetSym; EXPRESSION; CX1=CX; JPC CX2
    TEST(do); GetSym; STATEMENT
    LOD id; LIT 1; OPR op2; STO id; JMP CX1; CX2 = CX
  elif SYM == RETURN:        ★
    GetSym; GEN(OPR,0,0)
  ......""")

body('新增的 FOR 循环与 RETURN 在结构上与原有语句统一：先消化入口 token，再解析内部表达式/条件，最后生成 P-Code。这种「统一结构」是递归下降易于扩展的根本原因。')

h2('符号表与 P-Code 中间表示', num='2.4')
body('PL/0 使用扁平数组 TABLE[TXMAX] 表示整个程序符号表，TX=0 永远作为哨兵：POSITION 查不到标识符时返回 0，0 自然成为「未声明」或「越界」的统一信号。')
body('TABLE 每项保存三信息：NAME（至多 10 字符标识符）、KIND（CONSTANT/VARIABLE/PROCEDUR）、LEVEL 与 ADDR（对应嵌套深度与运行栈相对位置）。')
body('POSITION 采用自后向前线性搜索：从当前 TX 向位置 0 方向查找。与 PL/0「最近嵌套作用域」原则一致——内层标识符遮蔽外层同名标识符。')
body('P-Code 是栈式虚拟机指令，每条指令由三元组 (F, L, A) 构成：F 为操作码（FCT 枚举），L 为嵌套层级差，A 为操作数（含义随 F 变化）。')

body('下面给出 FORT01 程序编译完成后 TABLE 的实际内容。FORT01 源程序：')
code("""PROGRAM FORT01;
VAR I, S;
BEGIN
  S:=0;
  FOR I:=1 TO 5 DO S:=S+I;
  WRITE(S)
END.""")

body('编译后 TABLE 内容如下：')
TBL(['TX','NAME','KIND','LEVEL','ADR/VAL'],[
    ['0','(sentinel)','—','—','—'],
    ['1','FORT01','PROCEDUR','0','0'],
    ['2','I','VARIABLE','0','3'],
    ['3','S','VARIABLE','0','4'],
],'表 2.1 FORT01 编译后 TABLE 的实际内容')
body('主程序 FORT01 作为过程占据 TABLE[1]，LEVEL=0，ADR=0（主程序不通过 CAL 调用，ADR 无意义）。变量 I 和 S 依次占据 TABLE[2] 和 TABLE[3]，LEVEL=0，ADR=3 和 ADR=4（从栈位置 3 开始分配局部变量）。')

body('P-Code 指令集遵循「操作码+操作数」模式。FCT 枚举 8 种操作码：LIT（取常量）、OPR（算术/关系/返回）、LOD（取变量）、STO（存变量）、CAL（调用过程）、INT（栈分配）、JMP（无条件跳转）、JPC（条件跳转）。其中 OPR 的 A 域承载 17 种子操作，是本次扩展的主要复用点。')

h2('解释执行模型与栈帧管理', num='2.5')
body('解释器 Interpret 维护三个工作寄存器：P（下一条指令地址）、B（当前过程基址）、T（栈顶指针）。运行时栈 S[] 被划分为若干活动记录，每个过程调用占用一个活动记录。')

TBL(['栈位置','内容','说明'],[
    ['S[B-3]','SL（静态链）','调用者的过程层级基址'],
    ['S[B-2]','DL（动态链）','调用者的 B 值'],
    ['S[B-1]','RA（返回地址）','CAL 指令的下一条 P'],
    ['S[B+0..n-1]','参数/局部变量','DX 个变量依次存放'],
],'表 2.2 运行时栈活动记录布局')

body('BASE 函数沿静态链上溯 L 层解析非局部变量：每次上溯一层，读取 S[B1] 作为新 B1。例如主程序（LEVEL=0）调用 P（LEVEL=1）时，P 需访问主程序变量 X，LOD 的 L=1，BASE 执行一次上溯得到主程序 B 地址，最终变量位于 S[B+X.adr]。')
body('本次扩展未涉及栈帧管理的修改：FOR、复合赋值、RETURN 等新语句只在当前活动记录内读写局部变量。RETURN 等价于 OPR 0,0 指令，恢复栈顶、返回地址与基址，与现代 RISC-V/x86 的 ret 行为一致。')

PB()

# ================================================================
# 第三章 词法分析器扩展
# ================================================================
h1('词法分析器扩展', num='第三章')

h2('扩展保留字与运算符', num='3.1')
body('本次实验新增与调整了 16 个保留字/运算符。保留字方面在原有 14 个基础上新增 ELSE、FOR、TO、DOWNTO、RETURN 共 5 个。')
body('运算符方面新增 +=、-=、*=、/=、!= 五个双字符运算符，同时保留原版 # 作为 != 同义写法。+ 后跟 + 识别为 PLUSPLUS，- 后跟 - 识别为 MINUSMINUS。')

body('下面是完整的 token 枚举定义（「★」标记为新增）：')
code("""enum SYMBOLS {
  NUL, IDENT, NUMBER, PLUS, MINUS, TIMES, SLASH,
  ODDSYM, EQL, NEQ, LSS, LEQ, GTR, GEQ,
  LPAREN, RPAREN, COMMA, SEMICOLON, PERIOD,
  BECOMES, ENDSYM, ELSESYM, FORSYM,        ★
  TOSYM, DOWNTOSYM, RETURNSYM,             ★
  PLUSEQ, MINUSEQ, TIMESBECOMES, SLASHBECOMES, ★
  PLUSPLUS, MINUSMINUS,                    ★
  BEGINSYM, IFSYM, THENSYM, WHILESYM, DOSYM,
  CALLSYM, CONSTSYM, VARSYM, PROCSYM, READSYM, WRITESYM
};""")
body('新增 token 被插入枚举中间（BECOMES 之后、BEGINSYM 之前），以保持与原版 token 编号兼容。WSYM 数组中内部码与枚举值一一对应。')

TBL(['类别','扩展项','token 枚举值'],[
    ['保留字','ELSE','ELSESYM'],['保留字','FOR','FORSYM'],
    ['保留字','TO','TOSYM'],['保留字','DOWNTO','DOWNTOSYM'],
    ['保留字','RETURN','RETURNSYM'],['运算符','A += B','PLUSEQ'],
    ['运算符','A -= B','MINUSEQ'],['运算符','A *= B','TIMESBECOMES'],
    ['运算符','A /= B','SLASHBECOMES'],['运算符','A++','PLUSPLUS'],
    ['运算符','A--','MINUSMINUS'],['运算符','A != B','NEQ（与 <> 同义）'],
],'表 3.1 本次扩展的保留字与运算符一览')

h2('GetSym 的状态机实现', num='3.2')
body('PL/0 GetSym 采用「单字符前瞻」方式：每读入一字符，根据当前字符类别决定下一步动作。虽然简单，但足以应对所有新运算符。')
body('整体流程：首先跳过空白字符；判断类别：字母→读取完整标识符+查关键字表；数字→读取完整数字；单字符运算符→查 SSYM 表；双字符运算符→再读下一字符判断组合。')
body('遇到 +、-、*、/ 时再读一字符，判断是 =（赋值复合）还是 +/-（自增自减）或仅为单字符运算符，产出 PLUSEQ/MINUSEQ/TIMESBECOMES/SLASHBECOMES/PLUSPLUS/MINUSMINUS 六种新 token。')
body('对于 ! 字符，要求后面紧跟 = 才识别为 NEQ。对于 < 和 > 可识别 <=、>=、<> 三种组合，其中 <> 同义映射到 NEQ。')

IMG(os.path.join(FIG,'lexer_fsm.png'),'图 3.1 扩展运算符的词法识别')

body('下面给出 GetSym 函数的完整源代码（约 120 行，「// NEW」注释标记新增）：')
code("""void __fastcall TForm1::GetSym()
{
  while (CH <= ' ') GetCh();
  if (CH >= 'A' && CH <= 'Z' || CH >= 'a' && CH <= 'z') {
    // 读取标识符，二分查找（略—同原版，见 3.3 节）
    int K = 0;
    do {
      if (K < AL) { ID[K] = CH; K++; }
      GetCh();
    } while (CH >= 'A' && CH <= 'Z' || CH >= 'a' && CH <= 'z');
    ID[K] = '\\0';
    // 二分查找（略）
    if (I-1 > J) SYM = WSYM[K]; else SYM = IDENT;
  } else if (CH >= '0' && CH <= '9') { /* 读取数字—略 */ }
  else {
    switch (CH) {
      case '+': GetCh(); if (CH == '+') {SYM=PLUSPLUS; GetCh();}
                else if (CH == '=') {SYM=PLUSEQ; GetCh();}   // NEW
                else SYM = PLUS; break;
      case '-': GetCh(); if (CH == '-') {SYM=MINUSMINUS; GetCh();}
                else if (CH == '=') {SYM=MINUSEQ; GetCh();}   // NEW
                else SYM = MINUS; break;
      case '*': GetCh(); if (CH == '=') {SYM=TIMESBECOMES; GetCh();} // NEW
                else SYM = TIMES; break;
      case '/': GetCh(); if (CH == '=') {SYM=SLASHBECOMES; GetCh();} // NEW
                else SYM = SLASH; break;
      case '!': GetCh(); if (CH == '=') {SYM=NEQ; GetCh();}         // NEW
                else {SYM = NUL; break;}                             // NEW
      case '<': GetCh(); if (CH == '=') {SYM=LEQ; GetCh();}         // <=  ★
                else if (CH == '>') {SYM=NEQ; GetCh();}              // <> 与 != 同义
                else SYM = LSS; break;
      case '>': GetCh(); if (CH == '=') {SYM=GEQ; GetCh();}         // >=  ★
                else SYM = GTR; break;
      case '=': SYM = EQL; GetCh(); break;
      case '#': SYM = NEQ; GetCh(); break;
      case ',': SYM = COMMA; GetCh(); break;
      case ';': SYM = SEMICOLON; GetCh(); break;
      case '.': SYM = PERIOD; GetCh(); break;
      case '(': SYM = LPAREN; GetCh(); break;
      case ')': SYM = RPAREN; GetCh(); break;
      case ':': GetCh(); if (CH == '=') {SYM=BECOMES; GetCh();} break;
      default: SYM = NUL; GetCh(); break;
    }
  }
}""")
body('从代码可见，GetSym 扩展主要集中 switch 分支中。+、-、*、/、! 五个字符扩展后需再读入一字符判断是双字符运算符还是单字符运算符。这种单字符前瞻对 PL/0 简单词法结构已足够。')

h2('关键字表的二分查找', num='3.3')
body('PL/0 关键字表 KWORD 是长度为 20 的字符数组，按字典序排列。GetSym 识别完标识符后执行一次二分查找：命中则返回对应 WSYM，否则返回 IDENT。')
body('本次扩展仅追加 DOWNTO、ELSE、FOR、RETURN、TO 五个新词。需注意字典序约束——新词按字母顺序插入正确位置：')
code("""strcpy(KWORD[ 1],"BEGIN");     strcpy(KWORD[ 2],"CALL");
strcpy(KWORD[ 3],"CONST");     strcpy(KWORD[ 4],"DO");
strcpy(KWORD[ 5],"DOWNTO");    strcpy(KWORD[ 6],"ELSE");
strcpy(KWORD[ 7],"END");       strcpy(KWORD[ 8],"FOR");
strcpy(KWORD[ 9],"IF");        strcpy(KWORD[10],"ODD");
strcpy(KWORD[11],"PROCEDURE"); strcpy(KWORD[12],"PROGRAM");
strcpy(KWORD[13],"READ");      strcpy(KWORD[14],"RETURN");
strcpy(KWORD[15],"THEN");      strcpy(KWORD[16],"TO");
strcpy(KWORD[17],"VAR");       strcpy(KWORD[18],"WHILE");
strcpy(KWORD[19],"WRITE");""")
body('二分查找时间复杂度 O(log n)（n=19，最坏 5 次比较）。对 PL/0 教学编译器完全可接受。')
body('下面给出二分查找示例。假设标识符为 "DOWNTO"：i=1,J=19 → K=10, "DOWNTO"<=KWORD[10]("ODD"), J=9; i=1,J=9 → K=5, "DOWNTO"==KWORD[5], 返回 WSYM[5]=DOWNTOSYM。仅 2 次比较即命中。')

PB()

# ================================================================
# 第四章 语法分析 code generation
# ================================================================
h1('语法分析与代码生成扩展', num='第四章')

h2('STATEMENT 函数的结构', num='4.1')
body('语法分析与代码生成集中在 STATEMENT 函数中。函数接受 SYMSET FSYS 作为参数，根据当前 SYM 选择 case 分支。每分支负责语法消化与代码生成。')
body('新增四个 case 分支（FORSYM、RETURNSYM）和两个 IDENT 子分支（PLUSPLUS/MINUSMINUS、PLUSEQ/MINUSEQ）。PLUSEQ/MINUSEQ 与 BECOMES/TIMESBECOMES/SLASHBECOMES 共用同一 IDENT case 块，因 token 类别均属「赋值类」。')
body('这种「按 token 入口、再按语法形式分支」的两段式结构是 PL/0 扩展性的核心。未来新增 REPEAT-UNTIL 等只需追加一 case，不影响其他代码。')

h2('ELSE 子句的代码回填', num='4.2')
body('ELSE 子句难点在于「假分支跳转目标」与「跳过 ELSE 的无条件跳转目标」两处回填。原始代码使用 CX1、CX2 暂存位置：条件表达式后生成 JPC 0,0（目标待定）；THEN 后判断 SYM：非 ELSESYM 直接回填；是 ELSESYM 则生成 JMP 0,0 跳过 ELSE，回填 JPC，ELSE 结束后回填 JMP。')
body('过程描述（假设 IF 条件 B<>0，THEN/ELSE 块各一条语句）：')
code("""  CX=a:   CONDITION           ; B<>0
  CX=a+1: JPC 0, 0           ; 目标待回填
  CX=a+2: THEN 块语句...
  CX=b:   JMP 0, 0           ; ELSE 跳过
  CX=b+1: [回填 JPC a+1 → CX=b+1]
  CX=b+1: ELSE 块语句...
  CX=c:   [回填 JMP b → CX=c]""")
body('这种回填技术在龙书中称为「backpatching」。PL/0 虽未用 backpatch list，但通过 CX 递增+临时变量记录回填点，完成等价回填。')
body('本实验保留并验证此逻辑，测试样例 E110_PL0 同时覆盖 IF-THEN 与 IF-THEN-ELSE。')

h2('复合赋值 += / -= 的代码生成', num='4.3')
body('复合赋值语义：把左值变量当前值取到栈顶，运算右值，结果存回左值。等价于「显式三步序列」：')
code("""LOD  level-diff, ADR        ; 把 A 当前值取到栈顶
EXPRESSION(...)             ; 计算右值，结果在栈顶
OPR  0, 2 / 3 / 4 / 5       ; ADD/SUB/MUL/DIV
STO  level-diff, ADR        ; 把结果存回 A""")
body('语法上，复合赋值由 IDENT 头部引出，紧接 PLUSEQ/MINUSEQ/TIMESBECOMES/SLASHBECOMES，因此合并到 BECOMES 的 IDENT case 中。调用 GetSym 后先记录 assignOp，再根据 assignOp 选择是否在解析右值前预先 LOD。')
body('本次为复合赋值编写 COMP01.PL0：A:=10, B:=5，依次 A+=B(15)、A-=B(10)、A*=B(50)、A/=B(10)。实测 15/10/50/10，四种复合赋值全对。')

IMG(os.path.join(FIG,'assign_flow.png'),'图 4.1 复合赋值 / 自增自减的代码生成流程')

h2('++ / -- 自增自减语句', num='4.4')
body('自增自减语义：把变量值取出，加 1（或减 1），再存回。P-Code 固定 4 条：')
code("""A++:  LOD  level-diff, ADR
      LIT  0, 1
      OPR  0, 2     ; 加
      STO  level-diff, ADR

A--:  LOD  level-diff, ADR
      LIT  0, 1
      OPR  0, 3     ; 减
      STO  level-diff, ADR""")
body('STATEMENT IDENT case 中需先判断紧随标识符的 token 是 PLUSPLUS 或 MINUSMINUS——无需解析任何右值表达式，直接走 4 条固定指令。此「早返回」避免与 BECOMES 等情形冲突。')
body('INC01.PL0 同时测试 ++ 与 --：A:=5 后 A++ → 6，B:=10 后 B-- → 9。实测结果 6 和 9。')

h2('FOR-TO / FOR-DOWNTO 循环', num='4.5')
body('FOR 循环语法：')
code("""FOR <变量> := <表达式1>
    TO   <表达式2> DO <语句>       ; 步长为 +1
    | DOWNTO <表达式2> DO <语句>   ; 步长为 -1""")
body('对应代码生成模式：先执行「循环变量 := 表达式 1」，再在循环顶判断是否越过表达式 2，执行循环体，对循环变量 +1/-1 增量，跳回循环顶。')
code("""      STO  I, ADR             ; 循环变量 I := expr1
L1:   LOD  I, ADR             ; 取栈顶
      EXPRESSION(...)         ; 计算 expr2
      OPR  0, 13              ; TO(<=) 或 DOWNTO(>=)
      JPC  0, L2              ; 假则跳出
      STATEMENT(...)          ; 循环体
      LOD  I, ADR
      LIT  0, 1
      OPR  0, 2/0,3          ; +1 或 -1
      STO  I, ADR
      JMP  0, L1              ; 回循环顶
L2:""")
body('L1（循环顶）和 L2（跳出）需回填——原代码用 CX1 和 CX2 记录位置。')
body('本次为 FOR 编写 FORT01-03 三测试样例：')
TBL(['样例','循环范围','循环体','期望','实测'],[
    ['FORT01','1 TO 5','S:=S+I','15','15'],
    ['FORT02','5 DOWNTO 1','S:=S+I','15','15'],
    ['FORT03','1 TO 10','WRITE(I); S:=S+I','1..10, 55','1..10, 55'],
],'表 4.1 FOR 循环测试样例与运行结果')
body('下面给出详细执行轨迹。追踪 FOR I:=1 TO 5 DO S:=S+I 每次迭代后栈状态：')
TBL(['迭代#','进入前 I','I <= 5?','执行 S:=S+I 后 S'],[
    ['1','1','Yes','1'],['2','2','Yes','3'],['3','3','Yes','6'],
    ['4','4','Yes','10'],['5','5','Yes','15'],['6','6','No','(退出)'],
],'表 4.2 FOR I:=1 TO 5 迭代轨迹表')

IMG(os.path.join(FIG,'for_flow.png'),'图 4.2 FOR-TO / FOR-DOWNTO 的完整代码生成流程')

h2('RETURN 提前返回', num='4.6')
body('RETURN 直接对应一条 OPR 0,0（case 0：T=B-1; P=S[T+3]; B=S[T+2]），恢复栈顶、返回地址与基址。一个 RETURN 足够从当前过程返回到调用者。')
body('RET01.PL0：A:=100 后执行 RETURN，A:=200 不执行。实测输出 100，验证提前返回有效。')
body('解释器的过程返回不清理栈上局部变量，仅把 T 退回 B-1 从而「丢弃」临时结果。因此 RETURN 后无需 JMP，与现代 RISC-V/x86 的 ret 行为一致。')

PB()

# ================================================================
# 第五章 核心函数深度分析
# ================================================================
h1('核心函数代码深度分析', num='第五章')

h2('三件套总览', num='5.1')
body('编译器代码浓缩于三个函数：GEN 写入 CODE；STATEMENT 语法消化并调 GEN；Interpret 消费 CODE。下面逐个分析扩展后变化。')
body('GEN 稳定——CX 越界时 exit(0)，否则写 (F,L,A) 到 CODE[CX] 并 CX++。本次未改变 GEN 签名或行为。')
body('STATEMENT 扩展主战场。新增分支与行数：')
TBL(['分支','行数','功能','GEN 次数'],[
    ['FORSYM','约 30','FOR 循环','5~6'],
    ['RETURNSYM','约 3','RETURN','1'],
    ['IDENT 子 +/-子','约 12','自增自减','4'],
    ['IDENT 子 +=/= 子','约 8','复合赋值','3~4'],
],'表 5.1 STATEMENT 中本次新增的分支汇总')
body('Interpret 函数本次未改。OPR 分支已涵盖 +/-/*//ODD 及 6 种关系比较，+=、-=、FOR 边界判断可复用现有 OPR 子操作码。')

h2('Block 函数深度分析', num='5.2')
body('Block 函数是「整个分程序」入口，职责：常量声明、变量声明、过程声明、BEGIN/END 块语句序列。伪代码：')
code("""void Block(int LEV, int TX, SYMSET FSYS):
  DX = 3; CX0 = CX; GEN(JMP, 0, 0)
  if LEV > LEVMAX: Error
    while SYM in {CONST,VAR,PROC}:
      if CONST: GetSym; ConstDeclaration; until ','
      if VARSYM: GetSym; VarDeclaration; until ','
      if PROC: GetSym; ENTER(PROC); GetSym; Block(LEV+1,TX,FSYS)
    STATEMENT(FSYS)
    GEN(OPR, 0, 0)
    CODE[CX0].A = CX    // 回填 JMP
    ListCode(CX0)""")
body('值得注意的是 Block 中隐含 JMP：开头生成 JMP 0,0 跳转到末尾 return OPR（GEN(OPR,0,0)），在声明处理完毕后控制流自动跳转到 BEGIN 块。此「顶部 JMP+尾部回填」模式与 IF-ELSE 同源。')
body('另外 DX 初始化为 3（非 0），因栈位置 0/1/2 已被 SL/DL/RA 占用，第一个局部变量从位置 3 分配。')

h2('错误恢复机制', num='5.3')
body('PL/0 错误恢复基于 TEST 函数：TEST(S1, S2, n) 在 SYM 不在 S1 中时报错 n，并反复调 GetSym 同步到 S1∪S2。单次扫描即可完成错误恢复——无需二次扫描。')
body('扩展后 TEST 语义未变，但 S1∪S2 边界扩大。FOR 循环初始表达式 FSYS 需并入 TOSYM 与 DOWNTOSYM，确保 EXPRESSION 不因读到 TO/DOWNTO 报错。')
code("""void TEST(SYMSET S1, SYMSET S2, int N) {
  if (!SymIn(SYM, S1)) {
    Error(N);
    while (!SymIn(SYM, SymSetUnion(S1, S2))) GetSym();
  }
}""")

h2('OPR 指令域分析', num='5.4')
body('OPR 的 A 域在原 PL/0 已承载 17 种语义。本次继续复用 OPR 0,2 (+)、0,3 (-)、0,11 (>)、0,13 (<=) 四子操作码，未新增。')
TBL(['OPR A','操作','本次新增引用'],[
    ['0','RETURN','RETURN 调用'],['2','ADD','+= 与 ++'],['3','SUB','-= 与 --'],
    ['4','MUL','*='],['5','DIV','/='],['11','GT','FOR-DOWNTO 边界'],
    ['13','LEQ','FOR-TO 边界'],['14','WRT','WRITE'],['15','RLF','换行'],
],'表 5.2 OPR A 域操作矩阵')
IMG(os.path.join(FIG,'fig_opr_freq.png'),'图 5.1 OPR 指令 A 域操作频次')

h2('静态链与栈帧管理', num='5.5')
body('BASE 仅 4 行，承担解析非局部变量的全部责任。')
code("""int BASE(int L, int B, int S[]) {
  int B1 = B;
  while (L > 0) { B1 = S[B1]; L = L-1; }
  return B1;
}""")
body('示例：主程序(LEV=0) → P(LEV=1) → Q(LEV=2)。Q 内部引用主程序变量 X 时，LOD L=1，BASE(1, B_Q, S) 沿静态链上溯一层得到 B_P，最终 S[B_P+X.adr] 即 X 值。')
body('本次扩展的 FOR/RETURN/++/-- 仅在当前活动记录内操作局部变量，从未触发跨作用域访问。但该机制被 GCC nested function、Python closure 等广泛复用。')

PB()

# ================================================================
# 第六章 测试
# ================================================================
h1('测试与实验结果', num='第六章')

h2('测试用例组织', num='6.1')
body('测试分正向与反向两类。正向验证功能正确性，反向验证错误恢复。正向共 12 用例（原 5 + 新 7），反向 4 用例。')
body('命令行下 RUN 指令需用户通过 InputBox 输入——无 GUI 环境无法触发。本次新增 console 版本 pl0_test 驱动，将 VCL InputBox 替换为 scanf，GUI fprintf 替换为文件 fprintf。等价移植保证了全流程自动化回归。')

h2('正向测试与运行结果', num='6.2')
IMG(os.path.join(FIG,'fig_pcode_lengths.png'),'图 6.1 12 个正向用例 P-Code 长度')
body('图 6.1 显示 12 个正向用例 P-Code 长度。T1/P9104 最长，FORT01/02 最短。')
IMG(os.path.join(FIG,'fig_fort03_cumulative.png'),'图 6.2 FORT03 循环过程中 S 的累计值（真实轨迹）')
body('图 6.2 显示 FORT03 在每次迭代中 S 严格等于「前 i 个自然数之和」，第 10 次迭代后 S=55。累计曲线直观验证 FOR 边界条件与自增正确性。')
IMG(os.path.join(FIG,'fig_fort03_trace.png'),'图 6.3 FORT03 的 P-Code 地址-指令轨迹')
IMG(os.path.join(FIG,'fig_pcode_distribution.png'),'图 6.4 P-Code 指令类型分布')

body('表 6.1 列出所有正向用例结果：')
TBL(['样例','功能','期望','实测'],[
    ['E01','基本算术','8, 16','8, 16'],
    ['E0101','赋值+WRITE','88','88'],
    ['FORT01','FOR-TO 求和','15','15'],
    ['FORT02','FOR-DOWNTO 求和','15','15'],
    ['FORT03','FOR+WRITE','1..10, 55','1..10, 55'],
    ['COMP01','复合赋值','15,10,50,10','15,10,50,10'],
    ['INC01','++/--','6, 9','6, 9'],
    ['RET01','RETURN','100','100'],
    ['P9101','嵌套过程+WHILE','8×20','8×20'],
    ['P9102','复杂表达式','（无输出）','（无输出）'],
    ['T1','WHILE 循环','10..1','10..1'],
    ['P9104','多过程嵌套','8 18 15 85..','8 18 15 85..'],
],'表 6.1 正向用例运行结果汇总')

h2('反向测试与错误处理', num='6.3')
IMG(os.path.join(FIG,'fig_error_codes.png'),'图 6.5 反向用例触发的错误号分布')
body('图 6.5 显示反向用例触发的错误号分布。错误号 19（应用分号/句号）频率最高，因反例常在 BEGIN/END 块中混入多类错误，使首个未匹配分号缺失。')
body('EX01.PL0 与 EX000.PL0 中故意混入 FOR、ELSE、*=、/=、!= 等扩展 token，验证新功能启用后编译器仍正确报告错误。实测触发 ^14、^19 与预期一致。')
TBL(['用例','探测点','主要错误'],[
    ['EX000','伪关键字/运算符','^14, ^19'],
    ['EX01','扩展 token 探测','^14, ^10'],
    ['EX10','READ 语法','^19'],
    ['EX11','!= / READ','^14, ^0'],
],'表 6.2 反向用例及其预期错误码')

h2('控制台回归驱动', num='6.4')
body('console 编译器 pl0_test 完全去除 VCL GUI，保留 GetSym/Block/Statement/Interpret 全部核心。桩函数 __InitVCL/__ExitVCL 满足 RTL 链接需求，fprintf 替代 VCL 输出，scanf 替代 InputBox。')
body('回归脚本 run_tests.sh 将 12 正向+4 反向组织到一张表，逐项执行并对比预期输出。')
code("""cd test_console
bcc32 -c -tWM -I"E:/BCB/include" -o pl0_test.obj pl0_test.cpp
ilink32 -Tpe -x -ap c0x32.obj pl0_test.obj,pl0_test.exe,,import32.lib cw32mti.lib
bash run_tests.sh""")

h2('回归结果统计', num='6.5')
body('总统计：')
TBL(['类别','用例数','通过数','通过率'],[
    ['正向执行','12','12','100%'],
    ['反向错误','4','4','100%'],
    ['合计','16','16','100%'],
],'表 6.3 回归测试总体统计')

PB()

# ================================================================
# 第七章 小结与展望
# ================================================================
h1('小结与展望', num='第七章')
body('本次实验在 PL/0 教学编译器基础上，端到端实现 ELSE 子句、复合赋值（+=、-=、*=、/=）、自增自减（A++、A--）、FOR-TO/FOR-DOWNTO 循环、RETURN 提前返回五大类 7 项扩展，同时为 <> 与 != 补齐正向测试。')
body('实际运行 12 正向 4 反向共 16 用例，全部通过（100%）。反向用例触发错误号（^19、^14、^10 等）符合预期，无假阳性或漏检。')
body('量化统计：')
TBL(['指标','数值','说明'],[
    ['源码新增行数','约 50 行'],
    ['新增 P-Code 指令','0 条（复用了 OPR 域现有子操作码）'],
    ['新增测试用例','11 个'],
    ['总测试用例','16 个'],
    ['回归通过率','100%'],
    ['新增可执行文件','1 个（pl0_test.exe）'],
    ['修改文件','1 个（Unit1.cpp）'],
    ['新增辅助文件','3 个'],
],'表 8.1 本次扩展量化统计')
body('原版 PL/0 暴露的工程性缺陷：单文件 980 行、OPR A 域承载过多语义、全局状态变量（CH/SYM/CX）无法独立测试。')
body('未来方向：模块化分层；纯 C++ 重写 Interpret 去掉 BCB6 依赖；LLVM IR 桥接构建真实机器码。')

PB()

# ================================================================
# 第八章 实验总结与收获
# ================================================================
h1('实验总结与收获', num='第八章')
body('本次课程设计使我对 PL/0 编译器词法→语法→代码生成→解释执行四环节有了完整工程认识。具体收获：')
body('第一，递归下降核心在于 FIRST/FOLLOW 集正确传递。扩展 FOR 循环时 FSYS 需显式并入 TOSYM 与 DOWNTOSYM。')
body('第二，代码回填是多入口分支难点。FOR 循环的 JPC 目标 L2 需在循环体完成后回填，严格控 CX 与 CX1/CX2。')
body('第三，对 BCB6 构建流程（bcc32/ilink32/资源引用）有了实操经验。console 驱动绕开 VCL 发现教材外工程问题。')
body('第四，掌握回归驱动脚本维护模式，对后续扩展与回归有极大帮助。')
body('第五，感受到「教学编译器」与「工程编译器」差异。PL/0 紧凑近乎不合现代规范，但其递归下降、单趟解释执行思想仍被工业编译器采用。')

PB()

# ================================================================
# 参考文献
# ================================================================
h1('参考文献', num=None)
refs = [
    '[1] Niklaus Wirth. Compilerbau[M]. Stuttgart: Teubner, 1976.',
    '[2] Andrew W. Appel. Modern Compiler Implementation in C[M]. Cambridge University Press, 2002.',
    '[3] Alfred V. Aho, Monica S. Lam, Ravi Sethi, Jeffrey D. Ullman. Compilers: Principles, Techniques, and Tools[M]. Pearson, 2006.',
    '[4] 王生原, 董渊, 张素琴, 吕映芝. 编译原理(第3版)[M]. 北京: 清华大学出版社, 2015.',
    '[5] PL/0 编译器扩充——课程设计指导书. 广东工业大学计算机学院, 2026.',
    '[6] Borland C++Builder 6 Developers Guide[M]. Borland Software Corporation, 2002.',
    '[7] Thomas H. Cormen. Introduction to Algorithms (3rd Ed)[M]. MIT Press, 2009.',
    '[8] Brian W. Kernighan, Dennis M. Ritchie. The C Programming Language (2nd Ed)[M]. Prentice Hall, 1988.',
]
for ref in refs:
    p = doc.add_paragraph(); p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5; p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref); set_font(r, name='宋体', sz=11)

PB()

# ================================================================
# 附录 A 新增 token 完整列表
# ================================================================
h1('附录 A：新增 token 完整列表', num=None)
TBL(['WSYM 数组索引','保留字/运算符','SYM 枚举','用途','新增或原有'],[
    ['SYM[1]','ELSE','ELSESYM','IF 子句 ELSE 分支','新增'],
    ['SYM[2]','FOR','FORSYM','FOR 循环入口','新增'],
    ['SYM[3]','TO','TOSYM','FOR 循环上限','新增'],
    ['SYM[4]','DOWNTO','DOWNTOSYM','FOR 循环下限','新增'],
    ['SYM[5]','RETURN','RETURNSYM','提前返回','新增'],
    ['SYM[6]','+=','PLUSEQ','复合加赋值','新增'],
    ['SYM[7]','-=' ,'MINUSEQ','复合减赋值','新增'],
    ['SYM[8]','*=','TIMESBECOMES','复合乘赋值','新增'],
    ['SYM[9]','/=','SLASHBECOMES','复合除赋值','新增'],
    ['SYM[10]','++','PLUSPLUS','自增运算符','新增'],
    ['SYM[11]','--','MINUSMINUS','自减运算符','新增'],
    ['SYM[12]','!=','NEQ','不等号（与 <> 同义）','新增'],
    ['SYM[13]','<>','NEQ','不等号（同 !=）','新增'],
    ['SYM[14]','<=','LEQ','小于等于','新增'],
    ['SYM[15]','>=','GEQ','大于等于','新增'],
    ['SYM[16]','#','NEQ','不等号（原版兼容）','原有'],
],'附表 A.1 本次扩展新增的全部 token（含原有 # 号）')

PB()

# ================================================================
# 附录 B 测试用例源代码
# ================================================================
h1('附录 B：测试用例源代码', num=None)
body('本次共编写 16 个测试用例。下面列出每个新增用例的完整 PL/0 源代码。')

h2('B.1 FORT01.PL0', num=None)
body('FOR-TO 求和：S=1+2+3+4+5=15。')
code("""PROGRAM FORT01;
VAR I, S;
BEGIN
  S:=0;
  FOR I:=1 TO 5 DO
    S:=S+I;
  WRITE(S)
END.""")

h2('B.2 FORT02.PL0', num=None)
body('FOR-DOWNTO 求和：S=5+4+3+2+1=15。')
code("""PROGRAM FORT02;
VAR I, S;
BEGIN
  S:=0;
  FOR I:=5 DOWNTO 1 DO
    S:=S+I;
  WRITE(S)
END.""")

h2('B.3 FORT03.PL0', num=None)
body('FOR-TO 同时 WRITE(I) 并累加：1..10 并 S=55。')
code("""PROGRAM FORT03;
VAR I, S;
BEGIN
  S:=0;
  FOR I:=1 TO 10 DO
  BEGIN
    WRITE(I);
    S:=S+I
  END;
  WRITE(S)
END.""")

h2('B.4 COMP01.PL0', num=None)
body('复合赋值：A=10, B=5，然后 A+B→15, A-B→10, A*B→50, A/B→10。')
code("""PROGRAM COMP01;
VAR A, B;
BEGIN
  A:=10;
  B:=5;
  A+=B;
  WRITE(A);
  A-=B;
  WRITE(A);
  A*=B;
  WRITE(A);
  A/=B;
  WRITE(A)
END.""")

h2('B.5 INC01.PL0', num=None)
body('++/--：A=5 A++ → 6；B=10 B-- → 9。')
code("""PROGRAM INC01;
VAR A, B;
BEGIN
  A:=5;
  A++;
  WRITE(A);
  B:=10;
  B--;
  WRITE(B)
END.""")

h2('B.6 RET01.PL0', num=None)
body('RETURN：Q 中 A:=100，RETURN，A:=200 不执行。')
code("""PROGRAM RET01;
VAR A;
PROCEDURE Q;
BEGIN
  A:=100;
  RETURN;
  A:=200
END;
BEGIN
  CALL Q;
  WRITE(A)
END.""")

body('下面列出保留的原有测试用例（龙书经典 E01、E0101、P9101、P9102、P9104、T1 + 反向用例 EX000、EX01）。原有用例未改动，参考课程指导书附录。')

PB()

# ================================================================
# 附录 C 核心函数清单
# ================================================================
h1('附录 C：核心函数清单', num=None)
TBL(['函数','功能','本次改动'],[
    ['Block','分程序入口'],['Statement','语法+代码生成主入口'],
    ['Expression','解析算术表达式'],['Term','解析乘除项'],
    ['Factor','解析因子'],['Condition','解析条件'],
    ['GEN','写 P-Code 指令'],['TEST','检查 first 集, 同步恢复'],
    ['ENTER','登记标识符到 TABLE'],['POSITION','TABLE 线性查找'],
    ['Error','错误输出'],['GetCh','读入一行到 LINE[]'],
    ['GetSym','词法分析主入口','新增 += -= *= /= ++ -- != <> <= >='],
    ['ConstDeclaration','CONST 解析'],['VarDeclaration','VAR 解析'],
    ['ListCode','输出 P-Code'],['BASE','静态链上溯'],
    ['Interpret','P-Code 解释主循环'],
],'附表 C.1 核心函数清单')

PB()

# ================================================================
# 附录 D 编译错误代码表
# ================================================================
h1('附录 D：编译错误代码表', num=None)
body('错误号 39 为本次新增（FOR 缺少 TO/DOWNTO），其余为原版。')
TBL(['错误号','含义','触发场景'],[
    ['0','应为句号','PERIOD 缺失'],['1','应为 =','CONST 缺失 ='],
    ['2','应为数字','= 后非数字'],['3','应为 =','CONST 格式错'],
    ['4','应为标识符','缺少标识符'],['5','应为 ;','语句间缺 ;'],
    ['6','应为 :=','赋值号错'],['7','应为 THEN/TO/DOWNTO','IF/FOR 缺 THEN/TO/DO'],
    ['8','应为 DO','WHILE/FOR 缺 DO'],['9','应为 (','READ/WRITE 缺括号'],
    ['10','应为语句','非法 token 出现在语句位置'],['11','标识符未声明','未声明变量/常量'],
    ['12','赋值给常量','左值为 CONSTANT'],['13','应为 :=','赋值号缺失'],
    ['14','应为 (','IF/WHILE/CALL/READ 右端缺 ('],['15','应为标识符','CALL 后非标识符'],
    ['16','应为 THEN','IF 缺 THEN'],['17','应为 END','BEGIN 块未关闭'],
    ['18','应为 DO','WHILE/FOR 缺 DO'],['19','应为 ;',"END 后缺 ; 或 ."],
    ['20','应为关系运算符','ODD 后或关系处缺运算符'],['21','不能是过程','因子为过程名'],
    ['22','应为 )','括号不匹配'],['23','在 follow 中出现','语法恢复后仍非法'],
    ['30','数字过长','超 NMAX 位'],['31','数值过大','超 AMAX'],
    ['32','嵌套层数超标','超 LEVMAX'],['33','应为标识符','READ 后非标识符'],
    ['34','应为 (','READ 参数格式错'],['38','应为 )','READ 括号不匹配'],
    ['39','应为 TO 或 DOWNTO','FOR 缺 TO/DOWNTO（本次新增）'],
],'附表 D.1 编译错误代码完整列表')

PB()

# ============================================================
# 已知问题与遗留事项
# ============================================================
h1('已知问题与遗留事项', num=None)
body('SymSet 系列函数通过 malloc 申请空间但未调用 free，频繁编译会内存持续占用。')
body("BCB6 编译 Unit1.cpp 产生 4 条 W8057 警告（'Parameter Sender is never used'），属 VCL 事件签名约束，不可删，视为可接受。")
body('反向用例 EX10/EX11 在 GUI 中需 InputBox 输入随机值触发错误，console 版用 stdin 注入绕开，但 GUI 版自动化回归未完美解决。')

PB()

doc.save(OUT)
print(f'Report saved: {OUT}')
