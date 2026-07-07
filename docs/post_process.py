# -*- coding: utf-8 -*-
"""Post-process: add section break before Chapter 1, isolate footer with PAGE."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

DOC = r'C:\Users\xwmsu\Desktop\c_builder_b\PL0编译器扩充实验报告.docx'

doc = Document(DOC)

target = None
for p in doc.paragraphs:
    # Skip TOC entries (they have tab + page number)
    if '\t' in p.text:
        continue
    if p.text.strip() == '第一章  绪论' or p.text.strip() == '绪论':
        target = p
        break

if target is None:
    print('ERROR: Could not find Chapter 1 heading')
    exit(1)

print(f'Found chapter 1 at: {target.text}')

# Insert section break before Chapter 1
# Add a new section with start_type = NEW_PAGE
new_sect = OxmlElement('w:sectPr')
new_sect.set(qn('w:type'), 'newPage')

# Append sectPr to the paragraph's pPr
pPr = target._element.get_or_add_pPr()
pPr.append(new_sect)

print('Added section break before Chapter 1')

# Now unlink footers between sections
sections = doc.sections
print(f'Number of sections: {len(sections)}')

if len(sections) >= 2:
    sec0 = sections[0]
    sec1 = sections[1]
    
    # Unlink and clear section 0 footer
    sec0.footer.is_linked_to_previous = False
    for p in sec0.footer.paragraphs:
        for r in p.runs:
            r._element.getparent().remove(r._element)
    
    # Unlink section 1 footer and ensure PAGE field exists
    sec1.footer.is_linked_to_previous = False
    footer = sec1.footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = 2  # right
    
    # Check if PAGE field already exists
    has_page = False
    for r in p.runs:
        if 'PAGE' in r._element.xml:
            has_page = True
            break
    
    if not has_page:
        r = p.add_run()
        f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
        it.text = 'PAGE   \\* MERGEFORMAT'
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
        r._r.append(f1); r._r.append(it); r._r.append(f2)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(9)
    
    print('Section 0 footer cleared, Section 1 footer has PAGE field')

doc.save(DOC)
print('Saved')
