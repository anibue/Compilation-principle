from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "PL0仓库说明.docx"


def set_page_layout(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_title(document: Document) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PL/0 编译器仓库说明")
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "Microsoft YaHei"

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("基于 Borland C++Builder 6 VCL 的教学型编译器项目概览")
    run.italic = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(95, 95, 95)
    run.font.name = "Microsoft YaHei"


def add_heading(document: Document, text: str, level: int = 1) -> None:
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = RGBColor(31, 78, 121) if level == 1 else RGBColor(55, 55, 55)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)


def add_body(document: Document, text: str) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10.5)


def add_bullets(document: Document, items) -> None:
    for item in items:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(10.5)


def add_structure_blocks(document: Document) -> None:
    blocks = [
        ("./（根目录）", "PL01.bpr", "C++Builder 6 工程文件，位于根目录，路径相对于此文件解析。"),
        ("src/", "PL01.cpp、Unit1.cpp", "程序入口、VCL 主窗体、词法分析、语法分析、代码生成、解释执行都集中在这里。"),
        ("project/", "PL01.res", "资源文件及编译输出目录。"),
        ("test_cases/", "EX01.PL0、EX01.COD", "测试输入与输出样例，既包含正常样例，也包含错误触发样例。"),
        ("pascal_reference/", "PL0.PAS", "Pascal 参考实现，可用于对照当前 C++Builder 版本。"),
        ("docs/", "本说明文档", "存放仓库说明和实验文档相关材料。"),
    ]

    for folder, files, desc in blocks:
        p = document.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(f"{folder}  ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = "Microsoft YaHei"
        run.font.color.rgb = RGBColor(31, 78, 121)

        run = p.add_run(f"代表文件：{files}。")
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.name = "Microsoft YaHei"

        run = p.add_run(desc)
        run.font.size = Pt(10.5)
        run.font.name = "Microsoft YaHei"


def add_issue_box(document: Document) -> None:
    p = document.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("需要重点强调的问题")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(156, 0, 6)
    run.font.name = "Microsoft YaHei"

    issues = [
        "项目并不是完整模块化架构，编译器的大部分核心逻辑都堆叠在 Unit1.cpp 中，维护成本较高。",
        "FOR 虽然已经被识别为保留字，但在 STATEMENT 中没有对应实现，因此当前并不支持真正的 FOR 语句。",
        'test_cases 中不仅有正确样例，也有故意触发语法错误的输入；使用这些文件时需要区分"回归测试"和"错误测试"。',
    ]

    for issue in issues:
        p = document.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(issue)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(156, 0, 6)
        run.font.name = "Microsoft YaHei"


def build_doc() -> None:
    document = Document()
    set_page_layout(document)
    add_title(document)

    add_heading(document, "一、仓库定位")
    add_body(
        document,
        "这个仓库是一个基于 Borland C++Builder 6 与 VCL 图形界面的 PL/0 编译器实验项目。"
        "它的目标不是提供现代化命令行工具，而是通过一个可视化窗体，把词法分析、语法分析、P-Code 生成和解释执行串起来，方便教学演示与课程实验。"
    )

    add_heading(document, "二、目录与文件结构")
    add_structure_blocks(document)

    add_heading(document, "三、程序运行方式")
    add_bullets(
        document,
        [
            "使用 C++Builder 6 打开根目录下的 PL01.bpr，或使用命令行 bcc32 + ilink32 编译。",
            "编译后生成 project/PL01.exe 图形界面程序。",
            "在界面中输入测试文件名，例如 EX10，然后点击运行。",
            "程序会优先读取当前目录下的 <name>.PL0；如果不存在，再尝试 test_cases\\<name>.PL0。",
            "编译输出和运行结果会写入 Memo，同时生成对应的 .COD 文件。",
        ],
    )

    add_heading(document, "四、核心实现逻辑")
    add_bullets(
        document,
        [
            "PL01.cpp 只负责 VCL 应用初始化与创建主窗体。",
            "Unit1.h 定义了主窗体控件和输出辅助函数。",
            "Unit1.cpp 是整个编译器的核心，包含词法分析器 GetSym、递归下降语法分析器、符号表 TABLE、目标代码数组 CODE 和解释执行器 Interpret。",
            "ButtonRunClick 是整个流程的入口，负责初始化关键字表、符号集合、输入输出文件以及启动编译。",
        ],
    )

    add_heading(document, "五、语言与扩展特性")
    add_bullets(
        document,
        [
            "基础能力包括 PROGRAM、CONST、VAR、PROCEDURE、BEGIN/END、IF/THEN、WHILE/DO、CALL、READ、WRITE。",
            "表达式支持加减乘除、括号、常量和变量引用。",
            "条件判断支持 =、<、<=、>、>=、ODD。",
            "当前扩展了 ELSE、复合赋值 *= 与 /=，并支持 <> 与 != 两种不等号写法。",
            "虽然词法层面加入了 FOR，但语法与代码生成层面并未完成。",
        ],
    )

    add_heading(document, "六、测试材料的意义")
    add_body(
        document,
        'test_cases 目录不是单纯的"能运行样例"集合。部分文件用于验证正常编译和解释执行，部分文件则有意包含未完成语法或错误输入，用来观察报错行为。'
        "例如 EX01.PL0 就混入了 FOR、ELSE、*=、/=、!= 等探测性内容，对应输出会进入错误分支。"
    )

    add_heading(document, "七、整体评价")
    add_body(
        document,
        "从工程角度看，这是一个典型的课程实验仓库：规模不大，逻辑集中，便于直接阅读主文件理解编译过程。"
        "支持 IDE 和命令行两种构建方式，已修复中文编码问题。"
    )

    add_heading(document, "八、需要重点强调的问题")
    add_issue_box(document)

    document.save(str(OUT_PATH))


if __name__ == "__main__":
    build_doc()
