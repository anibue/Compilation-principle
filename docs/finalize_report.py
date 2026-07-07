# -*- coding: utf-8 -*-
"""Finalize report: fix pgNumType, page numbers."""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

DOC = r'C:\Users\xwmsu\Desktop\c_builder_b\PL0编译器扩充实验报告.docx'

doc = Document(DOC)
sections = doc.sections

print(f'Number of sections: {len(sections)}')

# Section 0 = Cover + TOC (no page numbers)
# Section 1 = Chapter 1 onwards (page numbers start from 1)

if len(sections) >= 2:
    # Add pgNumType to section 1, starting from 1
    sec1 = sections[1]
    sectPr = sec1._sectPr
    
    # Remove existing pgNumType if any
    existing = sectPr.find(qn('w:pgNumType'))
    if existing is not None:
        sectPr.remove(existing)
    
    # Add pgNumType with start=1
    pgNumType = OxmlElement('w:pgNumType')
    pgNumType.set(qn('w:start'), '1')
    sectPr.append(pgNumType)
    print('Added pgNumType start=1 to section 1')
    
    # Ensure section 1 footer has PAGE field
    footer = sec1.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    has_page = any('PAGE' in r._element.xml for r in p.runs)
    if not has_page:
        r = p.add_run()
        f1 = OxmlElement('w:fldChar'); f1.set(qn('w:fldCharType'), 'begin')
        it = OxmlElement('w:instrText'); it.set(qn('xml:space'), 'preserve')
        it.text = 'PAGE   \\* MERGEFORMAT'
        f2 = OxmlElement('w:fldChar'); f2.set(qn('w:fldCharType'), 'end')
        r._r.append(f1); r._r.append(it); r._r.append(f2)
        r.font.name = 'Times New Roman'; r.font.size = Pt(9)
        print('Added PAGE field to section 1 footer')
    else:
        print('PAGE field already exists in section 1 footer')
    
    # Ensure section 0 has no footer content
    sec0 = sections[0]
    sec0.footer.is_linked_to_previous = False
    for p in sec0.footer.paragraphs:
        for r in p.runs:
            r._element.getparent().remove(r._element)
    print('Cleared section 0 footer')

doc.save(DOC)
print()
print('=== Final report statistics ===')

# Analyze final report
doc = Document(DOC)
non_empty = [p.text for p in doc.paragraphs if p.text.strip()]
short = [p for p in non_empty if len(p) < 40]
mid = [p for p in non_empty if 40 <= len(p) < 200]
long = [p for p in non_empty if len(p) >= 200]
lines = len(short) * 2 + len(mid) * 2 + len(long) * 3.5
pages_text = lines / 36
pages_tbl = len(doc.tables) * 0.6
pages_img = len(doc.inline_shapes) * 0.7
total_pages = pages_text + pages_tbl + pages_img

print(f'Non-empty paragraphs: {len(non_empty)}')
print(f'Tables: {len(doc.tables)}')
print(f'Inline images: {len(doc.inline_shapes)}')
print(f'Total characters: {sum(len(t) for t in non_empty)}')
print(f'Sections: {len(doc.sections)}')
print(f'Estimated pages: ~{total_pages:.0f}')
print()
print('Done!')
